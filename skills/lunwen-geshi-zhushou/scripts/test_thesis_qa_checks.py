#!/usr/bin/env python3
"""Contract checks for reusable thesis structural QA diagnostics."""

from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Pt

import thesis_format_from_sample as fmt
import thesis_qa_checks as qa


def assert_equal(actual, expected, message: str) -> None:
    if actual != expected:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def assert_true(value, message: str) -> None:
    if not value:
        raise AssertionError(message)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.5


def make_profile() -> dict:
    profile = fmt.default_profile()
    profile["enabled_roles"] = [
        "table_header",
        "table_body_variable",
        "table_body_value",
    ]
    return profile


def make_captioned_table_doc() -> Document:
    doc = Document()
    doc.add_paragraph("Chapter 5 Empirical Process and Analysis")
    doc.add_paragraph("Table 5-1: Descriptive Statistics")
    table = doc.add_table(rows=2, cols=3)
    values = [
        ["", "count", "mean"],
        ["lnTC", "403", "10.1926"],
    ]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_text(cell, values[row_idx][col_idx])
    return doc


def main() -> int:
    profile = make_profile()
    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        unformatted_path = temp_path / "unformatted.docx"
        unformatted = make_captioned_table_doc()
        unformatted.save(unformatted_path)

        report = qa.run_structural_checks(unformatted_path, profile)
        assert_equal(report["status"], "fail", "unformatted captioned table fails QA")
        table_check = report["checks"]["captioned_tables"]
        assert_equal(table_check["captioned_table_count"], 1, "captioned table counted")
        issue = report["issues"][0]
        assert_equal(issue["code"], "captioned_table_cell_format_mismatch", "mismatch code")
        assert_equal(issue["table_index"], 1, "human-readable table index")
        assert_equal(issue["row"], 1, "human-readable row index")
        assert_equal(issue["column"], 2, "human-readable column index")
        assert_equal(issue["text"], "count", "failing cell text")
        fields = {difference["field"] for difference in issue["differences"]}
        assert_true("size_pt" in fields, "cell mismatch includes point size")
        assert_equal(issue["actual_signature"]["size_pt"], 12.0, "actual signature is preserved")
        assert_equal(issue["expected_signature"]["size_pt"], 9.0, "expected signature is preserved")

        formatted_path = temp_path / "formatted.docx"
        formatted = make_captioned_table_doc()
        fmt.format_tables(
            formatted,
            profile,
            skip_first_table=True,
            comment_tracker=fmt.CommentTracker(mode="none"),
            protected_table_elements=set(),
        )
        formatted.save(formatted_path)
        formatted_report = qa.run_structural_checks(formatted_path, profile)
        assert_equal(formatted_report["status"], "pass", "formatted captioned table passes QA")
        assert_equal(formatted_report["issues"], [], "formatted output has no mismatches")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
