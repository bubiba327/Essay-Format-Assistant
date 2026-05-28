#!/usr/bin/env python3
"""Contract checks for the thesis evaluation harness."""

from __future__ import annotations

import json
import subprocess
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Pt

import evaluate_thesis_skill as evaluator
import thesis_format_from_sample as fmt


SKILL_DIR = Path(__file__).resolve().parent.parent


def assert_equal(actual, expected, message: str) -> None:
    if actual != expected:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def assert_true(value, message: str) -> None:
    if not value:
        raise AssertionError(message)


def write_unformatted_captioned_table(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Table 5-1: Descriptive Statistics")
    table = doc.add_table(rows=2, cols=2)
    for row_index, values in enumerate((["", "count"], ["lnTC", "403"])):
        for column_index, text in enumerate(values):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.size = Pt(12)
            paragraph.paragraph_format.line_spacing = 1.5
    doc.save(path)


def main() -> int:
    manifest = SKILL_DIR / "fixtures" / "manifest.json"
    cases = evaluator.load_cases(manifest)
    assert_true("table-format1-sample2" in cases, "initial case is registered")
    assert_equal(
        evaluator.resolve_fixture_path("samples/format.doc", manifest),
        (manifest.parent / "samples" / "format.doc").resolve(),
        "relative packaged fixture path",
    )

    with TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        sample = temp_root / "sample.docx"
        target = temp_root / "target.docx"
        Document().save(sample)
        Document().save(target)
        raw_case = dict(cases["table-format1-sample2"])
        raw_case["sample"] = str(sample)
        raw_case["target"] = str(target)
        case = evaluator.resolve_case_paths(raw_case, manifest)
        layout = evaluator.case_layout(
            temp_root, "table-format1-sample2", case["sample"], case["target"]
        )
        analyze_cmd, apply_cmd = evaluator.build_formatter_commands(case, layout, "fast")
        assert_true("--analyze-only" in analyze_cmd, "stage 1 analyzes only")
        assert_true("--format-table" in apply_cmd, "stage 2 consumes the generated table")
        assert_true("--reuse-profile" in apply_cmd, "stage 2 reuses the learned profile")
        assert_true("--apply-engine" in apply_cmd and "direct" in apply_cmd, "case pins direct engine")
        isolated_cache = Path(temp_dir) / "isolated-shared-cache"
        cached_analyze_cmd, cached_apply_cmd = evaluator.build_formatter_commands(
            case, layout, "fast", shared_sample_cache_dir=isolated_cache
        )
        for command in (cached_analyze_cmd, cached_apply_cmd):
            assert_true("--shared-sample-cache-dir" in command, "evaluation can isolate shared sample cache")
            assert_true(str(isolated_cache) in command, "isolated cache path reaches formatter")
        uncached_analyze_cmd, uncached_apply_cmd = evaluator.build_formatter_commands(
            case, layout, "fast", no_shared_sample_cache=True
        )
        for command in (uncached_analyze_cmd, uncached_apply_cmd):
            assert_true("--no-shared-sample-cache" in command, "evaluation can disable shared sample cache")
        assert_equal(layout["source_docx"].name, "target.normalized.docx", "layout locates pre-format DOCX")
        assert_equal(
            layout["analyze_performance_json"].name,
            "performance-analyze.json",
            "layout locates Stage 1 timing report",
        )
        assert_equal(
            layout["apply_performance_json"].name,
            "performance-apply.json",
            "layout locates Stage 2 timing report",
        )

        layout["analysis_dir"].mkdir(parents=True, exist_ok=True)
        output_risk_dir = layout["visual_qa_dir"] / f"{layout['output'].stem}_output"
        output_risk_dir.mkdir(parents=True, exist_ok=True)
        source_doc = Document()
        fmt.add_page_number_field(source_doc.sections[0].footer.paragraphs[0])
        source_doc.save(layout["source_docx"])
        clean_doc = Document()
        fmt.add_page_number_field(clean_doc.sections[0].footer.paragraphs[0])
        clean_doc.save(layout["output"])
        profile = fmt.default_profile()
        profile["document_language"] = {"dominant": "english-dominant"}
        profile["enabled_roles"] = list(case["expected"]["required_roles"])
        layout["profile_json"].write_text(
            json.dumps(profile),
            encoding="utf-8",
        )
        layout["lint_json"].write_text(
            json.dumps({"issues": [{"severity": "warning", "code": "mixed_page_number_examples"}]}),
            encoding="utf-8",
        )
        layout["output_risk_json"].write_text(
            json.dumps({"qa_level": "fast", "blank_like_pages": [], "edge_overflow_pages": []}),
            encoding="utf-8",
        )
        layout["analyze_performance_json"].write_text(
            json.dumps({"elapsed_seconds": 8.25, "metrics": {"pdf_conversion_passes": 0}}),
            encoding="utf-8",
        )
        layout["apply_performance_json"].write_text(
            json.dumps(
                {
                    "elapsed_seconds": 13.5,
                    "metrics": {
                        "pdf_conversion_passes": 2,
                        "png_raster_passes": 1,
                        "toc_refreshed_entries": 57,
                    },
                }
            ),
            encoding="utf-8",
        )
        stages = [
            subprocess.CompletedProcess(["analyze"], 0, "", ""),
            subprocess.CompletedProcess(["apply"], 0, "", ""),
        ]
        result = evaluator.assess_case("table-format1-sample2", case, layout, "fast", stages)
        assert_equal(result["status"], "pass", "accepted artifacts pass")
        assert_equal(result["missing_roles"], [], "required roles observed")
        assert_equal(result["unexpected_lint_codes"], [], "allowed lint warning accepted")
        assert_equal(result["visual_check"], "skipped-fast", "fast mode is not presented as rendered QA")
        assert_equal(result["structural_status"], "pass", "clean output passes structural QA")
        assert_equal(result["structural_issues"], [], "clean output has no structural mismatches")
        assert_true(layout["structural_report_json"].exists(), "structural report artifact is written")
        assert_equal(
            result["structural_report"]["checks"]["page_numbers"]["source_page_number_count"],
            1,
            "source page-number presence reaches structural QA",
        )
        assert_equal(result["performance"]["stage1_seconds"], 8.25, "Stage 1 timing reaches result")
        assert_equal(result["performance"]["stage2_seconds"], 13.5, "Stage 2 timing reaches result")
        assert_equal(result["performance"]["total_seconds"], 21.75, "combined timing reaches result")
        assert_equal(result["performance"]["toc_refreshed_entries"], 57, "TOC refresh timing detail reaches result")

        write_unformatted_captioned_table(layout["output"])
        result = evaluator.assess_case("table-format1-sample2", case, layout, "fast", stages)
        assert_equal(result["status"], "fail", "unformatted captioned table fails evaluation")
        issue = result["structural_issues"][0]
        assert_equal(issue["table_index"], 1, "failure reports table location")
        assert_equal(issue["row"], 1, "failure reports row location")
        assert_equal(issue["column"], 2, "failure reports column location")
        report_paths = evaluator.write_reports(Path(temp_dir) / "failed-report", [result])
        markdown = report_paths["markdown"].read_text(encoding="utf-8")
        assert_true("table 1, row 1, column 2" in markdown, "failure location reaches report")
        assert_true("size_pt" in markdown, "failure format detail reaches report")

        clean_doc.save(layout["output"])
        layout["output_risk_json"].write_text(
            json.dumps({"qa_level": "review", "blank_like_pages": ["4"], "edge_overflow_pages": ["7"]}),
            encoding="utf-8",
        )
        result = evaluator.assess_case("table-format1-sample2", case, layout, "review", stages)
        assert_equal(result["status"], "pass", "review records visual risks without hard failure")
        assert_true(result["warnings"], "review visual risks are reported as warnings")
        result = evaluator.assess_case("table-format1-sample2", case, layout, "strict", stages)
        assert_equal(result["status"], "fail", "strict edge overflow is a hard failure")

        layout["profile_json"].write_text(
            json.dumps(
                {"document_language": {"dominant": "english-dominant"}, "enabled_roles": ["toc1"]}
            ),
            encoding="utf-8",
        )
        result = evaluator.assess_case("table-format1-sample2", case, layout, "fast", stages)
        assert_equal(result["status"], "fail", "missing learned roles fail")
        assert_true("table_header" in result["missing_roles"], "table-header omission is visible")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
