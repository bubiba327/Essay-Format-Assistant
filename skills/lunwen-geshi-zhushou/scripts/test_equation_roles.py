#!/usr/bin/env python3
"""Small regression checks for equation role detection and formatting."""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

import thesis_format_from_sample as fmt


def assert_equal(actual, expected, message: str) -> None:
    if actual != expected:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def assert_close(actual: float | None, expected: float, message: str) -> None:
    if actual is None or abs(actual - expected) > 0.01:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def build_math_paragraph():
    paragraph = Document().add_paragraph()
    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    omath.append(math_run)
    paragraph._p.append(omath)
    return paragraph


def main() -> int:
    assert_equal(fmt.role_for_target("(1)", "", False, False), "equation_number", "Arabic equation number")
    assert_equal(fmt.role_for_target("（3-1）", "", False, False), "equation_number", "Chinese punctuation equation number")
    assert_equal(
        fmt.role_for_target("share_ikt=X_ikt/X_it                                                              (2)", "", False, False),
        "equation_text_fallback",
        "Text fallback equation with trailing number",
    )
    assert_equal(
        fmt.role_for_target("edu_it = Σ (P_j,it × W_j) / P_it                                      (6)", "", False, False),
        "equation_text_fallback",
        "Symbol equation with trailing number",
    )

    math_paragraph = build_math_paragraph()
    assert_equal(fmt.equation_role_for_paragraph(math_paragraph, math_paragraph.text), "equation", "OMML equation")

    profile = fmt.default_profile()
    paragraph = Document().add_paragraph("share_ikt=X_ikt/X_it                                                              (2)")
    paragraph.paragraph_format.first_line_indent = Pt(24)
    fmt.format_equation_paragraph(paragraph, profile["roles"]["equation_text_fallback"], profile, "equation_text_fallback")
    assert_close(fmt.pt(paragraph.paragraph_format.first_line_indent), 0.0, "Equation fallback first-line indent")
    assert_equal(fmt.align_name(paragraph.alignment), "center", "Equation fallback alignment")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
