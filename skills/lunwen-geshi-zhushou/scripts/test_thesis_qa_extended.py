#!/usr/bin/env python3
"""Structural QA contracts beyond captioned-table formatting."""

from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import thesis_format_from_sample as fmt
import thesis_qa_checks as qa


def assert_equal(actual, expected, message: str) -> None:
    if actual != expected:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def assert_true(value, message: str) -> None:
    if not value:
        raise AssertionError(message)


def issue_for(report: dict, code: str) -> dict:
    for issue in report["issues"]:
        if issue["code"] == code:
            return issue
    raise AssertionError(f"Missing structural issue code: {code}; found {report['issues']!r}")


def make_profile() -> dict:
    profile = fmt.default_profile()
    profile["enabled_roles"] = [
        "toc1",
        "toc2",
        "toc3",
        "second",
        "figure_caption",
        "equation",
        "equation_text_fallback",
        "page_number",
    ]
    profile["header_footer"]["preserve_target_page_numbers"] = True
    return profile


def make_toc_paragraph(style_id: str, text: str):
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), style_id)
    properties.append(style)
    paragraph.append(properties)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    paragraph.append(run)
    return paragraph


def add_bad_toc(doc: Document) -> None:
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), "Table of Contents")
    properties.append(gallery)
    content = OxmlElement("w:sdtContent")
    content.append(make_toc_paragraph("TOC2", "1.1.1 Research Background 6"))
    control.append(properties)
    control.append(content)
    doc._element.body.insert(0, control)


def add_toc_entries(doc: Document, entries: list[tuple[str, str]]) -> None:
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), "Table of Contents")
    properties.append(gallery)
    content = OxmlElement("w:sdtContent")
    for style_id, text in entries:
        content.append(make_toc_paragraph(style_id, text))
    control.append(properties)
    control.append(content)
    doc._element.body.insert(0, control)


def add_math_paragraph(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x=1"
    math_run.append(math_text)
    omath.append(math_run)
    paragraph._p.append(omath)


def make_source_doc() -> Document:
    doc = Document()
    doc.add_paragraph("毕业论文（设计）")
    doc.add_paragraph("学生姓名：测试学生")
    declaration = doc.add_paragraph("原创性声明")
    declaration.runs[0].font.size = Pt(10)
    cover_table = doc.add_table(rows=1, cols=1)
    cover_cell = cover_table.cell(0, 0).paragraphs[0]
    cover_cell.text = "封面字段"
    cover_cell.runs[0].font.size = Pt(10)
    doc.add_paragraph("Abstract")
    add_math_paragraph(doc)
    fmt.add_page_number_field(doc.sections[0].footer.paragraphs[0])
    return doc


def make_changed_output_doc() -> Document:
    doc = Document()
    doc.add_paragraph("毕业论文（设计）")
    doc.add_paragraph("学生姓名：测试学生")
    declaration = doc.add_paragraph("原创性声明")
    declaration.runs[0].font.size = Pt(18)
    cover_table = doc.add_table(rows=1, cols=1)
    cover_cell = cover_table.cell(0, 0).paragraphs[0]
    cover_cell.text = "封面字段"
    cover_cell.runs[0].font.size = Pt(18)
    doc.add_paragraph("Abstract")
    return doc


def main() -> int:
    profile = make_profile()
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)

        toc_doc = Document()
        add_bad_toc(toc_doc)
        toc_path = root / "bad-toc.docx"
        toc_doc.save(toc_path)
        toc_issue = issue_for(qa.run_structural_checks(toc_path, profile), "toc_style_mismatch")
        assert_equal(toc_issue["role"], "toc3", "TOC issue identifies intended level")

        stale_toc = Document()
        add_toc_entries(stale_toc, [("TOC1", "Chapter 1 Introduction6")])
        toc_pages_check = qa.check_toc_displayed_pages(
            stale_toc,
            rendered_page_texts=[
                "CONTENTS Chapter 1 Introduction 6",
                "Chapter 1 Introduction",
            ],
        )
        stale_page_issue = issue_for({"issues": toc_pages_check["issues"]}, "toc_displayed_page_mismatch")
        assert_equal(stale_page_issue["displayed_page"], "6", "stale TOC page is exposed")
        assert_equal(stale_page_issue["actual_displayed_page"], "2", "rendered heading page is exposed")

        heading_doc = Document()
        heading_doc.add_paragraph("Abstract")
        heading_doc.add_paragraph("Chapter 1 Introduction")
        second = heading_doc.add_paragraph("1.1 Research Background")
        second.paragraph_format.first_line_indent = Pt(21)
        heading_path = root / "bad-heading.docx"
        heading_doc.save(heading_path)
        heading_issue = issue_for(
            qa.run_structural_checks(heading_path, profile), "heading_format_mismatch"
        )
        assert_equal(heading_issue["text"], "1.1 Research Background", "heading location retained")
        assert_equal(heading_issue["differences"][0]["field"], "first_line_pt", "indent is diagnosed")

        figure_doc = Document()
        figure_doc.add_paragraph("Figure 1-1: Missing Plot")
        figure_path = root / "orphan-figure.docx"
        figure_doc.save(figure_path)
        figure_issue = issue_for(
            qa.run_structural_checks(figure_path, profile), "figure_caption_without_image"
        )
        assert_equal(figure_issue["text"], "Figure 1-1: Missing Plot", "caption context retained")

        source_path = root / "source.docx"
        output_path = root / "changed.docx"
        source = make_source_doc()
        output = make_changed_output_doc()
        assert_true(fmt.detect_front_matter_protection(source)["enabled"], "source enables protection")
        source.save(source_path)
        output.save(output_path)
        source_comparison = qa.run_structural_checks(output_path, profile, source_path=source_path)
        front_issue = issue_for(source_comparison, "front_matter_format_changed")
        assert_equal(front_issue["text"], "原创性声明", "protected paragraph is located")
        table_issue = issue_for(source_comparison, "front_matter_table_format_changed")
        assert_equal(table_issue["text"], "封面字段", "protected table cell is located")
        issue_for(source_comparison, "equation_math_object_lost")
        issue_for(source_comparison, "page_number_lost")

        text_equation = Document()
        paragraph = text_equation.add_paragraph("x = y + z (1)")
        paragraph.paragraph_format.first_line_indent = Pt(21)
        equation_path = root / "indented-equation.docx"
        text_equation.save(equation_path)
        equation_issue = issue_for(
            qa.run_structural_checks(equation_path, profile), "equation_format_mismatch"
        )
        assert_equal(equation_issue["role"], "equation_text_fallback", "text equation role retained")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
