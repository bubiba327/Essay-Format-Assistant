#!/usr/bin/env python3
"""Reusable structural QA checks for generated thesis DOCX outputs."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

import thesis_format_from_sample as fmt


IMPORTANT_TABLE_FIELDS = {"size_pt", "line_spacing", "align", "bold", "italic"}
HEADING_FIELDS = {"align", "first_line_pt", "left_pt", "hanging_pt", "right_pt", "blank_before", "blank_after"}
PROTECTED_PARAGRAPH_FIELDS = {
    "size_pt",
    "line_spacing",
    "align",
    "bold",
    "italic",
    "first_line_pt",
    "left_pt",
    "hanging_pt",
    "right_pt",
}


def described_differences(
    actual: dict[str, Any],
    expected: dict[str, Any],
    field_names: set[str],
) -> list[dict[str, Any]]:
    differences: list[dict[str, Any]] = []
    for field in sorted(field_names):
        if field not in expected:
            continue
        actual_value = actual.get(field)
        expected_value = expected[field]
        if fmt.equivalent_format_value(field, actual_value, expected_value):
            continue
        differences.append(
            {
                "field": field,
                "actual": actual_value,
                "expected": expected_value,
                "actual_display": fmt.format_value(field, actual_value),
                "expected_display": fmt.format_value(field, expected_value),
            }
        )
    return differences


def table_cell_mismatch_issue(
    table_index: int,
    row_index: int,
    column_index: int,
    text: str,
    role: str,
    changes: list[tuple[str, Any, Any]],
    actual_signature: dict[str, Any],
    expected_signature: dict[str, Any],
) -> dict[str, Any]:
    differences = [
        {
            "field": field,
            "actual": actual,
            "expected": expected,
            "actual_display": fmt.format_value(field, actual),
            "expected_display": fmt.format_value(field, expected),
        }
        for field, actual, expected in changes
    ]
    return {
        "code": "captioned_table_cell_format_mismatch",
        "severity": "error",
        "table_index": table_index + 1,
        "row": row_index + 1,
        "column": column_index + 1,
        "text": text[:100],
        "role": role,
        "actual_signature": actual_signature,
        "expected_signature": expected_signature,
        "differences": differences,
    }


def check_captioned_table_formats(
    doc: Document,
    profile: dict[str, Any],
    max_cells_per_table: int = 30,
) -> dict[str, Any]:
    enabled_roles = fmt.get_enabled_roles(profile)
    issues: list[dict[str, Any]] = []
    captioned_table_count = 0
    checked_cell_count = 0
    roles = profile.get("roles", {})
    if not enabled_roles.intersection(fmt.TABLE_FORMAT_ROLES):
        return {
            "status": "pass",
            "captioned_table_count": 0,
            "checked_cell_count": 0,
            "issues": [],
        }

    for table_index, table in enumerate(doc.tables):
        if not fmt.has_preceding_table_caption(table):
            continue
        captioned_table_count += 1
        checked_for_table = 0
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    role = fmt.resolve_table_cell_format_role(
                        fmt.table_cell_role(row_index, column_index, text),
                        enabled_roles,
                        profile,
                    )
                    if role is None or role not in roles:
                        continue
                    checked_cell_count += 1
                    checked_for_table += 1
                    actual_signature = fmt.paragraph_signature(paragraph)
                    expected_signature = fmt.target_signature(roles[role], profile)
                    changes = [
                        change
                        for change in fmt.format_changes(
                            actual_signature,
                            roles[role],
                            profile,
                        )
                        if change[0] in IMPORTANT_TABLE_FIELDS
                    ]
                    if changes:
                        issues.append(
                            table_cell_mismatch_issue(
                                table_index,
                                row_index,
                                column_index,
                                text,
                                role,
                                changes,
                                actual_signature,
                                expected_signature,
                            )
                        )
                    if checked_for_table >= max_cells_per_table:
                        break
                if checked_for_table >= max_cells_per_table:
                    break
            if checked_for_table >= max_cells_per_table:
                break

    return {
        "status": "pass" if not issues else "fail",
        "captioned_table_count": captioned_table_count,
        "checked_cell_count": checked_cell_count,
        "issues": issues,
    }


def check_toc_structure(doc: Document, profile: dict[str, Any]) -> dict[str, Any]:
    enabled_roles = fmt.get_enabled_roles(profile)
    issues: list[dict[str, Any]] = []
    entry_count = 0
    for paragraph, role in fmt.iter_toc_content_paragraphs(doc):
        if role not in enabled_roles:
            continue
        entry_count += 1
        expected_style = fmt.TOC_ROLE_TO_STYLE_ID.get(role)
        actual_style = fmt.paragraph_style_id(paragraph._p)
        if not expected_style or actual_style.upper() == expected_style:
            continue
        issues.append(
            {
                "code": "toc_style_mismatch",
                "severity": "error",
                "text": fmt.paragraph_element_text(paragraph._p).strip()[:100],
                "role": role,
                "actual_style": actual_style or "no style",
                "expected_style": expected_style,
            }
        )
    return {"status": "pass" if not issues else "fail", "entry_count": entry_count, "issues": issues}


def check_toc_displayed_pages(
    doc: Document,
    rendered_pdf_path: Path | None = None,
    rendered_page_texts: list[str] | None = None,
) -> dict[str, Any]:
    if rendered_page_texts is None and rendered_pdf_path is None:
        return {"status": "not-checked", "entry_count": 0, "mismatch_count": 0, "issues": []}
    page_texts = rendered_page_texts
    if page_texts is None:
        page_texts = fmt.extract_pdf_page_texts(rendered_pdf_path)
    report = fmt.check_toc_page_results_against_rendered_text(doc, page_texts)
    issues = [
        {
            **mismatch,
            "severity": "error",
        }
        for mismatch in report["mismatches"]
    ]
    return {
        "status": report["status"],
        "entry_count": report["entry_count"],
        "mismatch_count": len(issues),
        "toc_page_indices": report["toc_page_indices"],
        "issues": issues,
    }


def check_heading_structure(doc: Document, profile: dict[str, Any]) -> dict[str, Any]:
    enabled_roles = fmt.get_enabled_roles(profile)
    issues: list[dict[str, Any]] = []
    heading_count = 0
    reached_main_text = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text or "\t" in text:
            continue
        if re.match(r"^Chapter\s+\d+\b", text, re.I) or re.match(r"^第[一二三四五六七八九十]+章", text):
            reached_main_text = True
            role = "chapter"
        elif not reached_main_text:
            continue
        elif fmt.is_third_heading_text(text):
            role = "third"
        elif fmt.is_second_heading_text(text):
            role = "second"
        else:
            continue
        if role not in enabled_roles or role not in profile.get("roles", {}):
            continue
        heading_count += 1
        actual = fmt.paragraph_signature(paragraph)
        actual["blank_before"] = len(fmt.previous_blank_siblings(paragraph))
        actual["blank_after"] = len(fmt.next_blank_siblings(paragraph))
        expected = fmt.target_signature(profile["roles"][role], profile)
        differences = described_differences(actual, expected, HEADING_FIELDS)
        if differences:
            issues.append(
                {
                    "code": "heading_format_mismatch",
                    "severity": "error",
                    "text": text[:100],
                    "role": role,
                    "actual_signature": actual,
                    "expected_signature": expected,
                    "differences": differences,
                }
            )
    return {"status": "pass" if not issues else "fail", "heading_count": heading_count, "issues": issues}


def paragraph_image_widths(paragraph) -> list[float]:
    widths: list[float] = []
    for element in paragraph._p.iter():
        if element.tag != fmt.qn("wp:extent") or not element.get("cx"):
            continue
        try:
            widths.append(round(int(element.get("cx")) / 12700, 2))
        except ValueError:
            continue
    return widths


def nearby_paragraph(paragraphs: list[Any], index: int, direction: int, max_steps: int = 4):
    steps = 0
    probe = index + direction
    while 0 <= probe < len(paragraphs) and steps < max_steps:
        candidate = paragraphs[probe]
        steps += 1
        if candidate.text.strip() or fmt.has_image(candidate):
            return candidate
        probe += direction
    return None


def check_figure_structure(doc: Document, profile: dict[str, Any]) -> dict[str, Any]:
    issues: list[dict[str, Any]] = []
    paragraphs = list(doc.paragraphs)
    image_count = sum(1 for paragraph in fmt.iter_document_paragraphs(doc) if fmt.has_image(paragraph))
    caption_count = 0
    expected_width = (profile.get("image_format") or {}).get("image_width_pt")
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if fmt.caption_role(text) != "figure_caption":
            continue
        caption_count += 1
        previous = nearby_paragraph(paragraphs, index, -1)
        following = nearby_paragraph(paragraphs, index, 1)
        if previous is None or not fmt.has_image(previous):
            issues.append(
                {
                    "code": (
                        "figure_caption_wrong_side"
                        if following is not None and fmt.has_image(following)
                        else "figure_caption_without_image"
                    ),
                    "severity": "error",
                    "text": text[:100],
                    "expected_relation": "image immediately before caption",
                }
            )
            continue
        widths = paragraph_image_widths(previous)
        if expected_width and widths and min(widths) < float(expected_width) * 0.65:
            issues.append(
                {
                    "code": "figure_image_unexpectedly_tiny",
                    "severity": "error",
                    "text": text[:100],
                    "actual_width_pt": min(widths),
                    "expected_width_pt": float(expected_width),
                }
            )
    return {
        "status": "pass" if not issues else "fail",
        "figure_caption_count": caption_count,
        "image_count": image_count,
        "issues": issues,
    }


def math_object_count(doc: Document) -> int:
    return sum(1 for paragraph in fmt.iter_document_paragraphs(doc) if fmt.paragraph_has_math_object(paragraph))


def check_equation_structure(
    doc: Document,
    profile: dict[str, Any],
    source_doc: Document | None = None,
) -> dict[str, Any]:
    issues: list[dict[str, Any]] = []
    enabled_roles = fmt.get_enabled_roles(profile)
    output_math_object_count = math_object_count(doc)
    source_math_object_count = math_object_count(source_doc) if source_doc is not None else None
    if source_math_object_count and output_math_object_count < source_math_object_count:
        issues.append(
            {
                "code": "equation_math_object_lost",
                "severity": "error",
                "source_math_object_count": source_math_object_count,
                "output_math_object_count": output_math_object_count,
            }
        )
    equation_paragraph_count = 0
    for paragraph in doc.paragraphs:
        role = fmt.equation_role_for_paragraph(paragraph, paragraph.text.strip())
        if role not in enabled_roles or role not in profile.get("roles", {}):
            continue
        equation_paragraph_count += 1
        actual = fmt.paragraph_signature(paragraph)
        expected = fmt.target_signature(profile["roles"][role], profile)
        differences = described_differences(actual, expected, {"first_line_pt"})
        if differences:
            issues.append(
                {
                    "code": "equation_format_mismatch",
                    "severity": "error",
                    "text": paragraph.text.strip()[:100],
                    "role": role,
                    "differences": differences,
                }
            )
    return {
        "status": "pass" if not issues else "fail",
        "source_math_object_count": source_math_object_count,
        "output_math_object_count": output_math_object_count,
        "equation_paragraph_count": equation_paragraph_count,
        "issues": issues,
    }


def nonblank_protected_paragraphs(doc: Document) -> tuple[dict[str, Any], list[Any]]:
    protection = fmt.detect_front_matter_protection(doc)
    paragraphs = [
        doc.paragraphs[index]
        for index in protection.get("protected_paragraph_indices", [])
        if doc.paragraphs[index].text.strip()
    ]
    return protection, paragraphs


def check_front_matter_unchanged(
    doc: Document,
    source_doc: Document | None = None,
) -> dict[str, Any]:
    if source_doc is None:
        return {
            "status": "not-checked",
            "protected_paragraph_count": 0,
            "protected_table_count": 0,
            "issues": [],
        }
    protection, source_paragraphs = nonblank_protected_paragraphs(source_doc)
    if not protection.get("enabled"):
        return {
            "status": "not-applicable",
            "protected_paragraph_count": 0,
            "protected_table_count": 0,
            "issues": [],
        }
    issues: list[dict[str, Any]] = []
    output_by_text: dict[str, list[Any]] = {}
    for paragraph in doc.paragraphs:
        key = fmt.compact_text(paragraph.text)
        if key:
            output_by_text.setdefault(key, []).append(paragraph)
    for paragraph in source_paragraphs:
        text = fmt.compact_text(paragraph.text)
        matches = output_by_text.get(text, [])
        if not matches:
            issues.append(
                {
                    "code": "front_matter_text_missing",
                    "severity": "error",
                    "text": text[:100],
                }
            )
            continue
        candidate = matches.pop(0)
        expected = fmt.paragraph_signature(paragraph)
        actual = fmt.paragraph_signature(candidate)
        differences = described_differences(actual, expected, PROTECTED_PARAGRAPH_FIELDS)
        if differences:
            issues.append(
                {
                    "code": "front_matter_format_changed",
                    "severity": "error",
                    "text": text[:100],
                    "actual_signature": actual,
                    "expected_signature": expected,
                    "differences": differences,
                }
            )
    source_tables = [
        table for table in source_doc.tables if table._tbl in protection.get("protected_table_elements", set())
    ]
    output_table_text = {fmt.table_compact_text(table): table for table in doc.tables}
    for table in source_tables:
        text = fmt.table_compact_text(table)
        if text not in output_table_text:
            issues.append(
                {
                    "code": "front_matter_table_missing",
                    "severity": "error",
                    "text": text[:100],
                }
            )
            continue
        output_table = output_table_text[text]
        expected = fmt.table_signature(table)
        actual = fmt.table_signature(output_table)
        differences = described_differences(
            actual,
            expected,
            {
                "table_alignment",
                "table_text_wrapping",
                "table_layout",
                "table_width_pct",
                "table_width_pt",
            },
        )
        if differences:
            issues.append(
                {
                    "code": "front_matter_table_format_changed",
                    "severity": "error",
                    "text": text[:100],
                    "differences": differences,
                }
            )
        for row_index, row in enumerate(table.rows):
            if row_index >= len(output_table.rows):
                break
            for column_index, cell in enumerate(row.cells):
                if column_index >= len(output_table.rows[row_index].cells):
                    break
                output_cell = output_table.rows[row_index].cells[column_index]
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    cell_text = fmt.compact_text(paragraph.text)
                    if not cell_text or paragraph_index >= len(output_cell.paragraphs):
                        continue
                    expected = fmt.paragraph_signature(paragraph)
                    actual = fmt.paragraph_signature(output_cell.paragraphs[paragraph_index])
                    differences = described_differences(actual, expected, PROTECTED_PARAGRAPH_FIELDS)
                    if differences:
                        issues.append(
                            {
                                "code": "front_matter_table_format_changed",
                                "severity": "error",
                                "table_text": text[:100],
                                "row": row_index + 1,
                                "column": column_index + 1,
                                "text": cell_text[:100],
                                "actual_signature": actual,
                                "expected_signature": expected,
                                "differences": differences,
                            }
                        )
    return {
        "status": "pass" if not issues else "fail",
        "protected_paragraph_count": len(source_paragraphs),
        "protected_table_count": len(source_tables),
        "issues": issues,
    }


def page_number_count(doc: Document) -> int:
    story_elements: set[int] = set()
    count = 0
    for section in doc.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            identity = id(story._element)
            if identity in story_elements:
                continue
            story_elements.add(identity)
            count += sum(1 for paragraph in story.paragraphs if fmt.paragraph_has_page_field(paragraph))
    return count


def check_page_numbers_preserved(
    doc: Document,
    profile: dict[str, Any],
    source_doc: Document | None = None,
) -> dict[str, Any]:
    output_count = page_number_count(doc)
    source_count = page_number_count(source_doc) if source_doc is not None else None
    preserve = (profile.get("header_footer") or {}).get("preserve_target_page_numbers", True)
    issues: list[dict[str, Any]] = []
    if preserve and source_count and output_count == 0:
        issues.append(
            {
                "code": "page_number_lost",
                "severity": "error",
                "source_page_number_count": source_count,
                "output_page_number_count": output_count,
            }
        )
    return {
        "status": "pass" if not issues else "fail",
        "source_page_number_count": source_count,
        "output_page_number_count": output_count,
        "preserve_target_page_numbers": bool(preserve),
        "issues": issues,
    }


def run_structural_checks(
    output_path: Path,
    profile: dict[str, Any],
    source_path: Path | None = None,
    rendered_pdf_path: Path | None = None,
) -> dict[str, Any]:
    doc = Document(output_path)
    source_doc = Document(source_path) if source_path else None
    checks = {
        "captioned_tables": check_captioned_table_formats(doc, profile),
        "toc": check_toc_structure(doc, profile),
        "toc_pages": check_toc_displayed_pages(doc, rendered_pdf_path=rendered_pdf_path),
        "headings": check_heading_structure(doc, profile),
        "figures": check_figure_structure(doc, profile),
        "equations": check_equation_structure(doc, profile, source_doc),
        "front_matter": check_front_matter_unchanged(doc, source_doc),
        "page_numbers": check_page_numbers_preserved(doc, profile, source_doc),
    }
    issues = [issue for check in checks.values() for issue in check["issues"]]
    metrics = {
        "table_count": checks["captioned_tables"]["captioned_table_count"],
        "table_cell_count": checks["captioned_tables"]["checked_cell_count"],
        "toc_entry_count": checks["toc"]["entry_count"],
        "toc_page_mismatch_count": checks["toc_pages"]["mismatch_count"],
        "heading_count": checks["headings"]["heading_count"],
        "figure_caption_count": checks["figures"]["figure_caption_count"],
        "image_count": checks["figures"]["image_count"],
        "equation_count": checks["equations"]["output_math_object_count"],
        "protected_paragraph_count": checks["front_matter"]["protected_paragraph_count"],
        "protected_table_count": checks["front_matter"]["protected_table_count"],
        "page_number_count": checks["page_numbers"]["output_page_number_count"],
    }
    return {
        "status": "pass" if not issues else "fail",
        "output": str(output_path),
        "source": str(source_path) if source_path else None,
        "metrics": metrics,
        "checks": checks,
        "issues": issues,
    }
