#!/usr/bin/env python3
"""Regression checks for thesis table formatting."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

import thesis_format_from_sample as fmt


def assert_close(actual: float | None, expected: float, message: str) -> None:
    if actual is None or abs(actual - expected) > 0.01:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def assert_equal(actual, expected, message: str) -> None:
    if actual != expected:
        raise AssertionError(f"{message}: expected {expected!r}, got {actual!r}")


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.5


def make_profile_without_enabled_table_header() -> dict:
    profile = fmt.default_profile()
    profile["enabled_roles"] = [
        "table_overall",
        "table_body_variable",
        "table_body_value",
    ]
    return profile


def make_captioned_first_table_doc() -> tuple[Document, object]:
    doc = Document()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("Chapter 5 Empirical Process and Analysis")
    doc.add_paragraph("Table 5-1: Descriptive Statistics")
    table = doc.add_table(rows=2, cols=3)
    values = [
        ["", "count", "mean"],
        ["lnTC", "403", "10.1926"],
    ]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_text(cell, values[row_idx][col_idx])
    return doc, table


def main() -> int:
    profile = make_profile_without_enabled_table_header()
    doc, table = make_captioned_first_table_doc()

    fmt.format_tables(
        doc,
        profile,
        skip_first_table=True,
        comment_tracker=fmt.CommentTracker(mode="none"),
        protected_table_elements=set(),
    )

    header_signature = fmt.paragraph_signature(table.cell(0, 1).paragraphs[0])
    variable_signature = fmt.paragraph_signature(table.cell(1, 0).paragraphs[0])
    value_signature = fmt.paragraph_signature(table.cell(1, 1).paragraphs[0])

    assert_close(header_signature["size_pt"], 9.0, "header fallback size")
    assert_close(header_signature["line_spacing"], 1.2, "header fallback line spacing")
    assert_equal(header_signature["bold"], True, "header fallback bold")
    assert_close(variable_signature["size_pt"], 9.0, "variable cell size")
    assert_close(variable_signature["line_spacing"], 1.2, "variable cell line spacing")
    assert_equal(variable_signature["italic"], True, "variable cell italic")
    assert_close(value_signature["size_pt"], 9.0, "value cell size")
    assert_close(value_signature["line_spacing"], 1.2, "value cell line spacing")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
