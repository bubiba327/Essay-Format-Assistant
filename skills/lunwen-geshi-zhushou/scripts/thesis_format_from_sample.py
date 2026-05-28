#!/usr/bin/env python3
"""Infer thesis formatting from a sample Word file and apply it to a target copy.

This script is intentionally heuristic: it recognizes common thesis roles
from text patterns (abstract, contents, Chapter headings, 1.1/1.1.1 headings,
captions, references, acknowledgements) and applies the detected/default
formatting to a copied target DOCX. It never edits the source target in place
unless --output explicitly points to that same file.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections import Counter, defaultdict
from contextlib import contextmanager
from dataclasses import asdict, dataclass, field
from pathlib import Path
from time import perf_counter
from typing import Any

from docx import Document
from docx.enum.dml import MSO_COLOR_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

from thesis_style_engine import apply_role_style, ensure_role_styles


LATIN_FONT = "Times New Roman"
EAST_ASIA_FONT = "宋体"
CACHE_VERSION = 2
PROFILE_LINT_REPORT_BASENAME = "profile-lint-report"
EAST_ASIA_FONT_ALIASES = {
    "simsun",
    "simhei",
    "kaiti",
    "fangsong",
    "microsoft yahei",
    "microsoft jhenghei",
    "pmingliu",
    "mingliu",
}
PDF_RENDERER = Path(__file__).resolve().with_name("render_pdf_pages.swift")
MACOS_SOFFICE = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
DEFAULT_SHARED_SAMPLE_CACHE_DIR = Path.home() / ".codex" / "cache" / "lunwen-geshi-zhushou" / "sample-analysis"
MATH_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
LEGACY_SUPPLEMENT_ROLES = (
    "toc1",
    "toc2",
    "toc3",
    "table_caption",
    "table_note",
    "note",
    "table_footnote",
)
CAPTION_NARRATIVE_WORDS = {
    "show",
    "shows",
    "showed",
    "showing",
    "provide",
    "provides",
    "provided",
    "present",
    "presents",
    "presented",
    "report",
    "reports",
    "reported",
    "examine",
    "examines",
    "examined",
    "illustrate",
    "illustrates",
    "illustrated",
    "indicate",
    "indicates",
    "demonstrate",
    "demonstrates",
    "summarize",
    "summarizes",
    "is",
    "are",
    "was",
    "were",
}
CHINESE_CAPTION_NARRATIVE_PREFIXES = (
    "显示",
    "表明",
    "说明",
    "展示",
    "呈现",
    "列示",
    "反映",
    "可见",
    "所示",
    "如下",
    "中",
)
CHINESE_NUMERAL_CHARS = "一二三四五六七八九十百千万零〇两"
CHINESE_TOC_HEADINGS = {"CONTENTS", "目  录", "目 录", "目录", "正文目录", "目次", "目  次"}
EQUATION_NUMBER_RE = re.compile(rf"^[（(]\s*[0-9{CHINESE_NUMERAL_CHARS}]+(?:\s*[-－–—.．]\s*[0-9{CHINESE_NUMERAL_CHARS}]+)?\s*[）)]$")
EQUATION_TRAILING_NUMBER_RE = re.compile(rf"\s*[（(]\s*[0-9{CHINESE_NUMERAL_CHARS}]+(?:\s*[-－–—.．]\s*[0-9{CHINESE_NUMERAL_CHARS}]+)?\s*[）)]\s*$")
EQUATION_MATH_SIGNAL_RE = re.compile(
    r"(?:"
    r"[=Σ∑×÷±≤≥√∞∫∂≈≠]|"
    r"[A-Za-zΑ-ω]+_[A-Za-z0-9,{}]+|"
    r"[A-Za-zΑ-ω]\s*(?:\^|/)\s*[A-Za-z0-9]|"
    r"(?:β|α|γ|δ|μ|λ|ε|θ|ρ|σ)"
    r")"
)
FRONT_MATTER_PROTECTION_KEYWORDS = (
    "原创性声明",
    "诚信承诺",
    "版权使用授权书",
    "本人郑重承诺",
    "承诺人签名",
    "保密学位论文",
    "不保密",
    "学位论文属于",
)
FRONT_COVER_KEYWORDS = (
    "毕业论文正文",
    "毕业论文（设计）",
    "题目：",
    "学    院",
    "专    业",
    "班    级",
    "学    号",
    "学生姓名",
    "指导教师",
)

ROLE_LABELS = {
    "front_heading": "摘要/Abstract 标题",
    "contents_heading": "目录标题",
    "paper_title": "论文题名",
    "paper_meta": "封面信息",
    "toc1": "目录一级条目",
    "toc2": "目录二级条目",
    "toc3": "目录三级条目",
    "chapter": "一级标题",
    "second": "二级标题",
    "third": "三级标题",
    "body": "正文",
    "keywords": "关键词段落",
    "keyword_label": "关键词标签",
    "figure_caption": "图题",
    "table_caption": "表格标题",
    "caption": "图表题注",
    "note": "注释",
    "figure_note": "图注",
    "table_note": "表注",
    "equation": "公式",
    "equation_number": "公式编号",
    "equation_text_fallback": "普通文本公式",
    "reference_heading": "参考文献标题",
    "reference_entry": "参考文献条目",
    "ack_heading": "致谢标题",
    "ack_body": "致谢正文",
    "signature": "署名/日期",
    "header_overall": "页眉整体设置",
    "header_text": "页眉文字",
    "footer_overall": "页脚整体设置",
    "footer_text": "页脚文字",
    "page_number": "页码",
    "table_overall": "表格整体设置",
    "table_header": "表头（列标题）",
    "table_body": "表格正文（兼容旧表）",
    "table_body_variable": "表格正文-变量名列",
    "table_body_definition": "表格正文-定义/说明列",
    "table_body_value": "表格正文-数值列",
    "table_footnote": "表尾标注",
    "table_note_cell": "表内注释",
    "table_text": "表格文字",
}

FIELD_LABELS = {
    "font": "西文字体",
    "east_asia_font": "中文字体",
    "font_color_name": "字体颜色名称",
    "font_color_hex": "字体颜色值",
    "size_pt": "字号",
    "bold": "加粗",
    "italic": "斜体",
    "all_caps": "全部大写",
    "small_caps": "小型大写",
    "align": "对齐方式",
    "line_spacing": "行距",
    "before_pt": "段前",
    "after_pt": "段后",
    "blank_before": "段前空行数",
    "blank_after": "段后空行数",
    "first_line_pt": "首行缩进",
    "left_pt": "左缩进",
    "hanging_pt": "悬挂缩进",
    "right_pt": "右缩进",
    "toc_page_number_right_aligned": "目录页码右对齐",
    "toc_leader": "目录点引导线",
    "toc_right_tab_pt": "目录右侧制表位",
    "table_alignment": "表格水平位置",
    "table_text_wrapping": "表格文字环绕",
    "table_layout": "表格布局方式",
    "table_width_pct": "表格宽度比例",
    "table_width_pt": "表格宽度",
    "cell_margin_top_pt": "单元格上边距",
    "cell_margin_bottom_pt": "单元格下边距",
    "cell_margin_left_pt": "单元格左边距",
    "cell_margin_right_pt": "单元格右边距",
}

ALIGN_LABELS = {
    "left": "左对齐",
    "center": "居中",
    "right": "右对齐",
    "justify": "两端对齐",
}

ALIGN_VALUES = {value: key for key, value in ALIGN_LABELS.items()} | {
    "left": "left",
    "center": "center",
    "right": "right",
    "justify": "justify",
}

TABLE_ALIGNMENT_LABELS = {
    "left": "左对齐",
    "center": "居中",
    "right": "右对齐",
}

TABLE_ALIGNMENT_VALUES = {value: key for key, value in TABLE_ALIGNMENT_LABELS.items()} | {
    "left": "left",
    "center": "center",
    "right": "right",
}

TABLE_WRAPPING_LABELS = {
    "none": "不环绕",
    "around": "环绕",
}

TABLE_WRAPPING_VALUES = {value: key for key, value in TABLE_WRAPPING_LABELS.items()} | {
    "none": "none",
    "around": "around",
}

TABLE_LAYOUT_LABELS = {
    "fixed": "固定列宽",
    "autofit": "自动调整",
}

TABLE_LAYOUT_VALUES = {value: key for key, value in TABLE_LAYOUT_LABELS.items()} | {
    "fixed": "fixed",
    "autofit": "autofit",
}

PAGE_NUMBER_LOCATION_LABELS = {
    "header": "页眉",
    "footer": "页脚",
}

PAGE_NUMBER_LOCATION_VALUES = {value: key for key, value in PAGE_NUMBER_LOCATION_LABELS.items()} | {
    "header": "header",
    "footer": "footer",
}

PAGE_NUMBER_FORMAT_LABELS = {
    "decimal": "阿拉伯数字(1,2,3)",
    "upperRoman": "大写罗马数字(I,II,III)",
    "lowerRoman": "小写罗马数字(i,ii,iii)",
    "upperLetter": "大写字母(A,B,C)",
    "lowerLetter": "小写字母(a,b,c)",
}

PAGE_NUMBER_FORMAT_VALUES = {value: key for key, value in PAGE_NUMBER_FORMAT_LABELS.items()} | {
    "decimal": "decimal",
    "upperRoman": "upperRoman",
    "lowerRoman": "lowerRoman",
    "upperLetter": "upperLetter",
    "lowerLetter": "lowerLetter",
}

PAGE_NUMBER_SECTION_MODE_LABELS = {
    "continue": "续前节",
    "restart": "重新编号",
}

PAGE_NUMBER_SECTION_MODE_VALUES = {value: key for key, value in PAGE_NUMBER_SECTION_MODE_LABELS.items()} | {
    "continue": "continue",
    "restart": "restart",
}

LEADER_LABELS = {
    "none": "无",
    "dots": "点引导线",
    "dashes": "短横线",
    "lines": "下划线",
    "heavy": "粗线",
    "middle_dot": "中点",
}

LEADER_VALUES = {value: key for key, value in LEADER_LABELS.items()} | {
    "none": "none",
    "dot": "dots",
    "dots": "dots",
    "dashes": "dashes",
    "lines": "lines",
    "heavy": "heavy",
    "middle_dot": "middle_dot",
}

COMMON_COLOR_NAMES = {
    "#000000": "黑色",
    "#FFFFFF": "白色",
    "#FF0000": "红色",
    "#00FF00": "绿色",
    "#0000FF": "蓝色",
    "#FFFF00": "黄色",
    "#808080": "灰色",
}

ROLE_ORDER = [
    "front_heading",
    "contents_heading",
    "toc1",
    "toc2",
    "toc3",
    "chapter",
    "second",
    "third",
    "body",
    "keywords",
    "keyword_label",
    "figure_caption",
    "table_caption",
    "caption",
    "figure_note",
    "table_note",
    "note",
    "equation",
    "equation_number",
    "equation_text_fallback",
    "reference_heading",
    "reference_entry",
    "ack_heading",
    "ack_body",
    "signature",
    "header_overall",
    "header_text",
    "footer_overall",
    "footer_text",
    "page_number",
    "table_overall",
    "table_header",
    "table_body_variable",
    "table_body_definition",
    "table_body_value",
    "table_footnote",
    "table_note_cell",
]

FORMAT_TABLE_COLUMNS = [
    "角色ID",
    "格式部分",
    "范文示例",
    "是否应用",
    "西文字体",
    "中文字体",
    "字体颜色名称",
    "字体颜色值",
    "字号(磅)",
    "加粗",
    "斜体",
    "全部大写",
    "小型大写",
    "对齐方式",
    "行距",
    "段前(磅)",
    "段后(磅)",
    "段前空行数",
    "段后空行数",
    "首行缩进(磅)",
    "左缩进(磅)",
    "悬挂缩进(磅)",
    "右缩进(磅)",
    "目录页码右对齐",
    "目录点引导线",
    "目录右侧制表位(磅)",
    "表格水平位置",
    "表格文字环绕",
    "表格布局方式",
    "表格宽度比例(%)",
    "表格宽度(磅)",
    "单元格上边距(磅)",
    "单元格下边距(磅)",
    "单元格左边距(磅)",
    "单元格右边距(磅)",
    "备注",
]

HEADER_FOOTER_NOTICE = "页眉页脚格式已更改。"
FORMAT_BOOL_FIELDS = {"bold", "italic", "all_caps", "small_caps"}
TOC_ROLE_TO_STYLE_ID = {"toc1": "TOC1", "toc2": "TOC2", "toc3": "TOC3"}
TOC_STYLE_ID_TO_ROLE = {
    "toc1": "toc1",
    "toc 1": "toc1",
    "toc2": "toc2",
    "toc 2": "toc2",
    "toc3": "toc3",
    "toc 3": "toc3",
}
HEADING_OUTLINE_LEVEL_BY_ROLE = {
    "front_heading": 0,
    "chapter": 0,
    "second": 1,
    "third": 2,
}
LEGACY_WORD_SUFFIXES = {".doc", ".rtf", ".odt"}
LEGACY_TOC_UNRELIABLE_BOOL_FIELDS = {"italic", "all_caps", "small_caps"}


@dataclass
class CommentTracker:
    mode: str = "role"
    author: str = "论文格式助手"
    max_comments: int = 120
    count: int = 0
    roles_seen: set[str] = field(default_factory=set)

    def should_comment(self, role: str, changes: list[tuple[str, Any, Any]], extra_notes: list[str] | None = None) -> bool:
        if self.mode == "none" or self.count >= self.max_comments:
            return False
        if not changes and not extra_notes:
            return False
        if self.mode == "role" and role in self.roles_seen:
            return False
        return True

    def mark_commented(self, role: str) -> None:
        self.count += 1
        self.roles_seen.add(role)


@dataclass
class PerformanceTracker:
    operation: str
    qa_level: str
    started: float = field(default_factory=perf_counter)
    phases: dict[str, float] = field(default_factory=dict)
    metrics: dict[str, Any] = field(default_factory=dict)

    @contextmanager
    def measure(self, name: str):
        phase_started = perf_counter()
        try:
            yield
        finally:
            elapsed = perf_counter() - phase_started
            self.phases[name] = round(self.phases.get(name, 0.0) + elapsed, 4)

    def increment(self, name: str, amount: int = 1) -> None:
        self.metrics[name] = int(self.metrics.get(name, 0)) + amount

    def set_metric(self, name: str, value: Any) -> None:
        self.metrics[name] = value

    def write(self, analysis_dir: Path) -> Path:
        analysis_dir.mkdir(parents=True, exist_ok=True)
        output = analysis_dir / f"performance-{self.operation}.json"
        payload = {
            "operation": self.operation,
            "qa_level": self.qa_level,
            "elapsed_seconds": round(perf_counter() - self.started, 4),
            "phases": self.phases,
            "metrics": self.metrics,
        }
        output.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return output


@dataclass
class RoleFormat:
    font: str = LATIN_FONT
    east_asia_font: str = EAST_ASIA_FONT
    font_color_name: str = "黑色"
    font_color_hex: str = "#000000"
    size_pt: float = 12.0
    bold: bool | None = None
    italic: bool | None = None
    all_caps: bool | None = None
    small_caps: bool | None = None
    align: str = "justify"
    line_spacing: float = 1.5
    before_pt: float = 0.0
    after_pt: float = 0.0
    blank_before: int | None = None
    blank_after: int | None = None
    first_line_pt: float | None = None
    left_pt: float | None = None
    hanging_pt: float | None = None
    right_pt: float | None = None
    toc_page_number_right_aligned: bool = False
    toc_leader: str = "none"
    toc_right_tab_pt: float | None = None
    table_alignment: str | None = None
    table_text_wrapping: str | None = None
    table_layout: str | None = None
    table_width_pct: float | None = None
    table_width_pt: float | None = None
    cell_margin_top_pt: float | None = None
    cell_margin_bottom_pt: float | None = None
    cell_margin_left_pt: float | None = None
    cell_margin_right_pt: float | None = None


def default_profile() -> dict[str, Any]:
    return {
        "latin_font": LATIN_FONT,
        "east_asia_font": EAST_ASIA_FONT,
        "page": {
            "preserve_target_page_size": True,
            "top_in": 1.0,
            "bottom_in": 1.0,
            "left_in": 1.25,
            "right_in": 1.25,
        },
        "header_footer": {
            "apply_header_footer": True,
            "preserve_target_text": True,
            "preserve_target_page_numbers": True,
            "header_distance_in": 0.5,
            "footer_distance_in": 0.5,
            "different_first_page": False,
            "odd_even_pages": False,
            "page_number_enabled": False,
            "page_number_location": "footer",
            "page_number_section_mode": "restart",
            "front_page_number_format": "upperRoman",
            "body_page_number_format": "decimal",
            "front_page_number_start": 1,
            "body_page_number_start": 1,
        },
        "roles": {
            "front_heading": asdict(RoleFormat(size_pt=15, bold=True, align="center", first_line_pt=None)),
            "contents_heading": asdict(RoleFormat(size_pt=16, bold=True, align="center", first_line_pt=None)),
            "paper_title": asdict(RoleFormat(size_pt=16, bold=True, align="center", first_line_pt=None)),
            "paper_meta": asdict(RoleFormat(size_pt=14, bold=False, align="center", first_line_pt=None)),
            "toc1": asdict(RoleFormat(size_pt=12, bold=True, italic=False, align="justify", before_pt=6, after_pt=6, left_pt=0, first_line_pt=None, toc_page_number_right_aligned=True, toc_leader="dots")),
            "toc2": asdict(RoleFormat(size_pt=10.5, bold=True, italic=False, align="justify", left_pt=10.5, first_line_pt=None, toc_page_number_right_aligned=True, toc_leader="dots")),
            "toc3": asdict(RoleFormat(size_pt=10.5, bold=False, italic=True, align="justify", left_pt=21, first_line_pt=None, toc_page_number_right_aligned=True, toc_leader="dots")),
            "chapter": asdict(RoleFormat(size_pt=16, bold=True, italic=False, align="justify", before_pt=15.6, after_pt=7.8, blank_before=2, blank_after=0, first_line_pt=None)),
            "second": asdict(RoleFormat(size_pt=14, bold=True, italic=False, align="justify", blank_before=0, blank_after=0, first_line_pt=None)),
            "third": asdict(RoleFormat(size_pt=12, bold=True, italic=False, align="justify", blank_before=0, blank_after=0, first_line_pt=None)),
            "body": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="justify", first_line_pt=24)),
            "keywords": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="justify", first_line_pt=None)),
            "keyword_label": asdict(RoleFormat(size_pt=14, bold=True, italic=None, align="justify", first_line_pt=None)),
            "figure_caption": asdict(RoleFormat(size_pt=10.5, bold=True, italic=False, align="center", first_line_pt=None)),
            "table_caption": asdict(RoleFormat(size_pt=10.5, bold=True, italic=False, align="center", first_line_pt=None)),
            "caption": asdict(RoleFormat(size_pt=10.5, bold=True, italic=False, align="center", first_line_pt=None)),
            "note": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="left", first_line_pt=None)),
            "figure_note": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="left", first_line_pt=None)),
            "table_note": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="left", first_line_pt=None)),
            "equation": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="center", first_line_pt=None)),
            "equation_number": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="right", first_line_pt=None)),
            "equation_text_fallback": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="center", first_line_pt=None)),
            "reference_heading": asdict(RoleFormat(size_pt=16, bold=True, italic=None, align="center", first_line_pt=None)),
            "reference_entry": asdict(RoleFormat(size_pt=11, bold=False, italic=False, align="justify", first_line_pt=None)),
            "ack_heading": asdict(RoleFormat(size_pt=16, bold=True, italic=None, align="center", first_line_pt=None)),
            "ack_body": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="justify", first_line_pt=24.8)),
            "signature": asdict(RoleFormat(size_pt=12, bold=None, italic=None, align="right", first_line_pt=None)),
            "header_overall": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="center", line_spacing=1.0, first_line_pt=None)),
            "header_text": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="center", line_spacing=1.0, first_line_pt=None)),
            "footer_overall": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="center", line_spacing=1.0, first_line_pt=None)),
            "footer_text": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="center", line_spacing=1.0, first_line_pt=None)),
            "page_number": asdict(RoleFormat(size_pt=10.5, bold=None, italic=None, align="center", line_spacing=1.0, first_line_pt=None)),
            "table_overall": asdict(RoleFormat(size_pt=9, align="center", first_line_pt=None, table_alignment="center", table_text_wrapping="none", table_layout="fixed", table_width_pct=100, cell_margin_top_pt=2.5, cell_margin_bottom_pt=2.5, cell_margin_left_pt=5, cell_margin_right_pt=5)),
            "table_header": asdict(RoleFormat(size_pt=9, bold=True, italic=False, align="center", line_spacing=1.2, first_line_pt=None)),
            "table_body": asdict(RoleFormat(size_pt=9, bold=None, italic=None, align="center", line_spacing=1.2, first_line_pt=None)),
            "table_body_variable": asdict(RoleFormat(size_pt=9, bold=None, italic=True, align="center", line_spacing=1.2, first_line_pt=None)),
            "table_body_definition": asdict(RoleFormat(size_pt=9, bold=None, italic=False, align="center", line_spacing=1.2, first_line_pt=None)),
            "table_body_value": asdict(RoleFormat(size_pt=9, bold=None, italic=False, align="center", line_spacing=1.2, first_line_pt=None)),
            "table_footnote": asdict(RoleFormat(size_pt=9, bold=None, italic=None, align="left", line_spacing=1.2, first_line_pt=None)),
            "table_note_cell": asdict(RoleFormat(size_pt=9, bold=None, italic=None, align="left", line_spacing=1.2, first_line_pt=None)),
            "table_text": asdict(RoleFormat(size_pt=9, bold=None, italic=None, align="center", line_spacing=1.2, first_line_pt=None)),
        },
        "source_observations": {},
        "enabled_roles": [],
        "image_format": {},
        "document_language": {
            "dominant": "unknown",
            "cjk_chars": 0,
            "latin_letters": 0,
            "cjk_ratio": 0.0,
            "latin_ratio": 0.0,
        },
    }


def infer_document_language_profile(doc: Document) -> dict[str, Any]:
    texts: list[str] = []
    texts.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    combined = "\n".join(texts)
    cjk_chars = len(re.findall(r"[\u4e00-\u9fff]", combined))
    latin_letters = len(re.findall(r"[A-Za-z]", combined))
    total = cjk_chars + latin_letters
    if total == 0:
        dominant = "unknown"
    elif cjk_chars >= latin_letters * 1.5:
        dominant = "chinese-dominant"
    elif latin_letters >= cjk_chars * 1.5:
        dominant = "english-dominant"
    else:
        dominant = "mixed"
    return {
        "dominant": dominant,
        "cjk_chars": cjk_chars,
        "latin_letters": latin_letters,
        "cjk_ratio": round(cjk_chars / total, 3) if total else 0.0,
        "latin_ratio": round(latin_letters / total, 3) if total else 0.0,
    }


def pt(value) -> float | None:
    if value is None:
        return None
    try:
        return round(value.pt, 2)
    except Exception:
        return None


def align_name(value) -> str | None:
    return {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    }.get(value)


def to_alignment(name: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(name, WD_ALIGN_PARAGRAPH.JUSTIFY)


def to_table_alignment(name: str):
    return {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }.get(name, WD_TABLE_ALIGNMENT.CENTER)


def to_tab_leader(name: str):
    return {
        "none": WD_TAB_LEADER.SPACES,
        "dots": WD_TAB_LEADER.DOTS,
        "dashes": WD_TAB_LEADER.DASHES,
        "lines": WD_TAB_LEADER.LINES,
        "heavy": WD_TAB_LEADER.HEAVY,
        "middle_dot": WD_TAB_LEADER.MIDDLE_DOT,
    }.get(name, WD_TAB_LEADER.SPACES)


def normalize_spacing(value) -> float:
    if isinstance(value, float):
        return value
    return 1.5


def contains_cjk(text: str) -> bool:
    return any("\u3400" <= char <= "\u9fff" for char in text)


def clean_font_part(value: str | None) -> str | None:
    text = normalize_text_cell(value)
    if not text:
        return None
    return text.strip("'\"“”‘’ ")


def sanitize_latin_font(value: str | None) -> str | None:
    text = clean_font_part(value)
    if not text:
        return None
    for part in [clean_font_part(part) for part in re.split(r"[;,]", text)]:
        if not part:
            continue
        if contains_cjk(part):
            continue
        if part.lower() in EAST_ASIA_FONT_ALIASES:
            continue
        return part
    return None


def sanitize_east_asia_font(value: str | None) -> str | None:
    text = clean_font_part(value)
    if not text:
        return None
    for part in [clean_font_part(part) for part in re.split(r"[;,]", text)]:
        if not part:
            continue
        if contains_cjk(part) or part.lower() in EAST_ASIA_FONT_ALIASES:
            return part
    return text


def first_text_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return None


def xml_toggle_value(r_pr, tag_name: str) -> bool | None:
    if r_pr is None:
        return None
    element = r_pr.find(qn(f"w:{tag_name}"))
    if element is None:
        return None
    value = element.get(qn("w:val"))
    if value is None:
        return True
    return str(value).lower() not in {"0", "false", "off"}


def run_element_text(run_element) -> str:
    return "".join(node.text or "" for node in run_element.iter(qn("w:t")))


def first_text_run_element(paragraph_element):
    for run_element in paragraph_element.iter(qn("w:r")):
        if run_element_text(run_element).strip():
            return run_element
    return None


def run_element_signature(run_element) -> dict[str, Any]:
    if run_element is None:
        return {}
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        return {}
    data: dict[str, Any] = {}
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is not None:
        latin = sanitize_latin_font(r_fonts.get(qn("w:ascii"))) or sanitize_latin_font(r_fonts.get(qn("w:hAnsi")))
        east_asia = sanitize_east_asia_font(r_fonts.get(qn("w:eastAsia"))) or sanitize_east_asia_font(r_fonts.get(qn("w:cs")))
        if latin:
            data["font"] = latin
        if east_asia:
            data["east_asia_font"] = east_asia
    color = r_pr.find(qn("w:color"))
    if color is not None:
        value = normalize_text_cell(color.get(qn("w:val"))).upper()
        if value and value not in {"AUTO", "AUTOMATIC"} and re.fullmatch(r"[0-9A-F]{6}", value):
            hex_value = f"#{value}"
            data["font_color_name"] = COMMON_COLOR_NAMES.get(hex_value, hex_value)
            data["font_color_hex"] = hex_value
    size = r_pr.find(qn("w:sz"))
    if size is not None:
        value = size.get(qn("w:val"))
        if value:
            try:
                data["size_pt"] = round(float(value) / 2, 2)
            except ValueError:
                pass
    for field_name, tag_name in {
        "bold": "b",
        "italic": "i",
        "all_caps": "caps",
        "small_caps": "smallCaps",
    }.items():
        value = xml_toggle_value(r_pr, tag_name)
        if value is not None:
            data[field_name] = value
    return data


def east_asia_font_name(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return sanitize_east_asia_font(r_pr.rFonts.get(qn("w:eastAsia"))) or sanitize_east_asia_font(r_pr.rFonts.get(qn("w:cs")))


def font_color_signature(run) -> tuple[str, str]:
    color = run.font.color
    if color.type == MSO_COLOR_TYPE.RGB and color.rgb is not None:
        hex_value = f"#{str(color.rgb).upper()}"
        return COMMON_COLOR_NAMES.get(hex_value, hex_value), hex_value
    if color.type == MSO_COLOR_TYPE.AUTO:
        return "自动（通常黑色）", "#000000"
    if color.type == MSO_COLOR_TYPE.THEME:
        return "主题色（按黑色套用）", "#000000"
    return "自动（通常黑色）", "#000000"


def tab_leader_name(value) -> str:
    mapping = {
        WD_TAB_LEADER.SPACES: "none",
        WD_TAB_LEADER.DOTS: "dots",
        WD_TAB_LEADER.DASHES: "dashes",
        WD_TAB_LEADER.LINES: "lines",
        WD_TAB_LEADER.HEAVY: "heavy",
        WD_TAB_LEADER.MIDDLE_DOT: "middle_dot",
    }
    return mapping.get(value, "none")


def tab_signature(paragraph) -> dict[str, Any]:
    tab_info = {
        "toc_page_number_right_aligned": None,
        "toc_leader": None,
        "toc_right_tab_pt": None,
    }
    try:
        tab_stops = list(paragraph.paragraph_format.tab_stops)
    except TypeError:
        tab_stops = []
    right_tabs = [tab for tab in tab_stops if tab.alignment == WD_TAB_ALIGNMENT.RIGHT]
    if right_tabs:
        tab = right_tabs[-1]
        tab_info["toc_page_number_right_aligned"] = True
        tab_info["toc_leader"] = tab_leader_name(tab.leader)
        tab_info["toc_right_tab_pt"] = pt(tab.position)
    return tab_info


def has_image(paragraph) -> bool:
    xml = paragraph._p.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def image_signature(paragraph) -> dict[str, Any] | None:
    if not has_image(paragraph):
        return None
    root = etree.fromstring(paragraph._p.xml.encode("utf-8"))
    ns = {
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }
    layout = "环绕式" if root.xpath(".//wp:anchor", namespaces=ns) else "嵌入式"
    extent = root.xpath("(.//wp:inline/wp:extent | .//wp:anchor/wp:extent)[1]", namespaces=ns)
    width_pt = height_pt = None
    if extent:
        cx = extent[0].get("cx")
        cy = extent[0].get("cy")
        if cx and cy:
            width_pt = round(int(cx) / 12700, 2)
            height_pt = round(int(cy) / 12700, 2)
    return {
        "image_layout": layout,
        "image_align": align_name(paragraph.alignment) or "center",
        "image_width_pt": width_pt,
        "image_height_pt": height_pt,
    }


def table_alignment_name(value) -> str | None:
    return {
        WD_TABLE_ALIGNMENT.LEFT: "left",
        WD_TABLE_ALIGNMENT.CENTER: "center",
        WD_TABLE_ALIGNMENT.RIGHT: "right",
    }.get(value)


def normalized_table_alignment(value: str | None) -> str:
    # Legacy .doc conversion often writes visually centered thesis tables as left-aligned XML.
    return "center" if value in {None, "left", "center"} else value


def twips_to_pt(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return round(int(value) / 20, 2)
    except ValueError:
        return None


def table_cell_margins(table) -> dict[str, float | None]:
    margins = {
        "cell_margin_top_pt": None,
        "cell_margin_bottom_pt": None,
        "cell_margin_left_pt": None,
        "cell_margin_right_pt": None,
    }
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar") if tbl_pr is not None else None
    if cell_mar is None:
        return margins
    for side, key in {
        "top": "cell_margin_top_pt",
        "bottom": "cell_margin_bottom_pt",
        "left": "cell_margin_left_pt",
        "right": "cell_margin_right_pt",
    }.items():
        node = cell_mar.find(qn(f"w:{side}"))
        if node is not None:
            margins[key] = twips_to_pt(node.get(qn("w:w")))
    return margins


def table_signature(table) -> dict[str, Any]:
    signature = {
        "table_alignment": normalized_table_alignment(table_alignment_name(table.alignment)),
        "table_text_wrapping": "none",
        "table_layout": None,
        "table_width_pct": 100.0,
        "table_width_pt": None,
        "sample_text": "表格整体设置",
    }
    tbl_pr = table._tbl.tblPr
    tblp_pr = tbl_pr.first_child_found_in("w:tblpPr") if tbl_pr is not None else None
    if tblp_pr is not None:
        signature["table_text_wrapping"] = "around"
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout") if tbl_pr is not None else None
    if tbl_layout is not None:
        signature["table_layout"] = tbl_layout.get(qn("w:type"))
    tbl_w = tbl_pr.first_child_found_in("w:tblW") if tbl_pr is not None else None
    if tbl_w is not None:
        width_type = tbl_w.get(qn("w:type"))
        width_value = tbl_w.get(qn("w:w"))
        if width_type == "pct" and width_value:
            signature["table_width_pct"] = round(int(width_value) / 50, 2)
        elif width_type == "dxa" and width_value:
            signature["table_width_pct"] = None
            signature["table_width_pt"] = twips_to_pt(width_value)
    signature.update(table_cell_margins(table))
    return signature


def is_numeric_table_value_text(text: str) -> bool:
    stripped = " ".join(text.strip().split())
    if not stripped:
        return False
    if stripped.lower() in {"yes", "no", "n/a", "na"}:
        return True
    if re.fullmatch(r"[-+−]?\(?\d[\d,]*(?:\.\d+)?\)?%?(?:\*+)?", stripped):
        return True
    if re.fullmatch(r"[-+−]?\(?\d[\d,]*(?:\.\d+)?\)?\s*[-–—]\s*[-+−]?\(?\d[\d,]*(?:\.\d+)?\)?", stripped):
        return True
    return False


def table_cell_role(row_idx: int, col_idx: int, text: str) -> str:
    stripped = text.strip()
    if caption_role(stripped) == "table_caption":
        return "table_caption"
    if is_table_footnote_text(stripped):
        return "table_footnote"
    if stripped.startswith("Note:") or stripped.startswith("注：") or stripped.startswith("Source:") or stripped.startswith("资料来源"):
        return "table_note_cell"
    if row_idx == 0:
        return "table_header"
    if col_idx == 0:
        return "table_body_variable"
    if is_numeric_table_value_text(stripped):
        return "table_body_value"
    return "table_body_definition"


def is_table_footnote_text(text: str) -> bool:
    stripped = " ".join(text.strip().split())
    if not stripped:
        return False
    if re.match(r"^(Note|Notes|Source|Sources)\s*[:：]", stripped, re.IGNORECASE):
        return True
    if stripped.startswith(("注：", "注:", "资料来源", "数据来源")):
        return True
    if re.search(r"\b(t|z|standard error|robust standard error)s?\s+statistics?\s+in\s+parentheses", stripped, re.IGNORECASE):
        return True
    if re.match(r"^\*+\s*p\s*[<≤]", stripped, re.IGNORECASE):
        return True
    if re.search(r"\*\*\*\s*p\s*[<≤]|\*\*\s*p\s*[<≤]|\*\s*p\s*[<≤]", stripped, re.IGNORECASE):
        return True
    return False


def looks_like_table_body_start(text: str) -> bool:
    return bool(re.match(r"^\d+([.)])?$", text.strip()))


def looks_like_section_or_caption(text: str) -> bool:
    if not text:
        return True
    if caption_role(text):
        return True
    if re.match(r"^Chapter\s+\d+\b", text, re.I) or re.match(r"^\d+\.\d+", text):
        return True
    if text in {"References", "参考文献", "Acknowledgements", "Acknowledgments", "致谢", "致  谢"}:
        return True
    return False


def add_flattened_table_observations(doc: Document, observations: dict[str, list[dict[str, Any]]]) -> None:
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if caption_role(text) != "table_caption":
            continue
        observations["table_overall"].append(
            {
                "table_alignment": "center",
                "table_text_wrapping": "none",
                "table_layout": "fixed",
                "table_width_pct": 100.0,
                "table_width_pt": None,
                "cell_margin_top_pt": 2.5,
                "cell_margin_bottom_pt": 2.5,
                "cell_margin_left_pt": 5.0,
                "cell_margin_right_pt": 5.0,
                "sample_text": "由表题后的扁平化表格内容推断",
            }
        )
        header_open = True
        for next_paragraph in paragraphs[index + 1 :]:
            next_text = next_paragraph.text.strip()
            if looks_like_section_or_caption(next_text):
                break
            if is_table_footnote_text(next_text):
                role = "table_footnote"
            elif header_open and not looks_like_table_body_start(next_text):
                role = "table_header"
            elif is_numeric_table_value_text(next_text):
                role = "table_body_value"
                header_open = False
            else:
                role = "table_body_variable"
                header_open = False
            signature = paragraph_signature(next_paragraph)
            signature["sample_text"] = next_text[:100]
            observations[role].append(signature)


def paragraph_signature(paragraph, run_override=None) -> dict[str, Any]:
    run = run_override or first_text_run(paragraph)
    fmt = paragraph.paragraph_format
    first_line_indent = pt(fmt.first_line_indent)
    first_line_pt = first_line_indent if first_line_indent is not None and first_line_indent > 0 else None
    hanging_pt = abs(first_line_indent) if first_line_indent is not None and first_line_indent < 0 else None
    data = {
        "align": align_name(paragraph.alignment) or "justify",
        "line_spacing": normalize_spacing(fmt.line_spacing),
        "before_pt": pt(fmt.space_before) or 0.0,
        "after_pt": pt(fmt.space_after) or 0.0,
        "first_line_pt": first_line_pt,
        "left_pt": pt(fmt.left_indent),
        "hanging_pt": hanging_pt,
        "right_pt": pt(fmt.right_indent),
        "font": LATIN_FONT,
        "east_asia_font": EAST_ASIA_FONT,
        "font_color_name": "自动（通常黑色）",
        "font_color_hex": "#000000",
        "size_pt": None,
        "bold": None,
        "italic": None,
        "all_caps": None,
        "small_caps": None,
    }
    data.update(tab_signature(paragraph))
    if run is not None:
        data["font"] = sanitize_latin_font(run.font.name) or LATIN_FONT
        data["east_asia_font"] = east_asia_font_name(run) or EAST_ASIA_FONT
        data["font_color_name"], data["font_color_hex"] = font_color_signature(run)
        data["size_pt"] = pt(run.font.size)
        data["bold"] = run.font.bold
        data["italic"] = run.font.italic
        data["all_caps"] = run.font.all_caps
        data["small_caps"] = run.font.small_caps
        for key, value in run_element_signature(run._element).items():
            if value is not None:
                data[key] = value
    return data


def paragraph_has_page_field(paragraph) -> bool:
    xml = paragraph._p.xml
    if re.search(r"<w:instrText[^>]*>[^<]*\bPAGE\b", xml):
        return True
    if re.search(r"<w:fldSimple[^>]+w:instr=\"[^\"]*\bPAGE\b", xml):
        return True
    return False


def story_content_paragraphs(story) -> list[Paragraph]:
    paragraphs = []
    for paragraph in story.paragraphs:
        if paragraph.text.strip() or has_image(paragraph) or paragraph_has_page_field(paragraph):
            paragraphs.append(paragraph)
    return paragraphs


def section_page_number_settings(section) -> dict[str, Any]:
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        return {}
    data: dict[str, Any] = {}
    fmt = pg_num.get(qn("w:fmt"))
    start = pg_num.get(qn("w:start"))
    if fmt:
        data["format"] = fmt
    if start:
        try:
            data["start"] = int(start)
        except ValueError:
            pass
    return data


def set_section_page_number_settings(section, number_format: str | None = None, start: int | None = None) -> None:
    if not number_format and start is None:
        return
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    if number_format:
        pg_num.set(qn("w:fmt"), number_format)
    if start is not None:
        pg_num.set(qn("w:start"), str(int(start)))


def header_footer_stories(section) -> list[tuple[str, str, Any]]:
    return [
        ("header", "default", section.header),
        ("header", "first", section.first_page_header),
        ("header", "even", section.even_page_header),
        ("footer", "default", section.footer),
        ("footer", "first", section.first_page_footer),
        ("footer", "even", section.even_page_footer),
    ]


def collect_header_footer_observations(doc: Document, profile: dict[str, Any], observations: dict[str, list[dict[str, Any]]]) -> None:
    header_footer = profile.setdefault("header_footer", {})
    if doc.sections:
        first_section = doc.sections[0]
        if first_section.header_distance is not None:
            header_footer["header_distance_in"] = round(first_section.header_distance.inches, 3)
        if first_section.footer_distance is not None:
            header_footer["footer_distance_in"] = round(first_section.footer_distance.inches, 3)
        header_footer["different_first_page"] = any(section.different_first_page_header_footer for section in doc.sections)
        header_footer["odd_even_pages"] = bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False))
        page_settings = [section_page_number_settings(section) for section in doc.sections]
        if page_settings:
            first_settings = page_settings[0]
            body_settings = page_settings[1] if len(page_settings) > 1 else page_settings[0]
            if first_settings.get("format"):
                header_footer["front_page_number_format"] = first_settings["format"]
            if body_settings.get("format"):
                header_footer["body_page_number_format"] = body_settings["format"]
            if first_settings.get("start") is not None:
                header_footer["front_page_number_start"] = first_settings["start"]
            if body_settings.get("start") is not None:
                header_footer["body_page_number_start"] = body_settings["start"]
            header_footer["page_number_section_mode"] = "restart" if any(settings.get("start") is not None for settings in page_settings[1:]) else "continue"

    page_number_found = False
    for section in doc.sections:
        for story_type, _story_variant, story in header_footer_stories(section):
            for paragraph in story_content_paragraphs(story):
                text = paragraph.text.strip()
                signature = paragraph_signature(paragraph)
                signature["sample_text"] = text[:100] if text else ("PAGE" if paragraph_has_page_field(paragraph) else "")
                observations[f"{story_type}_overall"].append(signature)
                if paragraph_has_page_field(paragraph):
                    observations["page_number"].append(signature)
                    header_footer["page_number_enabled"] = True
                    header_footer["page_number_location"] = story_type
                    header_footer["page_number_align"] = signature.get("align") or "center"
                    page_number_found = True
                elif text:
                    observations[f"{story_type}_text"].append(signature)
    if not page_number_found:
        header_footer["page_number_enabled"] = False


THIRD_HEADING_RE = re.compile(r"^\d+\s*\.\s*\d+\s*\.\s*\d+\b")
SECOND_HEADING_RE = re.compile(r"^\d+\s*\.\s*\d+\b(?!\s*\.)")
ARABIC_CHAPTER_HEADING_RE = re.compile(r"^\d+\s+[\w\u4e00-\u9fff]")
CHINESE_CHAPTER_HEADING_RE = re.compile(
    rf"^(?:第[{CHINESE_NUMERAL_CHARS}]+[章节]|[{CHINESE_NUMERAL_CHARS}]+[、.．])\s*\S+"
)
CHINESE_SECOND_HEADING_RE = re.compile(rf"^[（(][{CHINESE_NUMERAL_CHARS}]+[）)]\s*\S+")
CHINESE_THIRD_HEADING_RE = re.compile(r"^\d+\s*[.．、]\s*(?!\d)\S+")
HEADING_BLANK_ROLES = {"chapter", "second", "third"}


def compact_text(text: str) -> str:
    return " ".join(text.strip().split())


def xml_local_name(element) -> str:
    return etree.QName(element).localname


def paragraph_has_math_object(paragraph) -> bool:
    return any(xml_local_name(element) in {"oMath", "oMathPara"} for element in paragraph._p.iter())


def paragraph_math_text(paragraph) -> str:
    parts = []
    for element in paragraph._p.iter():
        if element.tag == f"{{{MATH_NAMESPACE}}}t" and element.text:
            parts.append(element.text)
    return "".join(parts)


def paragraph_text_with_math(paragraph, fallback_text: str | None = None) -> str:
    text_parts = [paragraph_math_text(paragraph)]
    fallback = fallback_text if fallback_text is not None else getattr(paragraph, "text", "")
    if fallback:
        text_parts.append(fallback)
    return compact_text(" ".join(part for part in text_parts if part))


def paragraph_has_equation_field(paragraph) -> bool:
    field_text = " ".join(node.text or "" for node in paragraph._p.iter(qn("w:instrText")))
    return bool(re.search(r"(?:^|\s)EQ(?:\s|$)|MERGEFORMAT", field_text, re.I))


def is_equation_number_text(text: str) -> bool:
    normalized = compact_text(text)
    return bool(EQUATION_NUMBER_RE.match(normalized))


def looks_like_math_expression_text(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized or is_equation_number_text(normalized):
        return False
    expression = EQUATION_TRAILING_NUMBER_RE.sub("", normalized).strip()
    if not expression:
        return False
    if not EQUATION_MATH_SIGNAL_RE.search(expression):
        return False
    prose_words = re.findall(r"[A-Za-z]{4,}", expression)
    if len(prose_words) > 10 and "=" not in expression:
        return False
    return True


def looks_like_text_equation(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized or len(normalized) > 260:
        return False
    if is_equation_number_text(normalized):
        return False
    has_trailing_number = bool(EQUATION_TRAILING_NUMBER_RE.search(normalized))
    return has_trailing_number and looks_like_math_expression_text(normalized)


def equation_role_from_text(text: str) -> str | None:
    normalized = compact_text(text)
    if not normalized:
        return None
    if is_equation_number_text(normalized):
        return "equation_number"
    if re.search(r"(?:^|\s)EQ(?:\s|$)|MERGEFORMAT", normalized, re.I):
        return "equation"
    if looks_like_text_equation(normalized):
        return "equation_text_fallback"
    return None


def equation_role_for_paragraph(paragraph, text: str) -> str | None:
    if paragraph_has_math_object(paragraph):
        math_text = paragraph_math_text(paragraph)
        if looks_like_math_expression_text(math_text):
            return "equation"
        return equation_role_from_text(text) or "equation"
    if paragraph_has_equation_field(paragraph):
        return "equation"
    return equation_role_from_text(text)


def is_plausible_heading_text(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized:
        return False
    if len(normalized) > 80:
        return False
    if normalized.endswith(("。", ".", "；", ";")):
        return False
    if normalized.count("，") + normalized.count(",") >= 2:
        return False
    return True


def is_third_heading_text(text: str) -> bool:
    normalized = compact_text(text)
    if not is_plausible_heading_text(normalized):
        return False
    return bool(THIRD_HEADING_RE.match(normalized) or CHINESE_THIRD_HEADING_RE.match(normalized))


def is_second_heading_text(text: str) -> bool:
    normalized = compact_text(text)
    if not is_plausible_heading_text(normalized):
        return False
    return bool(SECOND_HEADING_RE.match(normalized) or CHINESE_SECOND_HEADING_RE.match(normalized))


def is_chapter_heading_text(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized:
        return False
    explicit_chapter = bool(
        re.match(r"^Chapter\s+\d+\b", normalized, re.I) or CHINESE_CHAPTER_HEADING_RE.match(normalized)
    )
    if explicit_chapter:
        if len(normalized) > 160:
            return False
        if normalized.endswith(("。", "；", ";")):
            return False
        if normalized.count("，") + normalized.count(",") >= 2:
            return False
        return True
    if not is_plausible_heading_text(normalized):
        return False
    return bool(ARABIC_CHAPTER_HEADING_RE.match(normalized))


def contains_page_break_or_section_break(paragraph) -> bool:
    for element in paragraph._p.iter():
        if element.tag == qn("w:br") and element.get(qn("w:type")) == "page":
            return True
        if element.tag == qn("w:sectPr"):
            return True
    return False


def has_page_or_section_break_before(paragraph) -> bool:
    element = paragraph._p.getprevious()
    while element is not None and element.tag == qn("w:p"):
        candidate = paragraph_from_xml(element, paragraph._parent)
        if contains_page_break_or_section_break(candidate):
            return True
        if candidate.text.strip() or has_image(candidate):
            return False
        element = element.getprevious()
    return False


def page_break_paragraph_element():
    page_break_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    page_break_paragraph.append(run)
    return page_break_paragraph


def insert_page_break_before(paragraph) -> None:
    page_break_paragraph = page_break_paragraph_element()
    paragraph._p.addprevious(page_break_paragraph)


def insert_paragraph_after(paragraph, text: str = "") -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = paragraph_from_xml(new_element, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def run_contains_image(run) -> bool:
    xml = run._r.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def clear_run_text(run) -> None:
    for text_element in run._r.iter(qn("w:t")):
        text_element.text = ""


def xml_element_has_image(element) -> bool:
    return "<w:drawing" in element.xml or "<w:pict" in element.xml


def xml_element_text(element) -> str:
    return "".join(text_node.text or "" for text_node in element.iter(qn("w:t"))).strip()


def flatten_image_only_tables(doc: Document) -> int:
    flattened = 0
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:tbl"):
            continue
        cells = list(child.iter(qn("w:tc")))
        if len(cells) != 1:
            continue
        if xml_element_text(child) or not xml_element_has_image(child):
            continue
        image_paragraphs = [p for p in child.iter(qn("w:p")) if xml_element_has_image(p)]
        if not image_paragraphs:
            continue
        for paragraph_element in image_paragraphs:
            child.addprevious(paragraph_element)
        body.remove(child)
        flattened += 1
    return flattened


def split_captioned_image_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text or not has_image(paragraph):
            continue
        role = caption_role(text)
        if role != "figure_caption":
            continue
        for run in paragraph.runs:
            if run_contains_image(run):
                clear_run_text(run)
            else:
                run.text = ""
        insert_paragraph_after(paragraph, text)


def next_nonblank_paragraph(paragraph, max_blank: int = 3) -> tuple[Paragraph | None, list[Paragraph]]:
    blanks: list[Paragraph] = []
    element = paragraph._p.getnext()
    while element is not None and element.tag == qn("w:p"):
        candidate = paragraph_from_xml(element, paragraph._parent)
        if is_removable_blank_paragraph(candidate):
            blanks.append(candidate)
            if len(blanks) > max_blank:
                return None, blanks
            element = element.getnext()
            continue
        return candidate, blanks
    return None, blanks


def nearby_caption_role(paragraphs: list[Paragraph], index: int, expected_role: str, max_blank: int = 3) -> bool:
    for direction in (-1, 1):
        blanks = 0
        probe = index + direction
        while 0 <= probe < len(paragraphs):
            candidate = paragraphs[probe]
            text = compact_text(candidate.text)
            if not text:
                blanks += 1
                if blanks > max_blank:
                    break
                probe += direction
                continue
            if caption_role(text) == expected_role:
                return True
            break
    return False


def move_preceding_figure_captions_below_images(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if caption_role(text) != "figure_caption":
            continue
        next_paragraph, blanks = next_nonblank_paragraph(paragraph)
        if next_paragraph is None or not has_image(next_paragraph):
            continue
        for blank in blanks:
            remove_paragraph(blank)
        next_paragraph._p.addnext(paragraph._p)


def has_preceding_table_caption(table, max_paragraphs: int = 6) -> bool:
    seen_paragraphs = 0
    element = table._tbl.getprevious()
    while element is not None and seen_paragraphs < max_paragraphs:
        if element.tag == qn("w:p"):
            seen_paragraphs += 1
            text = compact_text(paragraph_element_text(element))
            if not text:
                element = element.getprevious()
                continue
            return caption_role(text) == "table_caption"
        if element.tag == qn("w:tbl"):
            return False
        element = element.getprevious()
    return False


def table_compact_text(table, max_chars: int = 2000) -> str:
    parts: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            text = compact_text(cell.text)
            if text:
                parts.append(text)
            if sum(len(part) for part in parts) >= max_chars:
                return compact_text(" ".join(parts))[:max_chars]
    return compact_text(" ".join(parts))[:max_chars]


def table_looks_like_front_cover(table) -> bool:
    text = table_compact_text(table)
    if not text:
        return False
    normalized = text.lower()
    markers = 0
    for keyword in (*FRONT_MATTER_PROTECTION_KEYWORDS, *FRONT_COVER_KEYWORDS):
        if keyword and keyword in text:
            markers += 1
    english_markers = (
        "student id",
        "student no",
        "student number",
        "student name",
        "class:",
        "major:",
        "supervisor",
        "advisor",
        "college",
        "school",
        "university",
        "thesis title",
    )
    markers += sum(1 for keyword in english_markers if keyword in normalized)
    return markers >= 2


def should_skip_table_for_formatting(table, index: int, skip_first_table: bool) -> bool:
    if not skip_first_table or index != 0:
        return False
    if has_preceding_table_caption(table):
        return False
    return table_looks_like_front_cover(table)


def is_front_matter_protection_text(text: str) -> bool:
    normalized = compact_text(text)
    return any(keyword in normalized for keyword in FRONT_MATTER_PROTECTION_KEYWORDS)


def is_front_cover_text(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized:
        return False
    if any(keyword in normalized for keyword in FRONT_COVER_KEYWORDS):
        return True
    return bool(re.search(r"大学|学院", normalized)) and len(normalized) <= 40


def is_academic_start_text(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized:
        return False
    if is_abstract_heading_text(normalized) or re.match(r"^摘\s*要[:：]", normalized) or re.match(r"^Abstract\s*[:：]", normalized, re.I):
        return True
    if KEYWORD_LABEL_RE.match(normalized) or is_contents_heading_text(normalized):
        return True
    return False


def looks_like_title_before_abstract(paragraphs: list[Paragraph], index: int, lookahead: int = 8) -> bool:
    text = compact_text(paragraphs[index].text)
    if not text or len(text) > 120:
        return False
    if re.search(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日", text) or re.search(r"[〇零○一二三四五六七八九十\d]{4}\s*年\s*[〇零○一二三四五六七八九十\d]{1,2}\s*月", text):
        return False
    if is_front_cover_text(text) or is_front_matter_protection_text(text):
        return False
    if re.match(r"^(题目|学院|专业|班级|学号|学生姓名|指导教师)\s*[:：]", text):
        return False
    probe = index + 1
    seen = 0
    while probe < len(paragraphs) and seen < lookahead:
        next_text = compact_text(paragraphs[probe].text)
        probe += 1
        if not next_text:
            continue
        seen += 1
        if re.match(r"^摘\s*要[:：]", next_text) or is_abstract_heading_text(next_text):
            return True
        if is_front_matter_protection_text(next_text) or is_contents_heading_text(next_text):
            return False
    return False


def academic_start_paragraph_index(paragraphs: list[Paragraph], after_index: int) -> int | None:
    for index in range(max(after_index + 1, 0), len(paragraphs)):
        text = compact_text(paragraphs[index].text)
        if not text:
            continue
        if looks_like_title_before_abstract(paragraphs, index):
            return index
        if is_academic_start_text(text):
            return index
    return None


def front_matter_candidate_index(paragraphs: list[Paragraph]) -> int | None:
    last_marker: int | None = None
    cover_markers = 0
    for index, paragraph in enumerate(paragraphs):
        text = compact_text(paragraph.text)
        if not text:
            continue
        if is_front_matter_protection_text(text):
            last_marker = index
        elif index < 60 and is_front_cover_text(text):
            cover_markers += 1
            last_marker = index if last_marker is None else max(last_marker, index)
    if last_marker is None or cover_markers + int(last_marker is not None) < 2:
        return None
    return last_marker


def paragraph_has_section_break(paragraph) -> bool:
    return paragraph._p.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None


def detect_front_matter_protection(doc: Document) -> dict[str, Any]:
    paragraphs = list(doc.paragraphs)
    marker_index = front_matter_candidate_index(paragraphs)
    if marker_index is None:
        return {
            "enabled": False,
            "start_paragraph_index": None,
            "protected_paragraph_indices": [],
            "protected_section_indices": [],
            "protected_paragraph_elements": set(),
            "protected_table_elements": set(),
            "reason": "no front matter protection markers",
        }
    start_index = academic_start_paragraph_index(paragraphs, marker_index)
    if start_index is None or start_index <= 0:
        return {
            "enabled": False,
            "start_paragraph_index": None,
            "protected_paragraph_indices": [],
            "protected_section_indices": [],
            "protected_paragraph_elements": set(),
            "protected_table_elements": set(),
            "reason": "no academic start after front matter markers",
        }

    protected_paragraph_indices: set[int] = set()
    protected_section_indices: set[int] = set()
    protected_paragraph_elements: set[Any] = set()
    protected_table_elements: set[Any] = set()
    body = doc._element.body
    paragraph_index = -1
    table_index = -1
    section_index = 0
    for child in body:
        if child.tag == qn("w:p"):
            paragraph_index += 1
            if paragraph_index < start_index:
                protected_paragraph_indices.add(paragraph_index)
                protected_paragraph_elements.add(child)
            if child.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
                if paragraph_index < start_index:
                    protected_section_indices.add(section_index)
                section_index += 1
        elif child.tag == qn("w:tbl"):
            table_index += 1
            if paragraph_index < start_index:
                protected_table_elements.add(child)

    for table in doc.tables:
        if table._tbl not in protected_table_elements:
            continue
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    protected_paragraph_elements.add(paragraph._p)

    return {
        "enabled": True,
        "start_paragraph_index": start_index,
        "protected_paragraph_indices": sorted(protected_paragraph_indices),
        "protected_section_indices": sorted(protected_section_indices),
        "protected_paragraph_elements": protected_paragraph_elements,
        "protected_table_elements": protected_table_elements,
        "reason": "front cover/declaration/authorization markers before academic content",
    }


def ensure_body_starts_after_toc_page_break(doc: Document) -> None:
    seen_contents = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text in {"CONTENTS", "目  录", "目录"}:
            seen_contents = True
            continue
        if not seen_contents:
            continue
        role = role_for_target(text, paragraph.style.name if paragraph.style else "", False, False)
        if role in {"toc1", "toc2", "toc3"}:
            continue
        if role in {"paper_title", "chapter"}:
            if not has_page_or_section_break_before(paragraph):
                insert_page_break_before(paragraph)
            return


def ensure_contents_starts_on_new_page(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() not in {"CONTENTS", "目  录", "目录"}:
            continue
        if not has_page_or_section_break_before(paragraph):
            insert_page_break_before(paragraph)
        return
    body = doc._element.body
    if body is None:
        return
    for child in list(body):
        text = paragraph_element_text(child).strip()
        if not any(text.startswith(marker) for marker in ("CONTENTS", "目  录", "目录")):
            continue
        element = child.getprevious()
        has_break = False
        while element is not None:
            if element.tag == qn("w:p"):
                candidate = paragraph_from_xml(element, doc._body)
                if contains_page_break_or_section_break(candidate):
                    has_break = True
                    break
                if candidate.text.strip() or has_image(candidate):
                    break
            elif paragraph_element_text(element).strip():
                break
            element = element.getprevious()
        if not has_break:
            child.addprevious(page_break_paragraph_element())
        return


def is_removable_blank_paragraph(paragraph) -> bool:
    return not paragraph.text.strip() and not has_image(paragraph) and not contains_page_break_or_section_break(paragraph)


def adjacent_blank_counts(paragraphs: list[Any], index: int) -> tuple[int, int]:
    before = 0
    probe = index - 1
    while probe >= 0 and is_removable_blank_paragraph(paragraphs[probe]):
        before += 1
        probe -= 1
    after = 0
    probe = index + 1
    while probe < len(paragraphs) and is_removable_blank_paragraph(paragraphs[probe]):
        after += 1
        probe += 1
    return before, after


def is_contents_heading_text(text: str) -> bool:
    return compact_text(text) in CHINESE_TOC_HEADINGS


def is_abstract_heading_text(text: str) -> bool:
    normalized = compact_text(text)
    without_spaces = re.sub(r"\s+", "", normalized).lower()
    return without_spaces in {"摘要", "abstract"}


def infer_format_sample_start_index(paragraphs: list[Any]) -> int:
    saw_reference_format_section = False
    for index, paragraph in enumerate(paragraphs):
        text = compact_text(getattr(paragraph, "text", ""))
        if not text:
            continue
        if any(marker in text for marker in ("参考格式", "参考范文", "参考样式", "论文模板")):
            saw_reference_format_section = True
            continue
        if saw_reference_format_section and is_abstract_heading_text(text):
            return index
    return 0


def looks_like_toc_entry(text: str) -> bool:
    normalized = compact_text(text)
    if not normalized:
        return False
    if caption_role(normalized):
        return False
    if not re.search(r"(?:\t|\.{2,}|-{3,}|…{2,}|\s+)\s*[IVXLCDMivxlcdm\d一二三四五六七八九十]+$", normalized):
        return False
    return bool(re.match(r"^(?:\d+|第|[一二三四五六七八九十]+[、.．]|[（(][一二三四五六七八九十\d]+[）)])", normalized))


def role_for_toc_text(text: str) -> str | None:
    normalized = compact_text(text)
    if not looks_like_toc_entry(normalized):
        return None
    entry = re.sub(r"(?:\t|\.{2,}|-{3,}|…{2,}|\s+)\s*[IVXLCDMivxlcdm\d一二三四五六七八九十]+$", "", normalized).strip()
    if re.match(r"^\d+\s*\.\s*\d+\s*\.\s*\d+\b", entry):
        return "toc3"
    if re.match(r"^\d+\s*\.\s*\d+\b", entry):
        return "toc2"
    if re.match(rf"^[（(][{CHINESE_NUMERAL_CHARS}0-9]+[）)]", entry):
        return "toc2"
    if re.match(r"^\d+\s+", entry) or re.match(rf"^(?:第[{CHINESE_NUMERAL_CHARS}]+[章节]|[{CHINESE_NUMERAL_CHARS}]+[、.．])", entry):
        return "toc1"
    return "toc1"


def role_from_style_name(style_name: str, text: str) -> str | None:
    lower_style = style_name.lower().strip()
    normalized = compact_text(text)
    if lower_style in {"toc 1", "toc1"}:
        return "toc1"
    if lower_style in {"toc 2", "toc2"}:
        return "toc2"
    if lower_style in {"toc 3", "toc3"}:
        return "toc3"
    if not is_plausible_heading_text(normalized):
        return None
    if lower_style in {"heading 1", "标题 1", "标题1", "一级", "chapter"}:
        return "chapter"
    if lower_style in {"heading 2", "标题 2", "标题2", "二级", "section"}:
        return "second"
    if lower_style in {"heading 3", "标题 3", "标题3", "三级", "subsection"}:
        return "third"
    return None


def role_for_sample(text: str, style_name: str, in_references: bool, in_ack: bool, in_toc: bool = False) -> str | None:
    if not text:
        return None
    if is_abstract_heading_text(text):
        return "front_heading"
    if is_contents_heading_text(text):
        return "contents_heading"
    style_role = role_from_style_name(style_name, text)
    if style_role:
        return style_role
    if in_toc:
        toc_role = role_for_toc_text(text)
        if toc_role:
            return toc_role
    if "PAGEREF" in text or text.startswith("HYPERLINK") or text.startswith("TOC "):
        if re.search(r"CHAPTER\s+\d+", text, re.I) or re.search(r"第[一二三四五六七八九十]+章", text):
            return "toc1"
        if re.search(r"\d+\s*\.\s*\d+\s*\.\s*\d+", text):
            return "toc3"
        if re.search(r"\d+\s*\.\s*\d+", text):
            return "toc2"
        return "toc1"
    if is_chapter_heading_text(text):
        return "chapter"
    if is_third_heading_text(text):
        return "third"
    if is_second_heading_text(text):
        return "second"
    if text in {"References", "参考文献"}:
        return "reference_heading"
    if text in {"Acknowledgements", "Acknowledgments", "致谢", "致  谢"}:
        return "ack_heading"
    if text.startswith("The Impact ") or text.startswith("A Study ") or text.startswith("Research on "):
        return "paper_title"
    if text.startswith("Class") or text.startswith("Student ID") or text.startswith("Name:") or "Student ID" in text:
        return "paper_meta"
    caption = caption_role(text)
    if caption:
        return caption
    if KEYWORD_LABEL_RE.match(text):
        return "keywords"
    if is_table_footnote_text(text):
        return "table_footnote"
    if text.startswith("Note:") or text.startswith("注："):
        return "note"
    equation_role = equation_role_from_text(text)
    if equation_role:
        return equation_role
    if in_references:
        return "reference_entry"
    if in_ack:
        return "ack_body"
    if len(text) > 80:
        return "body"
    return None


def english_caption_role(text: str) -> str | None:
    # Accept common thesis caption variants while excluding narrative references
    # like "Figure 3-1 shows..." and "Table 5-1 presents...".
    normalized = compact_text(text)
    match = re.match(r"^(Figure|Table)\s+\d+\s*[-－–—]\s*\d+\s*(.*)$", normalized, re.IGNORECASE)
    if not match:
        return None
    rest = match.group(2).strip()
    rest_without_punct = rest.lstrip(":.．。:：-－–— ").strip()
    if not rest_without_punct:
        return "figure_caption" if match.group(1).lower() == "figure" else "table_caption"
    word = rest_without_punct.split()[0].strip(".,:;!?()[]{}\"'“”‘’")
    if word.lower() in CAPTION_NARRATIVE_WORDS:
        return None
    if bool(word) and (word[0].isupper() or rest[:1] in {":", "：", ".", "．"}):
        return "figure_caption" if match.group(1).lower() == "figure" else "table_caption"
    return None


def chinese_caption_role(text: str) -> str | None:
    normalized = compact_text(text).replace("\u3000", " ")
    match = re.match(r"^(图|表)\s*(\d+)(?:\s*[-－–—.．]\s*(\d+))?\s*(.*)$", normalized)
    if not match:
        return None
    marker = match.group(1)
    rest = match.group(4).strip()
    rest_without_punct = rest.lstrip(":.．。:：、-－–— ").strip()
    if any(rest_without_punct.startswith(prefix) for prefix in CHINESE_CAPTION_NARRATIVE_PREFIXES):
        return None
    if not rest_without_punct:
        return "figure_caption" if marker == "图" else "table_caption"
    return "figure_caption" if marker == "图" else "table_caption"


def is_caption_text(text: str) -> bool:
    return bool(english_caption_role(text) or chinese_caption_role(text))


def caption_role(text: str) -> str | None:
    return english_caption_role(text) or chinese_caption_role(text)


def role_for_target(text: str, style_name: str, in_references: bool, in_ack: bool) -> str:
    role = role_for_sample(text, style_name, in_references, in_ack)
    if role:
        return role
    if in_ack and re.search(r"\d{4}$", text):
        return "signature"
    return "body"


def paragraph_element_text(paragraph_element) -> str:
    return "".join(node.text or "" for node in paragraph_element.iter(qn("w:t")))


def paragraph_style_id(paragraph_element) -> str:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return ""
    return p_style.get(qn("w:val")) or ""


def set_paragraph_style_id(paragraph_element, style_id: str) -> None:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    p_style.set(qn("w:val"), style_id)


def role_for_toc_entry(text: str, style_id: str) -> str | None:
    compact = "".join((text or "").split())
    if re.match(r"^\d+\.\d+\.\d+", compact):
        return "toc3"
    if re.match(r"^\d+\.\d+(?!\.)", compact):
        return "toc2"
    if re.match(r"^(CHAPTER|Chapter)\d+", compact) or compact.startswith(("摘要", "ABSTRACT", "Abstract")):
        return "toc1"
    return TOC_STYLE_ID_TO_ROLE.get((style_id or "").lower())


def is_toc_sdt(sdt_element) -> bool:
    for gallery in sdt_element.iter(qn("w:docPartGallery")):
        if gallery.get(qn("w:val")) == "Table of Contents":
            return True
    return False


def iter_toc_content_paragraphs(doc: Document):
    body = doc._element.body
    if body is None:
        return
    for sdt in body.iter(qn("w:sdt")):
        if not is_toc_sdt(sdt):
            continue
        content = sdt.find(qn("w:sdtContent"))
        if content is None:
            continue
        for paragraph_element in content.iter(qn("w:p")):
            text = paragraph_element_text(paragraph_element).strip()
            role = role_for_toc_entry(text, paragraph_style_id(paragraph_element))
            if role:
                yield Paragraph(paragraph_element, doc._body), role


TOC_DISPLAYED_PAGE_RE = re.compile(r"^(.*?)(\d+|[ivxlcdm]+)\s*$", re.I)
STANDALONE_PAGE_LABEL_RE = re.compile(r"^\s*(\d+|[ivxlcdm]+)\s*$", re.I)


def normalize_rendered_lookup_text(text: str) -> str:
    normalized = (text or "").lower().replace("’", "'").replace("–", "-").replace("—", "-")
    return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", normalized)


def parse_toc_display_entry(paragraph: Paragraph, role: str) -> dict[str, Any] | None:
    text_nodes = list(paragraph._p.iter(qn("w:t")))
    text = "".join(node.text or "" for node in text_nodes).strip()
    if not text:
        return None
    last_text = (text_nodes[-1].text or "").strip() if text_nodes else ""
    if len(text_nodes) > 1 and STANDALONE_PAGE_LABEL_RE.fullmatch(last_text):
        title = "".join(node.text or "" for node in text_nodes[:-1]).strip()
        displayed_page = last_text
    else:
        match = TOC_DISPLAYED_PAGE_RE.match(text)
        if not match:
            return None
        title = match.group(1).rstrip(" \t.．·…-－–—")
        displayed_page = match.group(2)
    if not title:
        return None
    return {
        "paragraph": paragraph,
        "role": role,
        "title": title,
        "displayed_page": displayed_page,
        "text": text,
    }


def rendered_page_label(text: str, physical_page: int) -> str:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    for line in reversed(lines[-6:]):
        if STANDALONE_PAGE_LABEL_RE.fullmatch(line):
            return line
    return str(physical_page)


def rendered_toc_page_indices(entries: list[dict[str, Any]], page_texts: list[str]) -> set[int]:
    normalized_pages = [normalize_rendered_lookup_text(text) for text in page_texts]
    queries = [normalize_rendered_lookup_text(entry["title"]) for entry in entries if entry.get("title")]
    toc_start_pages = []
    for page_index, normalized in enumerate(normalized_pages, start=1):
        matches = sum(1 for query in queries if query and query in normalized)
        has_contents_heading = any(
            normalize_rendered_lookup_text(heading) in normalized
            for heading in CHINESE_TOC_HEADINGS
            if heading
        )
        if has_contents_heading and matches:
            toc_start_pages.append(page_index)
    if not toc_start_pages:
        return set()
    toc_start = toc_start_pages[0]
    chapter_entry = next(
        (
            entry
            for entry in entries
            if is_chapter_heading_text(entry["title"])
        ),
        None,
    )
    if chapter_entry is not None:
        query = normalize_rendered_lookup_text(chapter_entry["title"])
        chapter_pages = [
            page_index
            for page_index, normalized in enumerate(normalized_pages, start=1)
            if query and query in normalized and page_index >= toc_start
        ]
        if len(chapter_pages) >= 2:
            return set(range(toc_start, chapter_pages[-1]))
    toc_pages: set[int] = {toc_start}
    for page_index in range(toc_start + 1, len(normalized_pages) + 1):
        matches = sum(1 for query in queries if query and query in normalized_pages[page_index - 1])
        if not matches:
            break
        toc_pages.add(page_index)
    return toc_pages


def toc_page_result_records(doc: Document, page_texts: list[str]) -> tuple[list[dict[str, Any]], set[int]]:
    parsed_entries = [
        entry
        for paragraph, role in iter_toc_content_paragraphs(doc)
        if (entry := parse_toc_display_entry(paragraph, role)) is not None
    ]
    normalized_pages = [normalize_rendered_lookup_text(text) for text in page_texts]
    toc_pages = rendered_toc_page_indices(parsed_entries, page_texts)
    records: list[dict[str, Any]] = []
    for entry in parsed_entries:
        query = normalize_rendered_lookup_text(entry["title"])
        matched_pages = [
            page_index
            for page_index, text in enumerate(normalized_pages, start=1)
            if query and query in text and page_index not in toc_pages
        ]
        actual_page_index = matched_pages[-1] if matched_pages else None
        actual_displayed_page = (
            rendered_page_label(page_texts[actual_page_index - 1], actual_page_index)
            if actual_page_index is not None
            else None
        )
        records.append(
            {
                **entry,
                "actual_page_index": actual_page_index,
                "actual_displayed_page": actual_displayed_page,
            }
        )
    return records, toc_pages


def check_toc_page_results_against_rendered_text(doc: Document, page_texts: list[str]) -> dict[str, Any]:
    records, toc_pages = toc_page_result_records(doc, page_texts)
    mismatches: list[dict[str, Any]] = []
    for record in records:
        if record["actual_displayed_page"] is None:
            mismatches.append(
                {
                    "code": "toc_heading_page_not_found",
                    "title": record["title"],
                    "displayed_page": record["displayed_page"],
                }
            )
        elif record["displayed_page"].lower() != record["actual_displayed_page"].lower():
            mismatches.append(
                {
                    "code": "toc_displayed_page_mismatch",
                    "title": record["title"],
                    "displayed_page": record["displayed_page"],
                    "actual_displayed_page": record["actual_displayed_page"],
                    "actual_page_index": record["actual_page_index"],
                }
            )
    public_entries = [
        {key: value for key, value in record.items() if key != "paragraph"}
        for record in records
    ]
    return {
        "status": "not-applicable" if not records else ("pass" if not mismatches else "fail"),
        "entry_count": len(records),
        "toc_page_indices": sorted(toc_pages),
        "entries": public_entries,
        "mismatches": mismatches,
    }


def replace_toc_displayed_page(record: dict[str, Any], new_page: str) -> None:
    paragraph = record["paragraph"]
    text_nodes = list(paragraph._p.iter(qn("w:t")))
    if not text_nodes:
        raise RuntimeError(f"No text nodes found for TOC entry: {record['title']}")
    last_node = text_nodes[-1]
    last_text = last_node.text or ""
    if len(text_nodes) > 1 and STANDALONE_PAGE_LABEL_RE.fullmatch(last_text.strip()):
        last_node.text = new_page
        return
    match = TOC_DISPLAYED_PAGE_RE.match(last_text)
    if not match:
        raise RuntimeError(f"Cannot replace displayed page result for TOC entry: {record['text']}")
    last_node.text = f"{match.group(1)}{new_page}"


def refresh_toc_page_results_from_rendered_text(doc: Document, page_texts: list[str]) -> dict[str, Any]:
    records, toc_pages = toc_page_result_records(doc, page_texts)
    changes: list[dict[str, Any]] = []
    unresolved: list[dict[str, Any]] = []
    for record in records:
        actual_page = record["actual_displayed_page"]
        if actual_page is None:
            unresolved.append({"title": record["title"], "displayed_page": record["displayed_page"]})
            continue
        if record["displayed_page"].lower() == actual_page.lower():
            continue
        replace_toc_displayed_page(record, actual_page)
        changes.append(
            {
                "title": record["title"],
                "from": record["displayed_page"],
                "to": actual_page,
                "actual_page_index": record["actual_page_index"],
            }
        )
    return {
        "status": "not-applicable" if not records else ("pass" if not unresolved else "fail"),
        "entry_count": len(records),
        "toc_page_indices": sorted(toc_pages),
        "updated_count": len(changes),
        "changes": changes,
        "unresolved": unresolved,
    }


def extract_pdf_page_texts(pdf_path: Path) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to validate rendered TOC page results.") from exc
    return [(page.extract_text() or "") for page in PdfReader(str(pdf_path)).pages]


def refresh_toc_page_results_from_pdf(docx_path: Path, pdf_path: Path) -> dict[str, Any]:
    doc = Document(docx_path)
    report = refresh_toc_page_results_from_rendered_text(doc, extract_pdf_page_texts(pdf_path))
    if report["updated_count"]:
        doc.save(docx_path)
    return report


def check_toc_page_results_against_pdf(docx_path: Path, pdf_path: Path) -> dict[str, Any]:
    return check_toc_page_results_against_rendered_text(Document(docx_path), extract_pdf_page_texts(pdf_path))


def write_toc_page_validation_report(report: dict[str, Any], render_dir: Path) -> None:
    json_path = render_dir / "toc-page-validation.json"
    txt_path = render_dir / "toc-page-validation.txt"
    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"TOC page validation: {report['status']}",
        f"TOC entries checked: {report['entry_count']}",
        "TOC rendered pages excluded: " + (
            ", ".join(str(index) for index in report.get("toc_page_indices", [])) or "none"
        ),
        f"Displayed-page mismatches: {len(report.get('mismatches', []))}",
    ]
    for issue in report.get("mismatches", []):
        lines.append(
            f"- {issue['title']}: displayed {issue.get('displayed_page', 'unknown')}, "
            f"rendered heading page {issue.get('actual_displayed_page', 'not found')}"
        )
    txt_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def normalize_docx_alignment_values(input_path: Path, output_path: Path | None = None) -> Path:
    """Map OOXML logical paragraph alignments to values python-docx can read."""
    output_path = output_path or input_path
    replacements = {
        b'w:val="start"': b'w:val="left"',
        b"w:val='start'": b"w:val='left'",
        b'w:val="end"': b'w:val="right"',
        b"w:val='end'": b"w:val='right'",
    }
    changed = False
    with zipfile.ZipFile(input_path, "r") as zin:
        members = []
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename.endswith(".xml"):
                normalized = data
                for old, new in replacements.items():
                    normalized = normalized.replace(old, new)
                if normalized != data:
                    changed = True
                data = normalized
            members.append((info, data))

    if not changed:
        if output_path.resolve() != input_path.resolve():
            shutil.copy2(input_path, output_path)
            return output_path
        return input_path

    destination = output_path
    temp_output = destination.with_suffix(destination.suffix + ".tmp")
    with zipfile.ZipFile(temp_output, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in members:
            zout.writestr(info, data)
    temp_output.replace(destination)
    return destination


LIBREOFFICE_INSTALL_HINT = "Install LibreOffice or set SOFFICE to the soffice executable path."


def convert_legacy_word(input_path: Path, work_dir: Path, soffice: Path | None) -> Path:
    if input_path.suffix.lower() == ".docx":
        normalized = work_dir / f"{input_path.stem}.normalized.docx"
        if normalized.exists() and normalized.stat().st_mtime >= input_path.stat().st_mtime:
            return normalized
        return normalize_docx_alignment_values(input_path, normalized)
    if input_path.suffix.lower() not in {".doc", ".rtf", ".odt"}:
        raise ValueError(f"Unsupported sample format: {input_path.suffix}. Use .doc, .docx, .rtf, or .odt.")
    if soffice is None:
        raise RuntimeError(
            f"LibreOffice is required to convert legacy Word file {input_path} to DOCX. "
            f"For no-LibreOffice trial runs, save the file as .docx first. {LIBREOFFICE_INSTALL_HINT}"
        )
    output = work_dir / f"{input_path.stem}.converted.docx"
    normalized = work_dir / f"{input_path.stem}.normalized.docx"
    if normalized.exists() and normalized.stat().st_mtime >= input_path.stat().st_mtime:
        return normalized
    if output.exists() and output.stat().st_mtime >= input_path.stat().st_mtime:
        return normalize_docx_alignment_values(output, normalized)
    with tempfile.TemporaryDirectory(prefix="legacy_doc_convert_") as convert_dir:
        env = os.environ.copy()
        env["PATH"] = f"{soffice.parent}{os.pathsep}{env.get('PATH', '')}"
        with tempfile.TemporaryDirectory(prefix="soffice_profile_") as profile_dir:
            proc = subprocess.run(
                [
                    str(soffice),
                    f"-env:UserInstallation=file://{profile_dir}",
                    "--headless",
                    "--invisible",
                    "--norestore",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    convert_dir,
                    str(input_path),
                ],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                env=env,
            )
        if proc.returncode != 0:
            raise RuntimeError(
                f"LibreOffice could not convert legacy Word file to DOCX: {input_path}\n"
                f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
            )
        converted = sorted(Path(convert_dir).glob("*.docx"))
        if not converted:
            raise RuntimeError(
                f"LibreOffice did not produce a DOCX for {input_path}\n"
                f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
            )
        if output.exists():
            output.unlink()
        shutil.move(str(converted[0]), output)
    return normalize_docx_alignment_values(output, normalized)


def find_soffice() -> Path | None:
    for candidate in (
        os.environ.get("SOFFICE"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        str(MACOS_SOFFICE) if MACOS_SOFFICE.exists() else None,
    ):
        if not candidate:
            continue
        path = Path(candidate).expanduser()
        if path.exists() and os.access(path, os.X_OK):
            return path.resolve()
    return None


def require_libreoffice() -> Path:
    soffice = find_soffice()
    if soffice is None:
        raise RuntimeError(
            "LibreOffice is required for review/strict visual QA and legacy .doc/.rtf/.odt conversion. "
            f"Use --qa-level fast with .docx files for a no-render trial run, or {LIBREOFFICE_INSTALL_HINT}"
        )
    if not PDF_RENDERER.exists():
        raise RuntimeError(f"PDFKit page renderer not found: {PDF_RENDERER}")
    if shutil.which("swift") is None:
        raise RuntimeError("Swift is required on macOS to rasterize LibreOffice PDFs into per-page PNGs.")
    return soffice


def run_requires_libreoffice(qa_level: str, input_paths: list[Path]) -> bool:
    if qa_level != "fast":
        return True
    return any(path.suffix.lower() in LEGACY_WORD_SUFFIXES for path in input_paths)


def libreoffice_optional_message(input_paths: list[Path]) -> str:
    names = ", ".join(path.name for path in input_paths) if input_paths else "no Word inputs"
    return (
        "LibreOffice: not found; continuing because --qa-level fast skips rendering and inputs are already .docx "
        f"({names}). Rerun with --qa-level strict on a machine with LibreOffice before final delivery."
    )


def emit(message: str, quiet: bool = False, important: bool = False) -> None:
    if important or not quiet:
        print(message)


def qa_policy(qa_level: str, render_width: int, render_height: int) -> dict[str, Any]:
    if qa_level == "fast":
        return {
            "qa_level": "fast",
            "render_enabled": False,
            "render_width": render_width,
            "render_height": render_height,
            "edge_overflow_is_error": False,
            "manual_review_required": False,
            "contact_sheet_max_pages": 0,
        }
    if qa_level == "review":
        return {
            "qa_level": "review",
            "render_enabled": True,
            "render_width": min(render_width, 1200),
            "render_height": min(render_height, 1700),
            "edge_overflow_is_error": False,
            "manual_review_required": True,
            "contact_sheet_max_pages": 16,
        }
    return {
        "qa_level": "strict",
        "render_enabled": True,
        "render_width": render_width,
        "render_height": render_height,
        "edge_overflow_is_error": True,
        "manual_review_required": True,
        "contact_sheet_max_pages": 20,
    }


def clean_render_dir(render_dir: Path) -> None:
    render_dir.mkdir(parents=True, exist_ok=True)
    for pattern in ("page-*.png", "*.pdf", "visual-report.*", "visual-risk-report.*", "contact-sheet.*"):
        for path in render_dir.glob(pattern):
            if path.is_file():
                path.unlink()


def page_sort_key(path: Path) -> tuple[int, str]:
    match = re.search(r"page-(\d+)\.png$", path.name)
    return (int(match.group(1)) if match else 0, path.name)


def inspect_rendered_pages(pages: list[Path], edge_overflow_is_error: bool = True) -> dict[str, Any]:
    try:
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("Pillow is required for visual QA page inspection.") from exc

    page_reports = []
    blank_pages = []
    edge_overflow_pages = []
    for page_index, page in enumerate(pages, start=1):
        with Image.open(page) as image:
            gray = image.convert("L")
            hist = gray.histogram()
            total = max(1, gray.width * gray.height)
            ink_ratio = sum(hist[:245]) / total
            dark_ratio = sum(hist[:80]) / total
            width, height = gray.size
            edge_width = max(8, int(width * 0.015))
            top = int(height * 0.08)
            bottom = int(height * 0.92)
            left_edge = gray.crop((0, top, edge_width, bottom))
            right_edge = gray.crop((width - edge_width, top, width, bottom))
            left_hist = left_edge.histogram()
            right_hist = right_edge.histogram()
            left_edge_ink_ratio = sum(left_hist[:245]) / max(1, left_edge.width * left_edge.height)
            right_edge_ink_ratio = sum(right_hist[:245]) / max(1, right_edge.width * right_edge.height)
        page_report = {
            "page_index": page_index,
            "file": str(page),
            "width": width,
            "height": height,
            "size_bytes": page.stat().st_size,
            "ink_ratio": round(ink_ratio, 6),
            "dark_ratio": round(dark_ratio, 6),
            "left_edge_ink_ratio": round(left_edge_ink_ratio, 6),
            "right_edge_ink_ratio": round(right_edge_ink_ratio, 6),
        }
        if width < 1000 or height < 1000:
            raise RuntimeError(f"Rendered page is too small for visual QA: {page} ({width}x{height})")
        if page_report["size_bytes"] < 5000:
            raise RuntimeError(f"Rendered page PNG is unexpectedly small: {page}")
        page_report["blank_like"] = ink_ratio < 0.001
        page_report["edge_overflow"] = left_edge_ink_ratio > 0.003 or right_edge_ink_ratio > 0.003
        page_reports.append(page_report)
        if ink_ratio < 0.001:
            blank_pages.append(str(page))
        if left_edge_ink_ratio > 0.003 or right_edge_ink_ratio > 0.003:
            edge_overflow_pages.append(str(page))

    if len(blank_pages) == len(pages):
        raise RuntimeError("All rendered pages appear blank; visual QA failed.")
    if edge_overflow_pages and edge_overflow_is_error:
        raise RuntimeError(
            "Rendered pages contain ink at the page side edge, which usually means table/image/text overflow: "
            + ", ".join(edge_overflow_pages)
        )

    return {
        "page_count": len(pages),
        "blank_like_pages": blank_pages,
        "edge_overflow_pages": edge_overflow_pages,
        "pages": page_reports,
    }


def select_review_page_indices(page_reports: list[dict[str, Any]], max_pages: int = 16) -> list[int]:
    if not page_reports or max_pages <= 0:
        return []
    page_indices = [int(report.get("page_index") or index) for index, report in enumerate(page_reports, start=1)]
    selected: list[int] = []

    def add(page_index: int) -> None:
        if page_index in page_indices and page_index not in selected and len(selected) < max_pages:
            selected.append(page_index)

    for page_index in page_indices[:3]:
        add(page_index)
    add(page_indices[-1])
    for report in page_reports:
        if report.get("blank_like") or report.get("edge_overflow"):
            add(int(report.get("page_index") or 0))
    if len(selected) < max_pages and len(page_indices) > 4:
        stride = max(1, len(page_indices) // max(1, max_pages - len(selected)))
        for page_index in page_indices[3:-1:stride]:
            add(page_index)
            if len(selected) >= max_pages:
                break
    return sorted(selected)


def create_contact_sheet(pages: list[Path], selected_indices: list[int], output_path: Path) -> Path | None:
    if not selected_indices:
        return None
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return None

    selected_pages = [pages[index - 1] for index in selected_indices if 1 <= index <= len(pages)]
    if not selected_pages:
        return None

    thumb_width = 260
    label_height = 28
    gutter = 18
    columns = min(4, len(selected_pages))
    thumbs = []
    for page_index, page in zip(selected_indices, selected_pages):
        with Image.open(page) as image:
            image = image.convert("RGB")
            ratio = thumb_width / max(1, image.width)
            thumb_height = max(1, int(image.height * ratio))
            thumb = image.resize((thumb_width, thumb_height))
        thumbs.append((page_index, thumb))

    max_thumb_height = max(thumb.height for _, thumb in thumbs)
    rows = (len(thumbs) + columns - 1) // columns
    sheet_width = columns * thumb_width + (columns + 1) * gutter
    sheet_height = rows * (max_thumb_height + label_height) + (rows + 1) * gutter
    sheet = Image.new("RGB", (sheet_width, sheet_height), "white")
    draw = ImageDraw.Draw(sheet)
    for item_index, (page_index, thumb) in enumerate(thumbs):
        row = item_index // columns
        col = item_index % columns
        x = gutter + col * (thumb_width + gutter)
        y = gutter + row * (max_thumb_height + label_height + gutter)
        draw.text((x, y), f"page {page_index}", fill="black")
        sheet.paste(thumb, (x, y + label_height))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(output_path)
    return output_path


def write_visual_risk_report(report: dict[str, Any], render_dir: Path, selected_pages: list[int], contact_sheet: Path | None) -> None:
    risk_report = {
        "qa_level": report.get("qa_level"),
        "page_count": report.get("page_count", 0),
        "selected_review_pages": selected_pages,
        "blank_like_pages": report.get("blank_like_pages", []),
        "edge_overflow_pages": report.get("edge_overflow_pages", []),
        "contact_sheet": str(contact_sheet) if contact_sheet else None,
    }
    (render_dir / "visual-risk-report.json").write_text(json.dumps(risk_report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"QA level: {risk_report['qa_level']}",
        f"Rendered pages: {risk_report['page_count']}",
        "Selected review pages: " + (", ".join(str(page) for page in selected_pages) if selected_pages else "none"),
        "Blank-like pages: " + (", ".join(risk_report["blank_like_pages"]) if risk_report["blank_like_pages"] else "none"),
        "Edge-overflow risk pages: " + (", ".join(risk_report["edge_overflow_pages"]) if risk_report["edge_overflow_pages"] else "none"),
        "Contact sheet: " + (str(contact_sheet) if contact_sheet else "not generated"),
    ]
    (render_dir / "visual-risk-report.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_environment(soffice: Path) -> dict[str, str]:
    env = os.environ.copy()
    env["PATH"] = f"{soffice.parent}{os.pathsep}{env.get('PATH', '')}"
    if sys.platform == "darwin" and Path("/private/tmp").exists():
        env["TMPDIR"] = "/private/tmp"
        env["TEMP"] = "/private/tmp"
        env["TMP"] = "/private/tmp"
    return env


def render_pdf_only(docx_path: Path, render_dir: Path, soffice: Path) -> dict[str, Any]:
    clean_render_dir(render_dir)
    env = render_environment(soffice)

    with tempfile.TemporaryDirectory(prefix="soffice_profile_") as profile_dir:
        convert_cmd = [
            str(soffice),
            f"-env:UserInstallation=file://{profile_dir}",
            "--invisible",
            "--headless",
            "--norestore",
            "--convert-to",
            "pdf",
            "--outdir",
            str(render_dir),
            str(docx_path),
        ]
        convert_proc = subprocess.run(convert_cmd, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)
    if convert_proc.returncode != 0:
        raise RuntimeError(
            f"LibreOffice PDF conversion failed for {docx_path}.\n"
            f"Command: {' '.join(convert_cmd)}\nSTDOUT:\n{convert_proc.stdout}\nSTDERR:\n{convert_proc.stderr}"
        )

    pdf_path = render_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        pdfs = sorted(render_dir.glob("*.pdf"))
        if not pdfs:
            raise RuntimeError(f"LibreOffice did not produce a PDF for {docx_path}")
        pdf_path = pdfs[0]
    return {
        "pdf": str(pdf_path),
        "libreoffice_stdout_tail": convert_proc.stdout[-2000:],
        "libreoffice_stderr_tail": convert_proc.stderr[-2000:],
    }


def render_pdf_pages_for_visual_qa(
    docx_path: Path,
    pdf_path: Path,
    render_dir: Path,
    soffice: Path,
    label: str,
    width: int,
    height: int,
    qa_level: str = "strict",
    edge_overflow_is_error: bool = True,
    contact_sheet_max_pages: int = 20,
    pdf_report: dict[str, Any] | None = None,
) -> dict[str, Any]:
    env = render_environment(soffice)
    raster_cmd = ["swift", str(PDF_RENDERER), str(pdf_path), str(render_dir), str(width), str(height)]
    raster_proc = subprocess.run(raster_cmd, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, env=env)
    if raster_proc.returncode != 0:
        raise RuntimeError(
            f"PDFKit page rasterization failed for {pdf_path}.\n"
            f"Command: {' '.join(raster_cmd)}\nSTDOUT:\n{raster_proc.stdout}\nSTDERR:\n{raster_proc.stderr}"
        )

    pages = sorted(render_dir.glob("page-*.png"), key=page_sort_key)
    if not pages:
        raise RuntimeError(f"LibreOffice rendered no page PNGs for {docx_path}")

    report = {
        "label": label,
        "docx": str(docx_path),
        "render_dir": str(render_dir),
        "renderer": str(PDF_RENDERER),
        "soffice": str(soffice),
        "pdf": str(pdf_path),
        "qa_level": qa_level,
        "manual_review_required": True,
        "visual_review_gate": (
            "strict full-page visual gate" if qa_level == "strict" else
            "review contact sheet and risk pages before delivery; rerun strict for final handoff"
        ),
        "manual_review_note": (
            "Open the rendered page-*.png files and inspect layout, TOC, tables, figures, comments, "
            "spacing, overlap, blank pages, and page breaks. If any issue is found, fix the responsible "
            "formatting/document logic and rerun rendering before delivery."
        ),
        "libreoffice_stdout_tail": (pdf_report or {}).get("libreoffice_stdout_tail", ""),
        "libreoffice_stderr_tail": (pdf_report or {}).get("libreoffice_stderr_tail", ""),
        "raster_stdout_tail": raster_proc.stdout[-2000:],
        "raster_stderr_tail": raster_proc.stderr[-2000:],
    }
    report.update(inspect_rendered_pages(pages, edge_overflow_is_error=edge_overflow_is_error))
    selected_pages = select_review_page_indices(report.get("pages", []), max_pages=contact_sheet_max_pages)
    contact_sheet = create_contact_sheet(pages, selected_pages, render_dir / "contact-sheet.png")
    if contact_sheet:
        report["contact_sheet"] = str(contact_sheet)
    report["selected_review_pages"] = selected_pages
    write_visual_risk_report(report, render_dir, selected_pages, contact_sheet)

    (render_dir / "visual-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"Visual QA label: {label}",
        f"QA level: {qa_level}",
        f"DOCX: {docx_path}",
        f"Rendered pages: {len(pages)}",
        f"Render directory: {render_dir}",
        f"PDF: {report['pdf']}",
        "Manual review required: yes",
        "Selected review pages: " + (", ".join(str(page) for page in selected_pages) if selected_pages else "none"),
        "Contact sheet: " + (str(contact_sheet) if contact_sheet else "not generated"),
        "Visual gate: if any page fails inspection, fix the document/format logic and rerender until it passes.",
    ]
    if report["blank_like_pages"]:
        lines.append("Blank-like pages to inspect: " + ", ".join(report["blank_like_pages"]))
    if report["edge_overflow_pages"]:
        lines.append("Edge-overflow risk pages to inspect: " + ", ".join(report["edge_overflow_pages"]))
    (render_dir / "visual-report.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def render_for_visual_qa(
    docx_path: Path,
    render_dir: Path,
    soffice: Path,
    label: str,
    width: int,
    height: int,
    qa_level: str = "strict",
    edge_overflow_is_error: bool = True,
    contact_sheet_max_pages: int = 20,
) -> dict[str, Any]:
    pdf_report = render_pdf_only(docx_path, render_dir, soffice)
    return render_pdf_pages_for_visual_qa(
        docx_path,
        Path(pdf_report["pdf"]),
        render_dir,
        soffice,
        label,
        width,
        height,
        qa_level=qa_level,
        edge_overflow_is_error=edge_overflow_is_error,
        contact_sheet_max_pages=contact_sheet_max_pages,
        pdf_report=pdf_report,
    )


def write_fast_visual_qa_report(docx_path: Path, render_dir: Path, soffice: Path | None, label: str) -> dict[str, Any]:
    clean_render_dir(render_dir)
    report = {
        "label": label,
        "qa_level": "fast",
        "docx": str(docx_path),
        "render_dir": str(render_dir),
        "soffice": str(soffice) if soffice else None,
        "pdf": None,
        "page_count": 0,
        "blank_like_pages": [],
        "edge_overflow_pages": [],
        "pages": [],
        "selected_review_pages": [],
        "contact_sheet": None,
        "manual_review_required": False,
        "visual_review_gate": "fast mode skips LibreOffice rendering; rerun with --qa-level review or strict before final delivery",
        "manual_review_note": "Fast QA validates document generation but intentionally skips page rendering.",
        "dependency_note": (
            "LibreOffice was unavailable, so this run is a no-render trial."
            if soffice is None
            else "LibreOffice was available, but fast QA intentionally skipped rendering."
        ),
    }
    (render_dir / "visual-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    write_visual_risk_report(report, render_dir, [], None)
    lines = [
        f"Visual QA label: {label}",
        "QA level: fast",
        f"DOCX: {docx_path}",
        f"Render directory: {render_dir}",
        "Rendered pages: 0",
        "Manual review required: no",
        "LibreOffice: " + (str(soffice) if soffice else "not available"),
        "Visual rendering skipped. Rerun with --qa-level review or --qa-level strict before final delivery.",
    ]
    (render_dir / "visual-report.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def render_output_with_toc_validation(
    output: Path,
    output_render_dir: Path,
    soffice: Path,
    policy: dict[str, Any],
    qa_level: str,
    expected_comments: int,
    comment_mode: str,
    tracker: PerformanceTracker | None = None,
) -> tuple[dict[str, Any], dict[str, Any]]:
    width = int(policy["render_width"])
    height = int(policy["render_height"])
    if tracker:
        with tracker.measure("output_pagination_pdf"):
            pdf_report = render_pdf_only(output.resolve(), output_render_dir, soffice)
        tracker.increment("pdf_conversion_passes")
    else:
        pdf_report = render_pdf_only(output.resolve(), output_render_dir, soffice)

    if tracker:
        with tracker.measure("toc_page_refresh"):
            toc_refresh_report = refresh_toc_page_results_from_pdf(output.resolve(), Path(pdf_report["pdf"]))
        tracker.set_metric("toc_entries_checked", int(toc_refresh_report.get("entry_count", 0)))
        tracker.set_metric("toc_refreshed_entries", int(toc_refresh_report.get("updated_count", 0)))
    else:
        toc_refresh_report = refresh_toc_page_results_from_pdf(output.resolve(), Path(pdf_report["pdf"]))

    if toc_refresh_report["updated_count"]:
        validate_docx_output(output.resolve(), expected_comments, comment_mode)
        if tracker:
            with tracker.measure("output_final_visual_qa"):
                output_report = render_for_visual_qa(
                    output.resolve(),
                    output_render_dir,
                    soffice,
                    "formatted-output-toc-refreshed",
                    width,
                    height,
                    qa_level=qa_level,
                    edge_overflow_is_error=bool(policy["edge_overflow_is_error"]),
                    contact_sheet_max_pages=int(policy["contact_sheet_max_pages"]),
                )
            tracker.increment("pdf_conversion_passes")
            tracker.increment("png_raster_passes")
        else:
            output_report = render_for_visual_qa(
                output.resolve(),
                output_render_dir,
                soffice,
                "formatted-output-toc-refreshed",
                width,
                height,
                qa_level=qa_level,
                edge_overflow_is_error=bool(policy["edge_overflow_is_error"]),
                contact_sheet_max_pages=int(policy["contact_sheet_max_pages"]),
            )
    else:
        if tracker:
            with tracker.measure("output_final_visual_qa"):
                output_report = render_pdf_pages_for_visual_qa(
                    output.resolve(),
                    Path(pdf_report["pdf"]),
                    output_render_dir,
                    soffice,
                    "formatted-output",
                    width,
                    height,
                    qa_level=qa_level,
                    edge_overflow_is_error=bool(policy["edge_overflow_is_error"]),
                    contact_sheet_max_pages=int(policy["contact_sheet_max_pages"]),
                    pdf_report=pdf_report,
                )
            tracker.increment("png_raster_passes")
        else:
            output_report = render_pdf_pages_for_visual_qa(
                output.resolve(),
                Path(pdf_report["pdf"]),
                output_render_dir,
                soffice,
                "formatted-output",
                width,
                height,
                qa_level=qa_level,
                edge_overflow_is_error=bool(policy["edge_overflow_is_error"]),
                contact_sheet_max_pages=int(policy["contact_sheet_max_pages"]),
                pdf_report=pdf_report,
            )

    if tracker:
        with tracker.measure("toc_page_validation"):
            toc_validation_report = check_toc_page_results_against_pdf(output.resolve(), Path(output_report["pdf"]))
    else:
        toc_validation_report = check_toc_page_results_against_pdf(output.resolve(), Path(output_report["pdf"]))
    toc_validation_report["refreshed_count"] = toc_refresh_report["updated_count"]
    toc_validation_report["refresh_changes"] = toc_refresh_report.get("changes", [])
    write_toc_page_validation_report(toc_validation_report, output_render_dir)
    if tracker:
        tracker.set_metric("output_rendered_pages", int(output_report.get("page_count", 0)))
    if toc_validation_report["status"] == "fail":
        examples = "; ".join(
            f"{issue['title']}: {issue.get('displayed_page')} -> {issue.get('actual_displayed_page', 'not found')}"
            for issue in toc_validation_report["mismatches"][:5]
        )
        raise RuntimeError("TOC displayed page validation failed after formatting: " + examples)
    return output_report, toc_validation_report


def validate_docx_output(docx_path: Path, expected_comments: int, comment_mode: str) -> None:
    with zipfile.ZipFile(docx_path) as zf:
        bad_member = zf.testzip()
        if bad_member:
            raise RuntimeError(f"DOCX zip validation failed at {bad_member}")
        names = set(zf.namelist())
    Document(docx_path)
    if comment_mode != "none" and expected_comments > 0 and "word/comments.xml" not in names:
        raise RuntimeError("Format-change comments were expected, but word/comments.xml was not found.")


def cleanup_after_successful_delivery(candidates: list[Path], protected: list[Path]) -> list[Path]:
    protected_paths = [path.resolve() for path in protected]
    removed: list[Path] = []
    seen: set[Path] = set()
    for candidate in candidates:
        path = candidate.resolve()
        if path in seen:
            continue
        seen.add(path)
        if any(
            path == protected_path
            or (path.is_dir() and path in protected_path.parents)
            or (protected_path.is_dir() and protected_path in path.parents)
            for protected_path in protected_paths
        ):
            continue
        if not path.exists():
            continue
        if path.is_dir():
            shutil.rmtree(path)
        else:
            path.unlink()
        removed.append(path)
    return removed


def delivery_cleanup_candidates(
    args: argparse.Namespace,
    target_docx: Path,
    profile_json: Path,
    lint_txt: Path,
    lint_json: Path,
) -> list[Path]:
    candidates = [target_docx, profile_json, lint_txt, lint_json]
    if args.format_table:
        candidates.append(args.format_table)
    target_stem = args.target.stem if args.target else target_docx.stem
    candidates.extend(
        [
            args.analysis_dir / f"{target_stem}.converted.docx",
            args.analysis_dir / f"{target_stem}.normalized.docx",
        ]
    )
    generated_docx_patterns = ("*.converted.docx", "*.normalized.docx", "*.textutil-fallback.docx")
    for pattern in generated_docx_patterns:
        candidates.extend(args.analysis_dir.resolve().glob(pattern))
    if args.reuse_profile:
        reuse_profile = args.reuse_profile.resolve()
        reuse_dir = reuse_profile.parent
        candidates.extend(
            [
                reuse_profile,
                reuse_dir / f"{PROFILE_LINT_REPORT_BASENAME}.txt",
                reuse_dir / f"{PROFILE_LINT_REPORT_BASENAME}.json",
                reuse_dir / "run-manifest.json",
            ]
        )
        for pattern in generated_docx_patterns:
            candidates.extend(reuse_dir.glob(pattern))
        manifest_path = reuse_dir / "run-manifest.json"
        if manifest_path.exists():
            try:
                manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                for key in ("sample_docx", "profile_json"):
                    value = manifest.get(key)
                    if value:
                        generated_path = Path(value)
                        candidates.append(generated_path)
                        if generated_path.name.endswith(".normalized.docx"):
                            candidates.append(
                                generated_path.with_name(
                                    generated_path.name.removesuffix(".normalized.docx") + ".converted.docx"
                                )
                            )
                visual_report = manifest.get("visual_report")
                if visual_report:
                    candidates.append(Path(visual_report).parent)
            except (OSError, ValueError, TypeError):
                pass
    return candidates


def validate_heading_formatting(doc: Document, profile: dict[str, Any], enabled_roles: set[str]) -> None:
    issues: list[str] = []
    reached_main_text = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if "\t" in text:
            continue
        if re.match(r"^Chapter\s+\d+\b", text, re.I) or re.match(r"^第[一二三四五六七八九十]+章", text):
            reached_main_text = True
            role = "chapter"
        elif not reached_main_text:
            continue
        elif is_third_heading_text(text):
            role = "third"
        elif is_second_heading_text(text):
            role = "second"
        else:
            continue
        if role not in enabled_roles:
            continue
        fmt = profile["roles"].get(role)
        if not fmt:
            continue
        signature = paragraph_signature(paragraph)
        signature["blank_before"] = len(previous_blank_siblings(paragraph))
        signature["blank_after"] = len(next_blank_siblings(paragraph))
        target = target_signature(fmt, profile)
        for field_name in ("align", "first_line_pt", "left_pt", "hanging_pt", "right_pt", "blank_before", "blank_after"):
            if field_name not in target:
                continue
            if equivalent_format_value(field_name, signature.get(field_name), target[field_name]):
                continue
            issues.append(
                f"{text[:60]}: {FIELD_LABELS.get(field_name, field_name)} "
                f"{format_value(field_name, signature.get(field_name))} != {format_value(field_name, target[field_name])}"
            )
            break
    if issues:
        raise RuntimeError("Heading structural QA failed; possible missed heading indentation/blank-line formatting:\n- " + "\n- ".join(issues[:12]))


def reconcile_legacy_toc_visible_booleans(signature: dict[str, Any], role: str) -> None:
    if role not in {"toc1", "toc2", "toc3"}:
        return
    # Legacy .doc TOC conversion can emit caps/smallCaps/italic OOXML markers
    # that do not match the rendered contents page. Keep bold, but require the
    # confirmation table to opt back into these visual toggles.
    for field_name in LEGACY_TOC_UNRELIABLE_BOOL_FIELDS:
        signature[field_name] = False


def update_profile_from_observations(profile: dict[str, Any], sample_docx: Path, legacy_source: bool = False) -> dict[str, Any]:
    doc = Document(sample_docx)
    profile["document_language"] = infer_document_language_profile(doc)
    observations: dict[str, list[dict[str, Any]]] = defaultdict(list)
    image_observations: list[dict[str, Any]] = []
    captioned_image_observations: list[dict[str, Any]] = []
    in_references = False
    in_ack = False
    in_toc_region = False
    last_caption_context: str | None = None
    collect_header_footer_observations(doc, profile, observations)

    paragraphs = doc.paragraphs
    sample_start_index = infer_format_sample_start_index(paragraphs)
    profile.setdefault("document_language", {})["sample_start_index"] = sample_start_index
    for paragraph_index, paragraph in enumerate(paragraphs):
        if paragraph_index < sample_start_index:
            continue
        text = paragraph.text.strip()
        if is_contents_heading_text(text):
            in_toc_region = True
        image = image_signature(paragraph)
        if image:
            image_observations.append(image)
            if nearby_caption_role(paragraphs, paragraph_index, "figure_caption"):
                captioned_image_observations.append(image)
        if text in {"References", "参考文献"}:
            in_references = True
            in_ack = False
        elif text in {"Acknowledgements", "Acknowledgments", "致谢", "致  谢"}:
            in_references = False
            in_ack = True
        equation_role = equation_role_for_paragraph(paragraph, text)
        if equation_role:
            role = equation_role
        elif is_table_footnote_text(text) and last_caption_context == "table":
            role = "table_footnote"
        elif text.startswith("Source:") or text.startswith("Note:") or text.startswith("注：") or text.startswith("资料来源"):
            role = f"{last_caption_context}_note" if last_caption_context in {"figure", "table"} else role_for_sample(text, paragraph.style.name if paragraph.style else "", in_references, in_ack, in_toc=in_toc_region)
        else:
            role = role_for_sample(text, paragraph.style.name if paragraph.style else "", in_references, in_ack, in_toc=in_toc_region)
        if in_toc_region and text and role not in {"contents_heading", "toc1", "toc2", "toc3"}:
            in_toc_region = False
        if role:
            if role == "figure_caption":
                last_caption_context = "figure"
            elif role == "table_caption":
                last_caption_context = "table"
            elif role in {"chapter", "second", "third", "body", "reference_heading", "ack_heading"}:
                last_caption_context = None
            if role == "keywords":
                text_runs = [run for run in paragraph.runs if run.text.strip()]
                label_run = text_runs[0] if text_runs else None
                content_run = text_runs[1] if len(text_runs) > 1 else None
                if label_run is not None:
                    label_signature = paragraph_signature(paragraph, label_run)
                    label_signature["sample_text"] = text[:100]
                    observations["keyword_label"].append(label_signature)
                if content_run is not None:
                    content_signature = paragraph_signature(paragraph, content_run)
                    content_signature["sample_text"] = text[:100]
                    observations["keywords"].append(content_signature)
                elif KEYWORD_LABEL_RE.match(text):
                    content_signature = dict(profile["roles"]["keywords"])
                    content_signature["sample_text"] = text[:100]
                    observations["keywords"].append(content_signature)
                continue
            signature = paragraph_signature(paragraph)
            if legacy_source:
                reconcile_legacy_toc_visible_booleans(signature, role)
            if role in HEADING_BLANK_ROLES:
                signature["blank_before"], signature["blank_after"] = adjacent_blank_counts(paragraphs, paragraph_index)
            sample_text = paragraph_text_with_math(paragraph, text) if role in {"equation", "equation_number", "equation_text_fallback"} else text
            signature["sample_text"] = sample_text[:100]
            observations[role].append(signature)

    for paragraph, role in iter_toc_content_paragraphs(doc):
        text = paragraph.text.strip()
        if not text:
            text = paragraph_element_text(paragraph._p).strip()
        signature = paragraph_signature(paragraph)
        xml_run_signature = run_element_signature(first_text_run_element(paragraph._p))
        for key, value in xml_run_signature.items():
            if value is not None:
                signature[key] = value
        if legacy_source:
            reconcile_legacy_toc_visible_booleans(signature, role)
        signature["sample_text"] = text[:100]
        observations[role].append(signature)

    if not doc.tables:
        add_flattened_table_observations(doc, observations)

    tables = list(doc.tables)
    captioned_tables = [table for table in tables if has_preceding_table_caption(table)]
    if captioned_tables:
        tables = captioned_tables
    for table in tables:
        observations["table_overall"].append(table_signature(table))
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    role = table_cell_role(row_idx, col_idx, text)
                    signature = paragraph_signature(paragraph)
                    signature["sample_text"] = text[:100]
                    observations[role].append(signature)

    for role, entries in observations.items():
        current = profile["roles"].get(role)
        if not current:
            continue
        chosen = choose_common_signature(entries)
        if role == "table_overall":
            chosen = normalize_table_overall_choice(chosen)
        if role in {"equation", "equation_number", "equation_text_fallback"}:
            chosen["first_line_pt"] = None
            chosen["hanging_pt"] = None
        # Only trust values that are visible in the sample. Keep defaults when conversion lost detail.
        for key in (
            "size_pt",
            "bold",
            "italic",
            "all_caps",
            "small_caps",
            "align",
            "before_pt",
            "after_pt",
            "blank_before",
            "blank_after",
            "first_line_pt",
            "left_pt",
            "hanging_pt",
            "right_pt",
            "east_asia_font",
            "font_color_name",
            "font_color_hex",
            "toc_page_number_right_aligned",
            "toc_leader",
            "toc_right_tab_pt",
            "table_alignment",
            "table_text_wrapping",
            "table_layout",
            "table_width_pct",
            "table_width_pt",
            "cell_margin_top_pt",
            "cell_margin_bottom_pt",
            "cell_margin_left_pt",
            "cell_margin_right_pt",
        ):
            if chosen.get(key) is not None:
                current[key] = chosen[key]
        if role == "table_overall" and chosen.get("table_width_pt") is None:
            current["table_width_pt"] = None
        if chosen.get("font"):
            current["font"] = chosen["font"]

    if doc.sections:
        section = doc.sections[0]
        # Preserve target page size by default; use sample margins if they look sane.
        margins = {
            "top_in": round(section.top_margin.inches, 3),
            "bottom_in": round(section.bottom_margin.inches, 3),
            "left_in": round(section.left_margin.inches, 3),
            "right_in": round(section.right_margin.inches, 3),
        }
        if all(0.4 <= value <= 2.0 for value in margins.values()):
            profile["page"].update(margins)
        usable_width_pt = round(section.page_width.pt - section.left_margin.pt - section.right_margin.pt, 2)
        if usable_width_pt > 0:
            for role in ("toc1", "toc2", "toc3"):
                if role in observations:
                    profile["roles"][role]["toc_page_number_right_aligned"] = True
                    profile["roles"][role]["toc_leader"] = profile["roles"][role].get("toc_leader") or "dots"
                    profile["roles"][role]["toc_right_tab_pt"] = profile["roles"][role].get("toc_right_tab_pt") or usable_width_pt

    profile["source_observations"] = {
        role: {
            "count": len(entries),
            "examples": [entry["sample_text"] for entry in entries[:5]],
        }
        for role, entries in observations.items()
    }
    preferred_image_observations = captioned_image_observations or image_observations
    if preferred_image_observations:
        profile["image_format"] = choose_common_image_signature(preferred_image_observations)
        profile["image_format"]["count"] = len(preferred_image_observations)
        profile["image_format"]["detected_from"] = "caption_nearby" if captioned_image_observations else "all_images"
    ensure_inferred_table_header_role(profile)
    profile["enabled_roles"] = observed_roles_for_table(profile)
    return profile


def supplement_missing_legacy_roles(profile: dict[str, Any], original_path: Path, work_dir: Path) -> list[str]:
    """Use textutil only as a narrow metadata fallback for legacy .doc samples."""
    if original_path.suffix.lower() not in {".doc", ".rtf"}:
        return []
    observed = set(profile.get("source_observations", {}))
    needed = [role for role in LEGACY_SUPPLEMENT_ROLES if role not in observed]
    if not needed or shutil.which("textutil") is None:
        return []

    fallback_docx = work_dir / f"{original_path.stem}.textutil-fallback.docx"
    proc = subprocess.run(
        ["textutil", "-convert", "docx", "-output", str(fallback_docx), str(original_path)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    if proc.returncode != 0 or not fallback_docx.exists():
        return []

    fallback_docx = normalize_docx_alignment_values(
        fallback_docx,
        work_dir / f"{original_path.stem}.textutil-fallback.normalized.docx",
    )
    fallback_profile = update_profile_from_observations(default_profile(), fallback_docx, legacy_source=True)
    fallback_observed = fallback_profile.get("source_observations", {})
    supplemented: list[str] = []
    for role in needed:
        if role not in fallback_observed:
            continue
        profile["roles"][role] = fallback_profile["roles"][role]
        if role in {"toc1", "toc2", "toc3"}:
            # textutil often exposes legacy TOC fields as italicized hyperlinks even
            # when the rendered LibreOffice sample is upright. Keep TOC supplements
            # structurally useful but do not import that false italic signal.
            profile["roles"][role]["italic"] = False
        profile.setdefault("source_observations", {})[role] = fallback_observed[role]
        supplemented.append(role)
    if supplemented:
        profile["enabled_roles"] = observed_roles_for_table(profile)
    return supplemented


def choose_common_signature(entries: list[dict[str, Any]]) -> dict[str, Any]:
    if not entries:
        return {}
    keys = [
        "size_pt",
        "bold",
        "italic",
        "all_caps",
        "small_caps",
        "align",
        "before_pt",
        "after_pt",
        "blank_before",
        "blank_after",
        "first_line_pt",
        "left_pt",
        "hanging_pt",
        "right_pt",
        "font",
        "east_asia_font",
        "font_color_name",
        "font_color_hex",
        "toc_page_number_right_aligned",
        "toc_leader",
        "toc_right_tab_pt",
        "table_alignment",
        "table_text_wrapping",
        "table_layout",
        "table_width_pct",
        "table_width_pt",
        "cell_margin_top_pt",
        "cell_margin_bottom_pt",
        "cell_margin_left_pt",
        "cell_margin_right_pt",
    ]
    chosen: dict[str, Any] = {}
    for key in keys:
        if key in {"bold", "italic", "all_caps", "small_caps"}:
            values = [entry.get(key) for entry in entries]
        else:
            values = [entry.get(key) for entry in entries if entry.get(key) is not None]
        if not values:
            chosen[key] = None
            continue
        chosen[key] = Counter(values).most_common(1)[0][0]
    return chosen


def normalize_table_overall_choice(chosen: dict[str, Any]) -> dict[str, Any]:
    chosen = dict(chosen)
    chosen["table_alignment"] = normalized_table_alignment(chosen.get("table_alignment"))
    if chosen.get("table_width_pct") is not None:
        chosen["table_width_pt"] = None
    elif chosen.get("table_width_pt") is None:
        chosen["table_width_pct"] = 100.0
    return chosen


def choose_common_image_signature(entries: list[dict[str, Any]]) -> dict[str, Any]:
    if not entries:
        return {}
    chosen: dict[str, Any] = {}
    for key in ("image_layout", "image_align"):
        values = [entry.get(key) for entry in entries if entry.get(key) is not None]
        if values:
            chosen[key] = Counter(values).most_common(1)[0][0]
    for key in ("image_width_pt", "image_height_pt"):
        values = [entry.get(key) for entry in entries if entry.get(key) is not None]
        if values:
            chosen[key] = Counter(values).most_common(1)[0][0]
    return chosen


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def script_sha256() -> str:
    return file_sha256(Path(__file__).resolve())


def sample_cache_manifest_path(analysis_dir: Path) -> Path:
    return analysis_dir / "run-manifest.json"


def sample_cache_metadata(
    sample_path: Path,
    render_width: int,
    render_height: int,
    use_sample_page_size: bool = False,
    qa_level: str = "strict",
) -> dict[str, Any]:
    resolved = sample_path.resolve()
    return {
        "cache_version": CACHE_VERSION,
        "sample_path": str(resolved),
        "sample_sha256": file_sha256(resolved),
        "script_sha256": script_sha256(),
        "render_width": int(render_width),
        "render_height": int(render_height),
        "use_sample_page_size": bool(use_sample_page_size),
        "qa_level": qa_level,
    }


def write_sample_cache_manifest(
    sample_path: Path,
    analysis_dir: Path,
    sample_docx: Path,
    profile_json: Path,
    visual_report: Path | None,
    render_width: int,
    render_height: int,
    use_sample_page_size: bool = False,
    qa_level: str = "strict",
) -> Path:
    analysis_dir.mkdir(parents=True, exist_ok=True)
    data = sample_cache_metadata(sample_path, render_width, render_height, use_sample_page_size, qa_level)
    data.update(
        {
            "sample_docx": str(sample_docx.resolve()),
            "profile_json": str(profile_json.resolve()),
            "visual_report": str(visual_report.resolve()) if visual_report else None,
        }
    )
    path = sample_cache_manifest_path(analysis_dir)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def load_cached_sample_analysis(
    sample_path: Path,
    analysis_dir: Path,
    profile_json: Path,
    render_dir: Path,
    render_width: int,
    render_height: int,
    use_sample_page_size: bool = False,
    qa_level: str = "strict",
) -> dict[str, Any] | None:
    manifest_path = sample_cache_manifest_path(analysis_dir)
    if not manifest_path.exists():
        return None
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        expected = sample_cache_metadata(sample_path, render_width, render_height, use_sample_page_size, qa_level)
    except (OSError, json.JSONDecodeError):
        return None
    for key, value in expected.items():
        if manifest.get(key) != value:
            return None
    cached_profile_json = Path(manifest.get("profile_json") or profile_json)
    cached_sample_docx = Path(manifest.get("sample_docx") or "")
    if not cached_profile_json.exists() or not cached_sample_docx.exists():
        return None
    try:
        profile = json.loads(cached_profile_json.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    visual_report = render_dir / "visual-report.json"
    if not visual_report.exists() and manifest.get("visual_report"):
        visual_report = Path(manifest["visual_report"])
    sample_report = None
    if visual_report.exists():
        try:
            sample_report = json.loads(visual_report.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            sample_report = None
    return {
        "profile": profile,
        "sample_docx": cached_sample_docx,
        "sample_report": sample_report,
        "manifest": manifest,
    }


def shared_sample_cache_entry_dir(
    shared_cache_dir: Path,
    sample_path: Path,
    render_width: int,
    render_height: int,
    use_sample_page_size: bool = False,
    qa_level: str = "strict",
) -> Path:
    metadata = sample_cache_metadata(sample_path, render_width, render_height, use_sample_page_size, qa_level)
    metadata.pop("sample_path", None)
    key = hashlib.sha256(json.dumps(metadata, sort_keys=True).encode("utf-8")).hexdigest()
    return shared_cache_dir.expanduser().resolve() / key


def rewrite_cached_render_paths(value: Any, source_render_dir: Path, cached_render_dir: Path) -> Any:
    if isinstance(value, dict):
        return {key: rewrite_cached_render_paths(item, source_render_dir, cached_render_dir) for key, item in value.items()}
    if isinstance(value, list):
        return [rewrite_cached_render_paths(item, source_render_dir, cached_render_dir) for item in value]
    if isinstance(value, str):
        for source_prefix in (str(source_render_dir), str(source_render_dir.resolve())):
            if value.startswith(source_prefix):
                return str(cached_render_dir.resolve()) + value[len(source_prefix):]
    return value


def publish_shared_sample_analysis(
    sample_path: Path,
    shared_cache_dir: Path,
    sample_docx: Path,
    profile_json: Path,
    sample_render_dir: Path,
    render_width: int,
    render_height: int,
    use_sample_page_size: bool = False,
    qa_level: str = "strict",
) -> Path:
    entry_dir = shared_sample_cache_entry_dir(
        shared_cache_dir,
        sample_path,
        render_width,
        render_height,
        use_sample_page_size,
        qa_level,
    )
    if entry_dir.exists():
        return entry_dir
    entry_dir.parent.mkdir(parents=True, exist_ok=True)
    temporary = entry_dir.with_name(entry_dir.name + ".tmp")
    if temporary.exists():
        shutil.rmtree(temporary)
    temporary.mkdir(parents=True)
    cached_docx = temporary / sample_docx.name
    cached_profile = temporary / profile_json.name
    cached_render_dir = temporary / "visual-qa" / sample_render_dir.name
    shutil.copy2(sample_docx, cached_docx)
    shutil.copy2(profile_json, cached_profile)
    if sample_render_dir.exists():
        shutil.copytree(sample_render_dir, cached_render_dir)
    temporary.replace(entry_dir)
    cached_docx = entry_dir / sample_docx.name
    cached_profile = entry_dir / profile_json.name
    cached_render_dir = entry_dir / "visual-qa" / sample_render_dir.name
    for report_path in cached_render_dir.glob("*.json"):
        try:
            report = json.loads(report_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        report = rewrite_cached_render_paths(report, sample_render_dir, cached_render_dir)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    visual_report = cached_render_dir / "visual-report.json"
    write_sample_cache_manifest(
        sample_path,
        entry_dir,
        cached_docx,
        cached_profile,
        visual_report if visual_report.exists() else None,
        render_width,
        render_height,
        use_sample_page_size,
        qa_level,
    )
    return entry_dir


def load_shared_sample_analysis(
    sample_path: Path,
    shared_cache_dir: Path,
    render_width: int,
    render_height: int,
    use_sample_page_size: bool = False,
    qa_level: str = "strict",
) -> dict[str, Any] | None:
    entry_dir = shared_sample_cache_entry_dir(
        shared_cache_dir,
        sample_path,
        render_width,
        render_height,
        use_sample_page_size,
        qa_level,
    )
    render_dir = entry_dir / "visual-qa" / f"{sample_path.stem}_sample"
    return load_cached_sample_analysis(
        sample_path,
        entry_dir,
        entry_dir / f"{sample_path.stem}.format-profile.json",
        render_dir,
        render_width,
        render_height,
        use_sample_page_size,
        qa_level,
    )


def load_profile_json(profile_path: Path) -> dict[str, Any]:
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise RuntimeError(f"Could not load reused profile JSON: {profile_path}") from exc
    if not isinstance(profile, dict) or "roles" not in profile:
        raise RuntimeError(f"Reused profile JSON does not look like a thesis format profile: {profile_path}")
    return profile


def page_number_token(text: str) -> str:
    return re.sub(r"[\s\-—–]+", "", text or "").strip()


def page_number_kind(text: str) -> str | None:
    token = page_number_token(text)
    if not token:
        return None
    if re.fullmatch(r"\d+", token):
        return "decimal"
    if re.fullmatch(r"[IVXLCDMivxlcdm]+", token):
        return "roman"
    return None


def lint_profile(profile: dict[str, Any]) -> list[dict[str, str]]:
    roles = profile.get("roles", {})
    enabled_roles = set(profile.get("enabled_roles") or roles.keys())
    issues: list[dict[str, str]] = []
    for role in ("figure_caption", "table_caption"):
        fmt = roles.get(role) or {}
        font = normalize_text_cell(fmt.get("font"))
        if role in enabled_roles and (not sanitize_latin_font(font) or sanitize_latin_font(font) != font):
            issues.append(
                {
                    "severity": "warning",
                    "code": "suspicious_caption_latin_font",
                    "role": role,
                    "message": f"{ROLE_LABELS.get(role, role)} 的西文字体看起来异常：{font or '未设置'}；建议使用 Times New Roman。",
                }
            )
    table_overall = roles.get("table_overall") or {}
    if "table_overall" in enabled_roles:
        if table_overall.get("table_alignment") not in (None, "center"):
            issues.append(
                {
                    "severity": "warning",
                    "code": "table_not_centered",
                    "role": "table_overall",
                    "message": "表格整体设置不是居中；正文表格通常应居中。",
                }
            )
        if table_overall.get("table_width_pct") not in (None, 0, 0.0) and table_overall.get("table_width_pt") not in (None, 0, 0.0):
            issues.append(
                {
                    "severity": "warning",
                    "code": "table_width_conflict",
                    "role": "table_overall",
                    "message": "表格整体设置同时包含百分比宽度和固定磅值宽度；建议只保留一种。",
                }
            )
    if "table_body_value" in enabled_roles and "table_header" not in enabled_roles:
        issues.append(
            {
                "severity": "warning",
                "code": "missing_table_header",
                "role": "table_header",
                "message": "已识别表格正文/数值列，但未识别表头（列标题）；请检查确认表是否需要补充表头格式。",
            }
        )
    page_examples = (profile.get("source_observations", {}).get("page_number", {}) or {}).get("examples") or []
    page_kinds = {kind for kind in (page_number_kind(example) for example in page_examples) if kind}
    if {"roman", "decimal"}.issubset(page_kinds):
        issues.append(
            {
                "severity": "warning",
                "code": "mixed_page_number_examples",
                "role": "page_number",
                "message": "页码样本同时包含罗马数字和阿拉伯数字；建议确认前置部分与正文部分的页码格式和分节重启设置。",
            }
        )
    header_footer = profile.get("header_footer", {})
    if header_footer.get("preserve_target_page_numbers") is False:
        issues.append(
            {
                "severity": "warning",
                "code": "target_page_numbers_not_preserved",
                "role": "page_number",
                "message": "保留目标页码为“否”；应用格式时可能改动目标文档页码。",
            }
        )
    return issues


def write_profile_lint_report(profile: dict[str, Any], analysis_dir: Path) -> tuple[Path, Path, list[dict[str, str]]]:
    analysis_dir.mkdir(parents=True, exist_ok=True)
    issues = lint_profile(profile)
    json_path = analysis_dir / f"{PROFILE_LINT_REPORT_BASENAME}.json"
    txt_path = analysis_dir / f"{PROFILE_LINT_REPORT_BASENAME}.txt"
    json_path.write_text(json.dumps({"issues": issues}, ensure_ascii=False, indent=2), encoding="utf-8")
    if issues:
        lines = ["Profile lint found possible issues:"]
        for issue in issues:
            lines.append(f"- [{issue['severity']}] {issue['code']} ({issue['role']}): {issue['message']}")
    else:
        lines = ["Profile lint found no high-risk issues."]
    txt_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return txt_path, json_path, issues


def observed_roles_for_table(profile: dict[str, Any]) -> list[str]:
    observed = set(profile.get("source_observations", {}).keys())
    roles = []
    for role in ROLE_ORDER:
        if (
            role in observed
            or (role == "keyword_label" and "keywords" in observed)
            or (role == "header_text" and "header_overall" in observed)
            or (role == "footer_text" and "footer_overall" in observed)
        ):
            roles.append(role)
    return roles


def ensure_inferred_table_header_role(profile: dict[str, Any]) -> bool:
    observations = profile.setdefault("source_observations", {})
    if "table_header" in observations:
        return False
    body_roles = ["table_body_value", "table_body_definition", "table_body_variable", "table_body"]
    fallback_role = next((role for role in body_roles if role in observations), None)
    if fallback_role is None:
        return False

    roles = profile.setdefault("roles", {})
    header_fmt = dict(roles.get("table_header") or {})
    fallback_fmt = roles.get(fallback_role) or {}
    for key in (
        "font",
        "east_asia_font",
        "font_color_name",
        "font_color_hex",
        "size_pt",
        "align",
        "line_spacing",
        "before_pt",
        "after_pt",
        "first_line_pt",
        "left_pt",
        "hanging_pt",
        "right_pt",
    ):
        if fallback_fmt.get(key) is not None:
            header_fmt[key] = fallback_fmt[key]
    header_fmt["bold"] = True
    if header_fmt.get("italic") is None:
        header_fmt["italic"] = False
    roles["table_header"] = header_fmt
    examples = observations.get(fallback_role, {}).get("examples") or []
    observations["table_header"] = {
        "count": 1,
        "examples": [f"由{ROLE_LABELS.get(fallback_role, fallback_role)}推断表头格式"] + examples[:1],
    }
    return True


def get_enabled_roles(profile: dict[str, Any]) -> set[str]:
    enabled = profile.get("enabled_roles")
    if enabled:
        return set(enabled)
    return set(profile.get("roles", {}).keys())


def make_format_booleans_explicit(profile: dict[str, Any], roles: set[str] | list[str] | None = None) -> None:
    selected_roles = set(roles) if roles is not None else set(profile.get("roles", {}))
    for role in selected_roles:
        fmt = profile.get("roles", {}).get(role)
        if not fmt:
            continue
        for field_name in FORMAT_BOOL_FIELDS:
            if fmt.get(field_name) is None:
                fmt[field_name] = False


def has_no(value: bool | None) -> str:
    return "有" if bool(value) else "没有"


def yes_no(value: bool | None, inherit_word: str = "未特别设置") -> str:
    if value is None:
        return inherit_word
    return "是" if value else "否"


def optional_cell(value: Any) -> Any:
    return 0 if value is None else value


def normalize_text_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def parse_optional_float(value: Any) -> float | None:
    text = normalize_text_cell(value)
    if text in {"", "None", "none", "null", "NULL", "未特别设置"}:
        return None
    if text.startswith("无"):
        return 0.0
    text = text.replace("磅", "").replace("倍", "").replace("英寸", "").replace("%", "").strip()
    return float(text)


def parse_optional_int(value: Any) -> int | None:
    parsed = parse_optional_float(value)
    if parsed is None:
        return None
    return max(0, int(round(parsed)))


def parse_optional_bool(value: Any) -> bool | None:
    text = normalize_text_cell(value).lower()
    if text in {"", "继承", "未特别设置", "none", "null"}:
        return None
    if text in {"是", "有", "yes", "y", "true", "1", "加粗", "斜体"}:
        return True
    if text in {"否", "没有", "无", "no", "n", "false", "0", "不", "不加粗", "不斜体"}:
        return False
    return None


def parse_apply(value: Any) -> bool:
    text = normalize_text_cell(value).lower()
    if text in {"否", "no", "n", "false", "0", "不应用", "停用", "跳过"}:
        return False
    return True


def parse_align(value: Any) -> str:
    text = normalize_text_cell(value)
    if not text:
        return "justify"
    return ALIGN_VALUES.get(text, ALIGN_VALUES.get(text.lower(), "justify"))


def parse_leader(value: Any) -> str:
    text = normalize_text_cell(value)
    if not text:
        return "none"
    return LEADER_VALUES.get(text, LEADER_VALUES.get(text.lower(), "none"))


def parse_table_wrapping(value: Any) -> str | None:
    text = normalize_text_cell(value)
    if not text:
        return None
    return TABLE_WRAPPING_VALUES.get(text, TABLE_WRAPPING_VALUES.get(text.lower(), None))


def parse_table_layout(value: Any) -> str | None:
    text = normalize_text_cell(value)
    if not text:
        return None
    return TABLE_LAYOUT_VALUES.get(text, TABLE_LAYOUT_VALUES.get(text.lower(), None))


def parse_page_number_location(value: Any) -> str:
    text = normalize_text_cell(value)
    if not text:
        return "footer"
    return PAGE_NUMBER_LOCATION_VALUES.get(text, PAGE_NUMBER_LOCATION_VALUES.get(text.lower(), "footer"))


def parse_page_number_format(value: Any) -> str:
    text = normalize_text_cell(value)
    if not text:
        return "decimal"
    return PAGE_NUMBER_FORMAT_VALUES.get(text, PAGE_NUMBER_FORMAT_VALUES.get(text.lower(), "decimal"))


def parse_page_number_section_mode(value: Any) -> str:
    text = normalize_text_cell(value)
    if not text:
        return "restart"
    return PAGE_NUMBER_SECTION_MODE_VALUES.get(text, PAGE_NUMBER_SECTION_MODE_VALUES.get(text.lower(), "restart"))


def parse_color_name_and_hex(name_value: Any, hex_value: Any) -> tuple[str, str]:
    name = normalize_text_cell(name_value)
    value = normalize_text_cell(hex_value).upper()
    if not value or value in {"自动", "AUTO", "AUTOMATIC"}:
        return name or "自动（通常黑色）", "#000000"
    if not value.startswith("#") and re.fullmatch(r"[0-9A-F]{6}", value):
        value = f"#{value}"
    if not re.fullmatch(r"#[0-9A-F]{6}", value):
        return name or "黑色", "#000000"
    return name or COMMON_COLOR_NAMES.get(value, value), value


def examples_for_role(profile: dict[str, Any], role: str) -> str:
    lookup_role = "keywords" if role == "keyword_label" else role
    examples = profile.get("source_observations", {}).get(lookup_role, {}).get("examples", [])
    return " / ".join(examples[:2])


def export_format_table(profile: dict[str, Any], xlsx_path: Path) -> None:
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "格式确认表"
    ws.append(FORMAT_TABLE_COLUMNS)

    ensure_inferred_table_header_role(profile)
    roles = observed_roles_for_table(profile)
    make_format_booleans_explicit(profile, roles)
    profile["enabled_roles"] = roles
    for role in roles:
        fmt = profile["roles"][role]
        ws.append(
            [
                role,
                ROLE_LABELS.get(role, role),
                examples_for_role(profile, role),
                "是",
                fmt.get("font") or profile["latin_font"],
                fmt.get("east_asia_font") or profile["east_asia_font"],
                fmt.get("font_color_name") or "黑色",
                fmt.get("font_color_hex") or "#000000",
                optional_cell(fmt.get("size_pt")),
                has_no(fmt.get("bold")),
                has_no(fmt.get("italic")),
                has_no(fmt.get("all_caps")),
                has_no(fmt.get("small_caps")),
                ALIGN_LABELS.get(fmt.get("align", "justify"), fmt.get("align", "justify")),
                optional_cell(fmt.get("line_spacing")),
                optional_cell(fmt.get("before_pt")),
                optional_cell(fmt.get("after_pt")),
                "" if fmt.get("blank_before") is None else optional_cell(fmt.get("blank_before")),
                "" if fmt.get("blank_after") is None else optional_cell(fmt.get("blank_after")),
                optional_cell(fmt.get("first_line_pt")),
                optional_cell(fmt.get("left_pt")),
                optional_cell(fmt.get("hanging_pt")),
                optional_cell(fmt.get("right_pt")),
                yes_no(fmt.get("toc_page_number_right_aligned"), "否"),
                LEADER_LABELS.get(fmt.get("toc_leader", "none"), fmt.get("toc_leader", "none")),
                optional_cell(fmt.get("toc_right_tab_pt")),
                TABLE_ALIGNMENT_LABELS.get(fmt.get("table_alignment") or "", fmt.get("table_alignment") or ""),
                TABLE_WRAPPING_LABELS.get(fmt.get("table_text_wrapping") or "", fmt.get("table_text_wrapping") or ""),
                TABLE_LAYOUT_LABELS.get(fmt.get("table_layout") or "", fmt.get("table_layout") or ""),
                "" if fmt.get("table_width_pt") is not None else optional_cell(fmt.get("table_width_pct")),
                optional_cell(fmt.get("table_width_pt")),
                optional_cell(fmt.get("cell_margin_top_pt")),
                optional_cell(fmt.get("cell_margin_bottom_pt")),
                optional_cell(fmt.get("cell_margin_left_pt")),
                optional_cell(fmt.get("cell_margin_right_pt")),
                "",
            ]
        )

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    widths = [18, 18, 42, 12, 18, 16, 18, 14, 12, 10, 10, 12, 12, 14, 10, 12, 12, 12, 12, 14, 12, 14, 12, 16, 16, 18, 14, 14, 14, 16, 16, 16, 16, 16, 16, 28]
    for column_index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(column_index)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    yes_no_validation = DataValidation(type="list", formula1='"是,否"', allow_blank=False)
    bool_validation = DataValidation(type="list", formula1='"有,没有"', allow_blank=True)
    align_validation = DataValidation(type="list", formula1='"左对齐,居中,右对齐,两端对齐"', allow_blank=True)
    leader_validation = DataValidation(type="list", formula1='"无,点引导线,短横线,下划线,粗线,中点"', allow_blank=True)
    table_align_validation = DataValidation(type="list", formula1='"左对齐,居中,右对齐"', allow_blank=True)
    table_wrapping_validation = DataValidation(type="list", formula1='"不环绕,环绕"', allow_blank=True)
    table_layout_validation = DataValidation(type="list", formula1='"固定列宽,自动调整"', allow_blank=True)
    ws.add_data_validation(yes_no_validation)
    ws.add_data_validation(bool_validation)
    ws.add_data_validation(align_validation)
    ws.add_data_validation(leader_validation)
    ws.add_data_validation(table_align_validation)
    ws.add_data_validation(table_wrapping_validation)
    ws.add_data_validation(table_layout_validation)
    if ws.max_row >= 2:
        header_letters = {name: get_column_letter(index) for index, name in enumerate(FORMAT_TABLE_COLUMNS, start=1)}
        yes_no_validation.add(f"{header_letters['是否应用']}2:{header_letters['是否应用']}{ws.max_row}")
        bool_validation.add(f"{header_letters['加粗']}2:{header_letters['小型大写']}{ws.max_row}")
        yes_no_validation.add(f"{header_letters['目录页码右对齐']}2:{header_letters['目录页码右对齐']}{ws.max_row}")
        align_validation.add(f"{header_letters['对齐方式']}2:{header_letters['对齐方式']}{ws.max_row}")
        leader_validation.add(f"{header_letters['目录点引导线']}2:{header_letters['目录点引导线']}{ws.max_row}")
        table_align_validation.add(f"{header_letters['表格水平位置']}2:{header_letters['表格水平位置']}{ws.max_row}")
        table_wrapping_validation.add(f"{header_letters['表格文字环绕']}2:{header_letters['表格文字环绕']}{ws.max_row}")
        table_layout_validation.add(f"{header_letters['表格布局方式']}2:{header_letters['表格布局方式']}{ws.max_row}")

    page_ws = wb.create_sheet("页面设置")
    page_ws.append(["项目ID", "项目", "值", "单位", "说明"])
    page = profile.get("page", {})
    header_footer = profile.get("header_footer", {})
    page_rows = [
        ("preserve_target_page_size", "保留目标论文原页面大小", yes_no(page.get("preserve_target_page_size", True), "是"), "", "填“是”则不强制改 A4；填“否”则按 A4。"),
        ("top_in", "上边距", page.get("top_in", 1.0), "英寸", ""),
        ("bottom_in", "下边距", page.get("bottom_in", 1.0), "英寸", ""),
        ("left_in", "左边距", page.get("left_in", 1.25), "英寸", ""),
        ("right_in", "右边距", page.get("right_in", 1.25), "英寸", ""),
        ("apply_header_footer", "应用页眉页脚页码设置", yes_no(header_footer.get("apply_header_footer", True), "是"), "", "填“否”则不调整页眉、页脚、页码。"),
        ("preserve_target_text", "保留目标页眉页脚文字", yes_no(header_footer.get("preserve_target_text", True), "是"), "", "默认保留目标论文原文字，只套范文格式。"),
        ("preserve_target_page_numbers", "保留目标页码", yes_no(header_footer.get("preserve_target_page_numbers", True), "是"), "", "默认不改目标论文已有页码、不插入页码、不改页码编号格式；填“否”才按下列页码设置处理。"),
        ("header_distance_in", "页眉距边界", header_footer.get("header_distance_in", 0.5), "英寸", ""),
        ("footer_distance_in", "页脚距边界", header_footer.get("footer_distance_in", 0.5), "英寸", ""),
        ("different_first_page", "首页不同", yes_no(header_footer.get("different_first_page", False), "否"), "", "控制 Word 的“首页不同”页眉页脚设置。"),
        ("odd_even_pages", "奇偶页不同", yes_no(header_footer.get("odd_even_pages", False), "否"), "", "控制 Word 的“奇偶页不同”页眉页脚设置。"),
        ("page_number_enabled", "应用页码", yes_no(header_footer.get("page_number_enabled", False), "否"), "", "仅当“保留目标页码”为“否”时生效；若目标没有页码，按页码位置插入 PAGE 域。"),
        ("page_number_location", "页码位置", PAGE_NUMBER_LOCATION_LABELS.get(header_footer.get("page_number_location", "footer"), "页脚"), "", "可填：页眉、页脚。"),
        ("page_number_section_mode", "分节页码续前节/重新编号", PAGE_NUMBER_SECTION_MODE_LABELS.get(header_footer.get("page_number_section_mode", "restart"), "重新编号"), "", "可填：续前节、重新编号。"),
        ("front_page_number_format", "摘要/目录页码格式", PAGE_NUMBER_FORMAT_LABELS.get(header_footer.get("front_page_number_format", "upperRoman"), "大写罗马数字(I,II,III)"), "", "用于第 1 节或正文前节。"),
        ("body_page_number_format", "正文页码格式", PAGE_NUMBER_FORMAT_LABELS.get(header_footer.get("body_page_number_format", "decimal"), "阿拉伯数字(1,2,3)"), "", "用于正文节。"),
        ("front_page_number_start", "摘要/目录页码起始值", header_footer.get("front_page_number_start", 1), "", ""),
        ("body_page_number_start", "正文页码起始值", header_footer.get("body_page_number_start", 1), "", ""),
    ]
    for row in page_rows:
        page_ws.append(row)
    for cell in page_ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    page_ws.column_dimensions["A"].width = 28
    page_ws.column_dimensions["B"].width = 28
    page_ws.column_dimensions["C"].width = 28
    page_ws.column_dimensions["D"].width = 10
    page_ws.column_dimensions["E"].width = 48
    page_ws.freeze_panes = "A2"
    wb.save(xlsx_path)


def apply_format_table(profile: dict[str, Any], xlsx_path: Path) -> dict[str, Any]:
    wb = load_workbook(xlsx_path, data_only=True)
    if "格式确认表" not in wb.sheetnames:
        raise ValueError(f"{xlsx_path} does not contain a '格式确认表' sheet.")
    ws = wb["格式确认表"]
    headers = {normalize_text_cell(cell.value): idx for idx, cell in enumerate(ws[1], start=1)}
    required = {"角色ID", "是否应用"}
    missing = required - set(headers)
    if missing:
        raise ValueError(f"Format table is missing required columns: {', '.join(sorted(missing))}")

    field_columns = {
        "font": "西文字体",
        "east_asia_font": "中文字体",
        "font_color_name": "字体颜色名称",
        "font_color_hex": "字体颜色值",
        "size_pt": "字号(磅)",
        "bold": "加粗",
        "italic": "斜体",
        "all_caps": "全部大写",
        "small_caps": "小型大写",
        "align": "对齐方式",
        "line_spacing": "行距",
        "before_pt": "段前(磅)",
        "after_pt": "段后(磅)",
        "blank_before": "段前空行数",
        "blank_after": "段后空行数",
        "first_line_pt": "首行缩进(磅)",
        "left_pt": "左缩进(磅)",
        "hanging_pt": "悬挂缩进(磅)",
        "right_pt": "右缩进(磅)",
        "toc_page_number_right_aligned": "目录页码右对齐",
        "toc_leader": "目录点引导线",
        "toc_right_tab_pt": "目录右侧制表位(磅)",
        "table_alignment": "表格水平位置",
        "table_text_wrapping": "表格文字环绕",
        "table_layout": "表格布局方式",
        "table_width_pct": "表格宽度比例(%)",
        "table_width_pt": "表格宽度(磅)",
        "cell_margin_top_pt": "单元格上边距(磅)",
        "cell_margin_bottom_pt": "单元格下边距(磅)",
        "cell_margin_left_pt": "单元格左边距(磅)",
        "cell_margin_right_pt": "单元格右边距(磅)",
    }
    enabled_roles = []
    for row_idx in range(2, ws.max_row + 1):
        role = normalize_text_cell(ws.cell(row_idx, headers["角色ID"]).value)
        if not role or role not in profile["roles"]:
            continue
        if not parse_apply(ws.cell(row_idx, headers["是否应用"]).value):
            continue
        enabled_roles.append(role)
        fmt = profile["roles"][role]
        for field_name, column_name in field_columns.items():
            col_idx = headers.get(column_name)
            if not col_idx:
                continue
            value = ws.cell(row_idx, col_idx).value
            if field_name in {"font", "east_asia_font"}:
                text = normalize_text_cell(value)
                if text and text != "未特别设置":
                    fmt[field_name] = text
            elif field_name == "font_color_name":
                name, hex_value = parse_color_name_and_hex(value, ws.cell(row_idx, headers.get("字体颜色值", 0)).value if headers.get("字体颜色值") else None)
                fmt["font_color_name"] = name
                fmt["font_color_hex"] = hex_value
            elif field_name == "font_color_hex":
                name, hex_value = parse_color_name_and_hex(ws.cell(row_idx, headers.get("字体颜色名称", 0)).value if headers.get("字体颜色名称") else None, value)
                fmt["font_color_name"] = name
                fmt["font_color_hex"] = hex_value
            elif field_name in {"bold", "italic", "all_caps", "small_caps"}:
                fmt[field_name] = parse_optional_bool(value)
            elif field_name == "toc_page_number_right_aligned":
                fmt[field_name] = parse_apply(value)
            elif field_name == "toc_leader":
                fmt[field_name] = parse_leader(value)
            elif field_name == "table_alignment":
                text = normalize_text_cell(value)
                fmt[field_name] = TABLE_ALIGNMENT_VALUES.get(text, TABLE_ALIGNMENT_VALUES.get(text.lower(), "center")) if text else None
            elif field_name == "table_text_wrapping":
                fmt[field_name] = parse_table_wrapping(value)
            elif field_name == "table_layout":
                fmt[field_name] = parse_table_layout(value)
            elif field_name == "align":
                fmt[field_name] = parse_align(value)
            elif field_name in {"blank_before", "blank_after"}:
                fmt[field_name] = parse_optional_int(value)
            else:
                fmt[field_name] = parse_optional_float(value)

    profile["enabled_roles"] = enabled_roles
    make_format_booleans_explicit(profile, enabled_roles)

    if "页面设置" in wb.sheetnames:
        page_ws = wb["页面设置"]
        header_footer = profile.setdefault("header_footer", {})
        for row in page_ws.iter_rows(min_row=2, values_only=True):
            key, _label, value, *_rest = row
            key = normalize_text_cell(key)
            if not key:
                continue
            if key == "preserve_target_page_size":
                profile["page"][key] = parse_apply(value)
            elif key in {"top_in", "bottom_in", "left_in", "right_in"}:
                parsed = parse_optional_float(value)
                if parsed is not None:
                    profile["page"][key] = parsed
            elif key in {"apply_header_footer", "preserve_target_text", "preserve_target_page_numbers", "different_first_page", "odd_even_pages", "page_number_enabled"}:
                header_footer[key] = parse_apply(value)
            elif key in {"header_distance_in", "footer_distance_in"}:
                parsed = parse_optional_float(value)
                if parsed is not None:
                    header_footer[key] = parsed
            elif key == "page_number_location":
                header_footer[key] = parse_page_number_location(value)
            elif key == "page_number_section_mode":
                header_footer[key] = parse_page_number_section_mode(value)
            elif key in {"front_page_number_format", "body_page_number_format"}:
                header_footer[key] = parse_page_number_format(value)
            elif key in {"front_page_number_start", "body_page_number_start"}:
                parsed = parse_optional_int(value)
                if parsed is not None:
                    header_footer[key] = max(1, parsed)
    return profile


def set_east_asia_font(run, latin: str, east_asia: str) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr, value in {
        "w:ascii": latin,
        "w:hAnsi": latin,
        "w:cs": latin,
        "w:eastAsia": east_asia,
    }.items():
        r_fonts.set(qn(attr), value)


def set_font_color(run, color_hex: str | None) -> None:
    value = normalize_text_cell(color_hex).upper()
    if not value:
        return
    if not value.startswith("#") and re.fullmatch(r"[0-9A-F]{6}", value):
        value = f"#{value}"
    if not re.fullmatch(r"#[0-9A-F]{6}", value):
        return
    run.font.color.rgb = RGBColor.from_string(value[1:])


def set_run_format(run, fmt: dict[str, Any], profile: dict[str, Any], force_bold=None) -> None:
    latin = fmt.get("font") or profile["latin_font"]
    east_asia = fmt.get("east_asia_font") or profile["east_asia_font"]
    set_east_asia_font(run, latin, east_asia)
    set_font_color(run, fmt.get("font_color_hex") or "#000000")
    if fmt.get("size_pt") is not None:
        run.font.size = Pt(float(fmt["size_pt"]))
    if force_bold is not None:
        run.font.bold = force_bold
    elif fmt.get("bold") is not None:
        run.font.bold = bool(fmt["bold"])
    if fmt.get("italic") is not None:
        run.font.italic = bool(fmt["italic"])
    if fmt.get("all_caps") is not None:
        run.font.all_caps = bool(fmt["all_caps"])
    if fmt.get("small_caps") is not None:
        run.font.small_caps = bool(fmt["small_caps"])


def get_or_add_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def set_xml_toggle(r_pr, tag_name: str, value: bool) -> None:
    element = get_or_add_child(r_pr, f"w:{tag_name}")
    element.set(qn("w:val"), "1" if value else "0")


def set_run_element_format(run_element, fmt: dict[str, Any], profile: dict[str, Any], force_bold=None) -> None:
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run_element.insert(0, r_pr)
    latin = fmt.get("font") or profile["latin_font"]
    east_asia = fmt.get("east_asia_font") or profile["east_asia_font"]
    r_fonts = get_or_add_child(r_pr, "w:rFonts")
    for attr, value in {
        "w:ascii": latin,
        "w:hAnsi": latin,
        "w:cs": latin,
        "w:eastAsia": east_asia,
    }.items():
        r_fonts.set(qn(attr), value)
    color_hex = normalize_text_cell(fmt.get("font_color_hex") or "#000000").upper()
    if color_hex.startswith("#"):
        color_hex = color_hex[1:]
    if re.fullmatch(r"[0-9A-F]{6}", color_hex):
        color = get_or_add_child(r_pr, "w:color")
        color.set(qn("w:val"), color_hex)
    if fmt.get("size_pt") is not None:
        size_value = str(int(round(float(fmt["size_pt"]) * 2)))
        get_or_add_child(r_pr, "w:sz").set(qn("w:val"), size_value)
        get_or_add_child(r_pr, "w:szCs").set(qn("w:val"), size_value)
    if force_bold is not None:
        set_xml_toggle(r_pr, "b", bool(force_bold))
        set_xml_toggle(r_pr, "bCs", bool(force_bold))
    elif fmt.get("bold") is not None:
        set_xml_toggle(r_pr, "b", bool(fmt["bold"]))
        set_xml_toggle(r_pr, "bCs", bool(fmt["bold"]))
    if fmt.get("italic") is not None:
        set_xml_toggle(r_pr, "i", bool(fmt["italic"]))
        set_xml_toggle(r_pr, "iCs", bool(fmt["italic"]))
    if fmt.get("all_caps") is not None:
        set_xml_toggle(r_pr, "caps", bool(fmt["all_caps"]))
    if fmt.get("small_caps") is not None:
        set_xml_toggle(r_pr, "smallCaps", bool(fmt["small_caps"]))


def target_signature(fmt: dict[str, Any], profile: dict[str, Any], force_bold=None) -> dict[str, Any]:
    target = {
        "font": fmt.get("font") or profile["latin_font"],
        "east_asia_font": fmt.get("east_asia_font") or profile["east_asia_font"],
        "font_color_name": fmt.get("font_color_name") or "黑色",
        "font_color_hex": fmt.get("font_color_hex") or "#000000",
        "align": fmt.get("align", "justify"),
        "line_spacing": float(fmt.get("line_spacing") or 1.5),
        "before_pt": float(fmt.get("before_pt") or 0),
        "after_pt": float(fmt.get("after_pt") or 0),
        "first_line_pt": float(fmt.get("first_line_pt") or 0),
        "left_pt": float(fmt.get("left_pt") or 0),
        "hanging_pt": float(fmt.get("hanging_pt") or 0),
        "right_pt": float(fmt.get("right_pt") or 0),
    }
    include_toc_controls = bool(fmt.get("toc_page_number_right_aligned")) or bool(fmt.get("toc_right_tab_pt")) or fmt.get("toc_leader") not in {None, "none"}
    if include_toc_controls and fmt.get("toc_page_number_right_aligned") is not None:
        target["toc_page_number_right_aligned"] = bool(fmt.get("toc_page_number_right_aligned"))
    if include_toc_controls and fmt.get("toc_leader") is not None:
        target["toc_leader"] = fmt.get("toc_leader")
    if include_toc_controls and fmt.get("toc_right_tab_pt") is not None:
        target["toc_right_tab_pt"] = float(fmt.get("toc_right_tab_pt") or 0)
    for field_name in ("blank_before", "blank_after"):
        if fmt.get(field_name) is not None:
            target[field_name] = int(fmt.get(field_name) or 0)
    if fmt.get("table_alignment") is not None:
        target["table_alignment"] = fmt.get("table_alignment")
    for field_name in ("table_text_wrapping", "table_layout"):
        if fmt.get(field_name) is not None:
            target[field_name] = fmt.get(field_name)
    for field_name in ("table_width_pct", "table_width_pt", "cell_margin_top_pt", "cell_margin_bottom_pt", "cell_margin_left_pt", "cell_margin_right_pt"):
        if fmt.get(field_name) is not None:
            target[field_name] = float(fmt.get(field_name) or 0)
    if fmt.get("size_pt") is not None:
        target["size_pt"] = float(fmt["size_pt"])
    if force_bold is not None:
        target["bold"] = bool(force_bold)
    elif fmt.get("bold") is not None:
        target["bold"] = bool(fmt["bold"])
    if fmt.get("italic") is not None:
        target["italic"] = bool(fmt["italic"])
    if fmt.get("all_caps") is not None:
        target["all_caps"] = bool(fmt["all_caps"])
    if fmt.get("small_caps") is not None:
        target["small_caps"] = bool(fmt["small_caps"])
    return target


def equivalent_format_value(field_name: str, before: Any, after: Any) -> bool:
    if field_name in {
        "before_pt",
        "after_pt",
        "first_line_pt",
        "left_pt",
        "hanging_pt",
        "right_pt",
        "toc_right_tab_pt",
        "table_width_pct",
        "table_width_pt",
        "cell_margin_top_pt",
        "cell_margin_bottom_pt",
        "cell_margin_left_pt",
        "cell_margin_right_pt",
    }:
        before = 0 if before is None else before
        after = 0 if after is None else after
    if field_name == "font_color_name":
        return True
    if field_name == "font_color_hex":
        before = normalize_text_cell(before).upper()
        after = normalize_text_cell(after).upper()
        return before == after
    if field_name in {"blank_before", "blank_after"}:
        before = 0 if before is None else int(before)
        after = 0 if after is None else int(after)
        return before == after
    if isinstance(before, (int, float)) or isinstance(after, (int, float)):
        if before is None or after is None:
            return before is after
        try:
            return abs(float(before) - float(after)) < 0.05
        except (TypeError, ValueError):
            return before == after
    return before == after


def format_value(field_name: str, value: Any) -> str:
    if value is None:
        if field_name in FORMAT_BOOL_FIELDS:
            return "没有"
        if field_name in {"blank_before", "blank_after"}:
            return "0 行"
        if field_name in {"before_pt", "after_pt", "first_line_pt", "left_pt", "hanging_pt", "right_pt", "toc_right_tab_pt", "table_width_pt", "cell_margin_top_pt", "cell_margin_bottom_pt", "cell_margin_left_pt", "cell_margin_right_pt"}:
            return "0 磅"
        if field_name == "table_width_pct":
            return "0%"
        return "未特别设置"
    if field_name in {"before_pt", "after_pt", "first_line_pt", "left_pt", "hanging_pt", "right_pt", "toc_right_tab_pt", "table_width_pt", "cell_margin_top_pt", "cell_margin_bottom_pt", "cell_margin_left_pt", "cell_margin_right_pt", "size_pt"}:
        return f"{float(value):g} 磅"
    if field_name == "table_width_pct":
        return f"{float(value):g}%"
    if field_name == "line_spacing":
        return f"{float(value):g} 倍"
    if field_name in {"blank_before", "blank_after"}:
        return f"{int(value)} 行"
    if field_name in FORMAT_BOOL_FIELDS:
        return has_no(value)
    if field_name == "align":
        return ALIGN_LABELS.get(str(value), str(value))
    if field_name == "toc_leader":
        return LEADER_LABELS.get(str(value), str(value))
    if field_name == "table_alignment":
        return TABLE_ALIGNMENT_LABELS.get(str(value), str(value))
    if field_name == "table_text_wrapping":
        return TABLE_WRAPPING_LABELS.get(str(value), str(value))
    if field_name == "table_layout":
        return TABLE_LAYOUT_LABELS.get(str(value), str(value))
    return str(value)


def format_changes(before: dict[str, Any], fmt: dict[str, Any], profile: dict[str, Any], force_bold=None) -> list[tuple[str, Any, Any]]:
    target = target_signature(fmt, profile, force_bold=force_bold)
    changes: list[tuple[str, Any, Any]] = []
    for field_name, after in target.items():
        before_value = before.get(field_name)
        if equivalent_format_value(field_name, before_value, after):
            continue
        changes.append((field_name, before_value, after))
    return changes


def format_changes_for_fields(before: dict[str, Any], fmt: dict[str, Any], profile: dict[str, Any], field_names: set[str]) -> list[tuple[str, Any, Any]]:
    target = target_signature(fmt, profile)
    changes: list[tuple[str, Any, Any]] = []
    for field_name in field_names:
        if field_name not in target:
            continue
        before_value = before.get(field_name)
        after = target[field_name]
        if equivalent_format_value(field_name, before_value, after):
            continue
        changes.append((field_name, before_value, after))
    return changes


def add_format_comment(
    doc: Document,
    paragraph,
    role: str,
    changes: list[tuple[str, Any, Any]],
    tracker: CommentTracker,
    extra_notes: list[str] | None = None,
) -> None:
    extra_notes = [note for note in (extra_notes or []) if note]
    if not tracker.should_comment(role, changes, extra_notes):
        return
    if getattr(getattr(paragraph, "part", None), "partname", None) != "/word/document.xml":
        return
    runs = [run for run in paragraph.runs if run.text]
    if not runs:
        return
    role_label = ROLE_LABELS.get(role, role)
    text_parts = []
    if changes:
        parts = []
        for field_name, before, after in changes:
            label = FIELD_LABELS.get(field_name, field_name)
            parts.append(f"{label}：{format_value(field_name, before)} -> {format_value(field_name, after)}")
        text_parts.append(f"已调整格式（{role_label}）。原格式 -> 调整后格式：{'；'.join(parts)}。")
    text_parts.extend(extra_notes)
    text = "".join(text_parts)
    doc.add_comment(runs, text=text, author=tracker.author, initials="FG")
    tracker.mark_commented(role)


def format_paragraph_layout(paragraph, fmt: dict[str, Any], force_no_first_line: bool = False) -> None:
    paragraph.alignment = to_alignment(fmt.get("align", "justify"))
    pf = paragraph.paragraph_format
    pf.line_spacing = float(fmt.get("line_spacing") or 1.5)
    pf.space_before = Pt(float(fmt.get("before_pt") or 0))
    pf.space_after = Pt(float(fmt.get("after_pt") or 0))
    hanging_pt = 0 if force_no_first_line else float(fmt.get("hanging_pt") or 0)
    first_line_pt = 0 if force_no_first_line else float(fmt.get("first_line_pt") or 0)
    pf.first_line_indent = Pt(-hanging_pt) if hanging_pt > 0 else Pt(first_line_pt)
    pf.left_indent = Pt(float(fmt.get("left_pt") or 0))
    pf.right_indent = Pt(float(fmt.get("right_pt") or 0))
    if fmt.get("toc_page_number_right_aligned") and fmt.get("toc_right_tab_pt"):
        pf.tab_stops.clear_all()
        pf.tab_stops.add_tab_stop(
            Pt(float(fmt["toc_right_tab_pt"])),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=to_tab_leader(fmt.get("toc_leader", "dots")),
        )


def keep_paragraph_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def keep_table_row_together(row, with_next: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True
            if with_next:
                paragraph.paragraph_format.keep_with_next = True


def format_paragraph(paragraph, fmt: dict[str, Any], profile: dict[str, Any]) -> None:
    format_paragraph_layout(paragraph, fmt)
    page_field = paragraph_has_page_field(paragraph)
    for run in paragraph.runs:
        if run.text or page_field:
            set_run_format(run, fmt, profile)
    for run_element in paragraph._p.iter(qn("w:r")):
        if run_element_text(run_element).strip() or page_field:
            set_run_element_format(run_element, fmt, profile)


def format_equation_paragraph(paragraph, fmt: dict[str, Any], profile: dict[str, Any], role: str) -> None:
    format_paragraph_layout(paragraph, fmt, force_no_first_line=True)
    if role == "equation":
        return
    for run in paragraph.runs:
        if run.text:
            set_run_format(run, fmt, profile)
    for run_element in paragraph._p.iter(qn("w:r")):
        if run_element_text(run_element).strip():
            set_run_element_format(run_element, fmt, profile)


def format_toc_content_controls(doc: Document, profile: dict[str, Any], enabled_roles: set[str]) -> None:
    for paragraph, role in iter_toc_content_paragraphs(doc):
        if role not in enabled_roles:
            continue
        fmt = profile["roles"][role]
        style_id = TOC_ROLE_TO_STYLE_ID.get(role)
        if style_id:
            set_paragraph_style_id(paragraph._p, style_id)
        format_paragraph(paragraph, fmt, profile)


def validate_toc_content_controls(doc: Document, enabled_roles: set[str] | None = None) -> None:
    issues: list[str] = []
    for paragraph, role in iter_toc_content_paragraphs(doc):
        if enabled_roles is not None and role not in enabled_roles:
            continue
        expected_style = TOC_ROLE_TO_STYLE_ID.get(role)
        actual_style = paragraph_style_id(paragraph._p)
        if expected_style and actual_style.upper() != expected_style:
            text = paragraph_element_text(paragraph._p).strip()
            issues.append(f"{text[:80]}: {actual_style or 'no style'} != {expected_style}")
    if issues:
        raise RuntimeError("TOC structure validation failed:\n" + "\n".join(issues[:20]))


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def story_has_page_number(story) -> bool:
    return any(paragraph_has_page_field(paragraph) for paragraph in story.paragraphs)


def story_target_paragraph(story, prefer_empty: bool = True) -> Paragraph:
    if prefer_empty:
        for paragraph in story.paragraphs:
            if not paragraph.text.strip() and not has_image(paragraph) and not paragraph_has_page_field(paragraph):
                return paragraph
    return story.paragraphs[0] if story.paragraphs else story.add_paragraph()


def page_number_stories(section, location: str, odd_even_pages: bool) -> list[Any]:
    if location == "header":
        stories = [section.header]
        if odd_even_pages:
            stories.append(section.even_page_header)
        return stories
    stories = [section.footer]
    if odd_even_pages:
        stories.append(section.even_page_footer)
    return stories


def inches_or_none(value) -> float | None:
    return round(value.inches, 3) if value is not None else None


def apply_page_numbering_to_sections(doc: Document, header_footer: dict[str, Any], protected_section_indices: set[int] | None = None) -> None:
    if not header_footer.get("page_number_enabled"):
        return
    protected_section_indices = protected_section_indices or set()
    sections = list(doc.sections)
    if not sections:
        return
    mode = header_footer.get("page_number_section_mode", "restart")
    front_format = header_footer.get("front_page_number_format", "upperRoman")
    body_format = header_footer.get("body_page_number_format", "decimal")
    front_start = int(header_footer.get("front_page_number_start", 1) or 1)
    body_start = int(header_footer.get("body_page_number_start", 1) or 1)
    if mode == "restart" and len(sections) >= 2:
        if 0 not in protected_section_indices:
            set_section_page_number_settings(sections[0], front_format, front_start)
        for index, section in enumerate(sections[1:], start=1):
            if index in protected_section_indices:
                continue
            set_section_page_number_settings(section, body_format, body_start if index == 1 else None)
    elif mode == "restart":
        if 0 not in protected_section_indices:
            set_section_page_number_settings(sections[0], body_format, body_start)
    else:
        for index, section in enumerate(sections):
            if index in protected_section_indices:
                continue
            set_section_page_number_settings(section, body_format, None)


def apply_header_footer_profile(doc: Document, profile: dict[str, Any], enabled_roles: set[str], protected_section_indices: set[int] | None = None) -> bool:
    header_footer = profile.get("header_footer", {})
    if not header_footer.get("apply_header_footer", True):
        return False
    protected_section_indices = protected_section_indices or set()
    preserve_page_numbers = header_footer.get("preserve_target_page_numbers", True)
    changed = False
    if not protected_section_indices:
        try:
            target_odd_even = bool(header_footer.get("odd_even_pages", False))
            if bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False)) != target_odd_even:
                changed = True
            doc.settings.odd_and_even_pages_header_footer = target_odd_even
        except Exception:
            pass
    if not preserve_page_numbers:
        page_settings_before = [section_page_number_settings(section) for section in doc.sections]
        apply_page_numbering_to_sections(doc, header_footer, protected_section_indices)
        if page_settings_before != [section_page_number_settings(section) for section in doc.sections]:
            changed = True
    location = header_footer.get("page_number_location", "footer")
    odd_even_pages = bool(header_footer.get("odd_even_pages", False))
    for section_index, section in enumerate(doc.sections):
        if section_index in protected_section_indices:
            continue
        if header_footer.get("header_distance_in") is not None:
            target_header_distance = float(header_footer.get("header_distance_in") or 0.5)
            if not equivalent_format_value("header_distance_in", inches_or_none(section.header_distance), target_header_distance):
                changed = True
            section.header_distance = Inches(target_header_distance)
        if header_footer.get("footer_distance_in") is not None:
            target_footer_distance = float(header_footer.get("footer_distance_in") or 0.5)
            if not equivalent_format_value("footer_distance_in", inches_or_none(section.footer_distance), target_footer_distance):
                changed = True
            section.footer_distance = Inches(target_footer_distance)
        target_different_first = bool(header_footer.get("different_first_page", False))
        if bool(section.different_first_page_header_footer) != target_different_first:
            changed = True
        section.different_first_page_header_footer = target_different_first
        if header_footer.get("page_number_enabled") and not preserve_page_numbers:
            for story in page_number_stories(section, location, odd_even_pages):
                if not story_has_page_number(story):
                    changed = True
                    paragraph = story_target_paragraph(story)
                    add_page_number_field(paragraph)

        for story_type, _story_variant, story in header_footer_stories(section):
            for paragraph in story_content_paragraphs(story):
                if paragraph_has_page_field(paragraph):
                    if preserve_page_numbers:
                        continue
                    role = "page_number"
                elif story_type == "header":
                    role = "header_text" if "header_text" in enabled_roles else "header_overall"
                else:
                    role = "footer_text" if "footer_text" in enabled_roles else "footer_overall"
                if role not in enabled_roles:
                    continue
                fmt = profile["roles"][role]
                before = paragraph_signature(paragraph)
                changes = format_changes(before, fmt, profile)
                if changes:
                    changed = True
                format_paragraph(paragraph, fmt, profile)
    return changed


def paragraph_from_xml(element, parent) -> Paragraph:
    return Paragraph(element, parent)


def set_paragraph_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def clear_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for num_pr in list(p_pr.findall(qn("w:numPr"))):
        p_pr.remove(num_pr)


def previous_blank_siblings(paragraph) -> list[Paragraph]:
    blanks: list[Paragraph] = []
    element = paragraph._p.getprevious()
    while element is not None and element.tag == qn("w:p"):
        candidate = paragraph_from_xml(element, paragraph._parent)
        if not is_removable_blank_paragraph(candidate):
            break
        blanks.append(candidate)
        element = element.getprevious()
    return blanks


def next_blank_siblings(paragraph) -> list[Paragraph]:
    blanks: list[Paragraph] = []
    element = paragraph._p.getnext()
    while element is not None and element.tag == qn("w:p"):
        candidate = paragraph_from_xml(element, paragraph._parent)
        if not is_removable_blank_paragraph(candidate):
            break
        blanks.append(candidate)
        element = element.getnext()
    return blanks


def remove_paragraph(paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_blank_paragraph_before(paragraph) -> None:
    paragraph._p.addprevious(OxmlElement("w:p"))


def insert_blank_paragraph_after(paragraph) -> None:
    paragraph._p.addnext(OxmlElement("w:p"))


def enforce_blank_paragraphs(paragraph, fmt: dict[str, Any]) -> None:
    for direction, getter, inserter in (
        ("blank_before", previous_blank_siblings, insert_blank_paragraph_before),
        ("blank_after", next_blank_siblings, insert_blank_paragraph_after),
    ):
        target = fmt.get(direction)
        if target is None:
            continue
        target = max(0, int(target))
        blanks = getter(paragraph)
        for extra in blanks[target:]:
            remove_paragraph(extra)
        for _ in range(max(0, target - len(blanks))):
            inserter(paragraph)


KEYWORD_LABEL_RE = re.compile(r"^(\s*(?:关键词|Key\s*words|Keywords)\s*[:：]\s*)(.*)$", re.IGNORECASE)


def is_keyword_label_text(text: str) -> bool:
    return bool(re.fullmatch(r"\s*(?:关键词|Key\s*words|Keywords)\s*[:：]\s*", text or "", re.IGNORECASE))


def split_keyword_runs(paragraph) -> None:
    match = KEYWORD_LABEL_RE.match(paragraph.text or "")
    if not match:
        return
    label, content = match.groups()
    if len(paragraph.runs) == 2 and is_keyword_label_text(paragraph.runs[0].text):
        return
    had_page_break = contains_page_break_or_section_break(paragraph)
    paragraph.clear()
    paragraph.add_run(label)
    if content:
        paragraph.add_run(content)
    if had_page_break:
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def format_keywords(paragraph, profile: dict[str, Any], enabled_roles: set[str]) -> None:
    roles = profile["roles"]
    split_keyword_runs(paragraph)
    format_paragraph(paragraph, roles["keywords"], profile)
    for idx, run in enumerate(paragraph.runs):
        if not run.text:
            continue
        is_label = idx == 0 and is_keyword_label_text(run.text)
        label_enabled = "keyword_label" in enabled_roles
        set_run_format(run, roles["keyword_label"] if is_label and label_enabled else roles["keywords"], profile)


def ensure_style_font(style, profile: dict[str, Any], fmt: dict[str, Any]) -> None:
    style.font.name = fmt.get("font") or profile["latin_font"]
    if fmt.get("size_pt") is not None:
        style.font.size = Pt(float(fmt["size_pt"]))
    if fmt.get("bold") is not None:
        style.font.bold = bool(fmt["bold"])
    if fmt.get("italic") is not None:
        style.font.italic = bool(fmt["italic"])
    if fmt.get("all_caps") is not None:
        style.font.all_caps = bool(fmt["all_caps"])
    if fmt.get("small_caps") is not None:
        style.font.small_caps = bool(fmt["small_caps"])
    set_font_color(style, fmt.get("font_color_hex") or "#000000")
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr, value in {
        "w:ascii": fmt.get("font") or profile["latin_font"],
        "w:hAnsi": fmt.get("font") or profile["latin_font"],
        "w:cs": fmt.get("font") or profile["latin_font"],
        "w:eastAsia": fmt.get("east_asia_font") or profile["east_asia_font"],
    }.items():
        r_fonts.set(qn(attr), value)


def pt_to_twips(value: float | None, default_pt: float = 5.0) -> int:
    return int(round(float(default_pt if value is None else value) * 20))


def set_cell_margins(
    table,
    top_pt: float | None = None,
    bottom_pt: float | None = None,
    left_pt: float | None = None,
    right_pt: float | None = None,
) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    values = {
        "top": pt_to_twips(top_pt, 2.5),
        "bottom": pt_to_twips(bottom_pt, 2.5),
        "left": pt_to_twips(left_pt, 5),
        "right": pt_to_twips(right_pt, 5),
    }
    for side, margin_twips in values.items():
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_table_width_percent(table, width_pct: float | None) -> None:
    if width_pct is None:
        return
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(int(round(float(width_pct) * 50))))


def set_table_width_points(table, width_pt: float | None) -> None:
    if width_pt is None:
        return
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(round(float(width_pt) * 20))))


def clear_table_indent(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)


def clear_table_row_heights(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.trPr
        if tr_pr is None:
            continue
        for height in list(tr_pr.findall(qn("w:trHeight"))):
            tr_pr.remove(height)


def set_table_layout(table, layout: str | None) -> None:
    if layout not in {"fixed", "autofit"}:
        return
    table.autofit = layout == "autofit"
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), layout)


def set_table_text_wrapping(table, wrapping: str | None) -> None:
    if wrapping not in {"none", "around"}:
        return
    tbl_pr = table._tbl.tblPr
    existing = [child for child in tbl_pr if child.tag == qn("w:tblpPr")]
    if wrapping == "none":
        for child in existing:
            tbl_pr.remove(child)
        return
    if existing:
        return
    tblp_pr = OxmlElement("w:tblpPr")
    tblp_pr.set(qn("w:horzAnchor"), "margin")
    tblp_pr.set(qn("w:tblpXSpec"), "center")
    tblp_pr.set(qn("w:vertAnchor"), "text")
    tblp_pr.set(qn("w:leftFromText"), "180")
    tblp_pr.set(qn("w:rightFromText"), "180")
    tbl_pr.append(tblp_pr)


def apply_table_overall(table, fmt: dict[str, Any]) -> None:
    alignment = normalized_table_alignment(fmt.get("table_alignment") or "center")
    table.alignment = to_table_alignment(alignment)
    if alignment == "center":
        clear_table_indent(table)
    set_table_text_wrapping(table, fmt.get("table_text_wrapping") or "none")
    set_table_layout(table, fmt.get("table_layout"))
    if fmt.get("table_width_pt") is not None:
        set_table_width_points(table, fmt.get("table_width_pt"))
    else:
        set_table_width_percent(table, fmt.get("table_width_pct"))
    clear_table_row_heights(table)
    set_cell_margins(
        table,
        top_pt=fmt.get("cell_margin_top_pt"),
        bottom_pt=fmt.get("cell_margin_bottom_pt"),
        left_pt=fmt.get("cell_margin_left_pt"),
        right_pt=fmt.get("cell_margin_right_pt"),
    )


TABLE_FORMAT_ROLES = {
    "table_overall",
    "table_caption",
    "table_header",
    "table_body",
    "table_body_variable",
    "table_body_definition",
    "table_body_value",
    "table_footnote",
    "table_note_cell",
    "table_text",
}
TABLE_BODY_CELL_ROLES = {"table_body_variable", "table_body_definition", "table_body_value"}


def resolve_table_cell_format_role(role: str, enabled_roles: set[str], profile: dict[str, Any]) -> str | None:
    if role in enabled_roles:
        return role
    if role == "table_header" and TABLE_BODY_CELL_ROLES.intersection(enabled_roles):
        return "table_header" if "table_header" in profile.get("roles", {}) else None
    if role in TABLE_BODY_CELL_ROLES and "table_body" in enabled_roles:
        return "table_body"
    if role.startswith("table_") and "table_text" in enabled_roles:
        return "table_text"
    return None


def format_tables(
    doc: Document,
    profile: dict[str, Any],
    skip_first_table: bool,
    comment_tracker: CommentTracker,
    protected_table_elements: set[Any] | None = None,
) -> None:
    enabled_roles = get_enabled_roles(profile)
    if not enabled_roles.intersection(TABLE_FORMAT_ROLES):
        return
    protected_table_elements = protected_table_elements or set()
    roles = profile["roles"]
    for index, table in enumerate(doc.tables):
        if table._tbl in protected_table_elements:
            continue
        if should_skip_table_for_formatting(table, index, skip_first_table):
            continue
        if "table_overall" in enabled_roles:
            before_overall = table_signature(table)
            overall_fmt = roles["table_overall"]
            changes = format_changes_for_fields(
                before_overall,
                overall_fmt,
                profile,
                {
                    "table_alignment",
                    "table_text_wrapping",
                    "table_layout",
                    "table_width_pct",
                    "table_width_pt",
                    "cell_margin_top_pt",
                    "cell_margin_bottom_pt",
                    "cell_margin_left_pt",
                    "cell_margin_right_pt",
                },
            )
            apply_table_overall(table, overall_fmt)
            first_paragraph = next((cell.paragraphs[0] for row in table.rows for cell in row.cells if cell.paragraphs), None)
            if first_paragraph is not None:
                add_format_comment(doc, first_paragraph, "table_overall", changes, comment_tracker)
        for row_idx, row in enumerate(table.rows):
            keep_table_row_together(row, with_next=(row_idx == 0 and len(table.rows) > 1))
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    role = table_cell_role(row_idx, col_idx, text)
                    role = resolve_table_cell_format_role(role, enabled_roles, profile)
                    if role is None:
                        continue
                    fmt = roles[role]
                    before = paragraph_signature(paragraph)
                    changes = format_changes(before, fmt, profile)
                    format_paragraph(paragraph, fmt, profile)
                    add_format_comment(doc, paragraph, role, changes, comment_tracker)


def validate_captioned_table_formatting(
    doc: Document,
    profile: dict[str, Any],
    enabled_roles: set[str],
    protected_table_elements: set[Any] | None = None,
    max_cells_per_table: int = 30,
) -> None:
    if not enabled_roles.intersection(TABLE_FORMAT_ROLES):
        return
    protected_table_elements = protected_table_elements or set()
    important_fields = {"size_pt", "line_spacing", "align", "bold", "italic"}
    roles = profile.get("roles", {})
    for table_index, table in enumerate(doc.tables):
        if table._tbl in protected_table_elements or not has_preceding_table_caption(table):
            continue
        checked = 0
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    role = resolve_table_cell_format_role(table_cell_role(row_idx, col_idx, text), enabled_roles, profile)
                    if role is None or role not in roles:
                        continue
                    mismatches = [
                        change
                        for change in format_changes(paragraph_signature(paragraph), roles[role], profile)
                        if change[0] in important_fields
                    ]
                    if mismatches:
                        detail = "；".join(
                            f"{FIELD_LABELS.get(field, field)} {format_value(field, before)} != {format_value(field, after)}"
                            for field, before, after in mismatches[:4]
                        )
                        raise RuntimeError(
                            f"Captioned table formatting QA failed at table {table_index + 1}, "
                            f"row {row_idx + 1}, column {col_idx + 1}: {detail}"
                        )
                    checked += 1
                    if checked >= max_cells_per_table:
                        break
                if checked >= max_cells_per_table:
                    break
            if checked >= max_cells_per_table:
                break


def set_image_extents(paragraph, width_pt: float | None) -> None:
    if width_pt is None:
        return
    new_cx = int(round(float(width_pt) * 12700))
    wp_extents = [element for element in paragraph._p.iter() if element.tag == qn("wp:extent")]
    a_extents = [element for element in paragraph._p.iter() if element.tag == qn("a:ext")]
    for extent in wp_extents:
        old_cx = extent.get("cx")
        old_cy = extent.get("cy")
        if not old_cx or not old_cy:
            continue
        try:
            ratio = int(old_cy) / int(old_cx)
        except (ValueError, ZeroDivisionError):
            ratio = None
        extent.set("cx", str(new_cx))
        if ratio:
            extent.set("cy", str(int(round(new_cx * ratio))))
    for extent in a_extents:
        old_cx = extent.get("cx")
        old_cy = extent.get("cy")
        if not old_cx or not old_cy:
            continue
        try:
            ratio = int(old_cy) / int(old_cx)
        except (ValueError, ZeroDivisionError):
            ratio = None
        extent.set("cx", str(new_cx))
        if ratio:
            extent.set("cy", str(int(round(new_cx * ratio))))


def iter_document_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def convert_anchor_images_to_inline(paragraph) -> None:
    anchors = [
        element
        for element in paragraph._p.iter()
        if element.tag == qn("wp:anchor")
    ]
    for anchor in anchors:
        inline = OxmlElement("wp:inline")
        for attr_name in ("distT", "distB", "distL", "distR"):
            value = anchor.get(attr_name)
            if value is not None:
                inline.set(attr_name, value)
        for child in list(anchor):
            local_name = etree.QName(child).localname
            if local_name in {"extent", "effectExtent", "docPr", "cNvGraphicFramePr", "graphic"}:
                inline.append(child)
        parent = anchor.getparent()
        if parent is not None:
            parent.replace(anchor, inline)


def format_images(doc: Document, profile: dict[str, Any], protected_paragraph_elements: set[Any] | None = None) -> None:
    image_fmt = profile.get("image_format") or {}
    if not image_fmt:
        return
    protected_paragraph_elements = protected_paragraph_elements or set()
    image_align = image_fmt.get("image_align", "center")
    width_pt = image_fmt.get("image_width_pt")
    image_layout = image_fmt.get("image_layout")
    for paragraph in iter_document_paragraphs(doc):
        if paragraph._p in protected_paragraph_elements:
            continue
        if not has_image(paragraph):
            continue
        if image_layout == "嵌入式":
            convert_anchor_images_to_inline(paragraph)
        paragraph.alignment = to_alignment(image_align)
        set_image_extents(paragraph, width_pt)


def mark_fields_for_update(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        items = {info.filename: zin.read(info.filename) for info in zin.infolist()}
    settings = "word/settings.xml"
    if settings not in items:
        return
    root = etree.fromstring(items[settings])
    update = root.find(qn("w:updateFields"))
    if update is None:
        update = etree.Element(qn("w:updateFields"))
        root.insert(0, update)
    update.set(qn("w:val"), "true")
    items[settings] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in items.items():
            zout.writestr(name, data)


def apply_profile(
    target_path: Path,
    output_path: Path,
    profile: dict[str, Any],
    skip_first_table: bool,
    comment_mode: str,
    comment_author: str,
    max_comments: int,
    apply_engine: str = "styles",
    quiet: bool = False,
) -> int:
    if target_path.resolve() != output_path.resolve():
        shutil.copy2(target_path, output_path)
    normalize_docx_alignment_values(output_path)

    doc = Document(output_path)
    comment_tracker = CommentTracker(mode=comment_mode, author=comment_author, max_comments=max_comments)
    enabled_roles = get_enabled_roles(profile)
    page = profile["page"]
    flatten_image_only_tables(doc)
    split_captioned_image_paragraphs(doc)
    move_preceding_figure_captions_below_images(doc)
    ensure_contents_starts_on_new_page(doc)
    ensure_body_starts_after_toc_page_break(doc)
    front_matter_protection = detect_front_matter_protection(doc)
    protected_paragraph_elements: set[Any] = front_matter_protection.get("protected_paragraph_elements", set())
    protected_table_elements: set[Any] = front_matter_protection.get("protected_table_elements", set())
    protected_section_indices: set[int] = set(front_matter_protection.get("protected_section_indices", []))
    if front_matter_protection.get("enabled"):
        emit(
            "Protected front matter skipped: "
            f"{len(front_matter_protection.get('protected_paragraph_indices', []))} paragraph(s), "
            f"{len(protected_section_indices)} section(s). "
            f"Formatting starts near paragraph {front_matter_protection.get('start_paragraph_index')}.",
            quiet=quiet,
        )
    for section_index, section in enumerate(doc.sections):
        if section_index in protected_section_indices:
            continue
        if not page.get("preserve_target_page_size", True):
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        section.top_margin = Inches(float(page.get("top_in", 1.0)))
        section.bottom_margin = Inches(float(page.get("bottom_in", 1.0)))
        section.left_margin = Inches(float(page.get("left_in", 1.25)))
        section.right_margin = Inches(float(page.get("right_in", 1.25)))

    for style_name, role in (("Normal", "body"), ("Heading 1", "chapter"), ("Heading 2", "second"), ("Heading 3", "third"), ("toc 1", "toc1"), ("toc 2", "toc2"), ("toc 3", "toc3")):
        if role not in enabled_roles:
            continue
        try:
            ensure_style_font(doc.styles[style_name], profile, profile["roles"][role])
        except KeyError:
            pass
    if apply_engine == "styles":
        ensure_role_styles(doc, profile, enabled_roles)

    header_footer_notice_pending = apply_header_footer_profile(doc, profile, enabled_roles, protected_section_indices)
    formatting_started = False
    in_references = False
    in_ack = False
    last_caption_context: str | None = None
    for paragraph in doc.paragraphs:
        if paragraph._p in protected_paragraph_elements:
            continue
        text = paragraph.text.strip()
        if not formatting_started:
            if is_academic_start_text(text):
                formatting_started = True
            else:
                continue
        if paragraph._p in protected_paragraph_elements:
            continue

        if text in {"References", "参考文献"}:
            in_references = True
            in_ack = False
        elif text in {"Acknowledgements", "Acknowledgments", "致谢", "致  谢"}:
            in_references = False
            in_ack = True

        equation_role = equation_role_for_paragraph(paragraph, text)
        if equation_role:
            role = equation_role
        elif is_table_footnote_text(text) and last_caption_context == "table":
            role = "table_footnote"
        elif text.startswith("Source:") or text.startswith("Note:") or text.startswith("注：") or text.startswith("资料来源"):
            role = f"{last_caption_context}_note" if last_caption_context in {"figure", "table"} else role_for_target(text, paragraph.style.name if paragraph.style else "", in_references, in_ack)
        else:
            role = role_for_target(text, paragraph.style.name if paragraph.style else "", in_references, in_ack)
        if role == "empty":
            continue
        if role == "table_footnote" and role not in enabled_roles and "table_note" in enabled_roles:
            role = "table_note"
        if role not in enabled_roles:
            if role == "front_heading" and header_footer_notice_pending:
                add_format_comment(doc, paragraph, role, [], comment_tracker, extra_notes=[HEADER_FOOTER_NOTICE])
                header_footer_notice_pending = False
            continue
        if role == "figure_caption":
            last_caption_context = "figure"
        elif role == "table_caption":
            last_caption_context = "table"
        elif role in {"chapter", "second", "third", "body", "reference_heading", "ack_heading"}:
            last_caption_context = None
        if role == "chapter":
            paragraph.style = doc.styles["Heading 1"]
        elif role == "second":
            paragraph.style = doc.styles["Heading 2"]
        elif role == "third":
            paragraph.style = doc.styles["Heading 3"]
        elif role == "reference_heading":
            role = "reference_heading"
        elif role == "ack_heading":
            role = "ack_heading"

        if role in HEADING_OUTLINE_LEVEL_BY_ROLE:
            set_paragraph_outline_level(paragraph, HEADING_OUTLINE_LEVEL_BY_ROLE[role])
            clear_paragraph_numbering(paragraph)

        fmt = profile["roles"].get(role, profile["roles"]["body"])
        before = paragraph_signature(paragraph)
        if role in HEADING_BLANK_ROLES:
            before["blank_before"] = len(previous_blank_siblings(paragraph))
            before["blank_after"] = len(next_blank_siblings(paragraph))
        changes = format_changes(before, fmt, profile)
        if role in HEADING_BLANK_ROLES:
            enforce_blank_paragraphs(paragraph, fmt)
        if role == "keywords":
            format_keywords(paragraph, profile, enabled_roles)
        elif role in {"equation", "equation_number", "equation_text_fallback"}:
            format_equation_paragraph(paragraph, fmt, profile, role)
        else:
            if apply_engine == "styles":
                apply_role_style(paragraph, role, doc, profile, enabled_roles)
            format_paragraph(paragraph, fmt, profile)
            if role == "table_caption":
                keep_paragraph_with_next(paragraph)
        extra_notes = []
        if role == "front_heading" and header_footer_notice_pending:
            extra_notes.append(HEADER_FOOTER_NOTICE)
            header_footer_notice_pending = False
        add_format_comment(doc, paragraph, role, changes, comment_tracker, extra_notes=extra_notes)

    format_toc_content_controls(doc, profile, enabled_roles)
    validate_toc_content_controls(doc, enabled_roles)
    format_tables(
        doc,
        profile,
        skip_first_table=skip_first_table,
        comment_tracker=comment_tracker,
        protected_table_elements=protected_table_elements,
    )
    validate_captioned_table_formatting(doc, profile, enabled_roles, protected_table_elements)
    format_images(doc, profile, protected_paragraph_elements=protected_paragraph_elements)
    validate_heading_formatting(doc, profile, enabled_roles)

    doc.save(output_path)
    mark_fields_for_update(output_path)
    return comment_tracker.count


def summarize_profile(profile: dict[str, Any]) -> str:
    lines = []
    roles = profile["roles"]
    summary_roles = [role for role in ROLE_ORDER if role in get_enabled_roles(profile)]
    if not summary_roles:
        summary_roles = ["front_heading", "contents_heading", "toc1", "toc2", "toc3", "chapter", "second", "third", "body", "caption", "reference_entry", "ack_heading"]
    for role in summary_roles:
        if role not in roles:
            continue
        fmt = roles[role]
        lines.append(
            f"- {role}: {fmt.get('size_pt')}pt, color={fmt.get('font_color_hex')}, "
            f"bold={fmt.get('bold')}, italic={fmt.get('italic')}, align={fmt.get('align')}, "
            f"first={fmt.get('first_line_pt')}, left={fmt.get('left_pt')}, hanging={fmt.get('hanging_pt')}"
        )
    return "\n".join(lines)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--sample", type=Path, help="Sample thesis/reference Word file (.doc/.docx/.rtf/.odt). Required when creating a format table.")
    parser.add_argument("--target", type=Path, help="Target Word file (.doc/.docx/.rtf/.odt) to format. Required when applying a confirmed format table.")
    parser.add_argument("--output", type=Path, help="Output DOCX path. Defaults to '<target>_formatted_from_sample.docx'.")
    parser.add_argument("--analysis-dir", type=Path, default=Path(".thesis_format_analysis"), help="Directory for converted sample and JSON profile.")
    parser.add_argument("--profile-json", type=Path, help="Optional explicit path for the JSON profile report.")
    parser.add_argument("--reuse-profile", type=Path, help="Reuse an existing format-profile.json instead of re-analyzing the sample during apply.")
    parser.add_argument("--export-format-table", type=Path, help="Write an editable XLSX format confirmation table inferred from the sample.")
    parser.add_argument("--format-table", type=Path, help="Read a user-confirmed XLSX format table and use it as the source of truth.")
    parser.add_argument("--analyze-only", action="store_true", help="Only infer and write the sample profile/format table; do not format a target.")
    parser.add_argument("--use-sample-page-size", action="store_true", help="Also force target page size to A4 instead of preserving target size.")
    parser.add_argument("--format-cover-table", action="store_true", help="Also format the first table, which is often a cover form and is skipped by default.")
    parser.add_argument(
        "--comment-mode",
        choices=("role", "all", "none"),
        default="role",
        help="Add Word comments explaining formatting changes: one per role (default), every changed paragraph/table cell, or none.",
    )
    parser.add_argument("--comment-author", default="论文格式助手", help="Author name shown on generated format-change comments.")
    parser.add_argument("--max-comments", type=int, default=120, help="Safety cap for generated comments.")
    parser.add_argument(
        "--apply-engine",
        choices=("styles", "direct"),
        default="styles",
        help="Paragraph formatting engine: style-first with direct fallback (default) or the legacy direct-only path.",
    )
    parser.add_argument("--visual-qa-dir", type=Path, help="Directory for mandatory LibreOffice-rendered PNG/PDF visual QA artifacts.")
    parser.add_argument(
        "--shared-sample-cache-dir",
        type=Path,
        default=DEFAULT_SHARED_SAMPLE_CACHE_DIR,
        help="Cross-run cache directory for unchanged sample analysis and sample visual-QA artifacts.",
    )
    parser.add_argument(
        "--no-shared-sample-cache",
        action="store_true",
        help="Disable reuse and publication of cross-run sample analysis artifacts for diagnosis.",
    )
    parser.add_argument(
        "--qa-level",
        choices=("fast", "review", "strict"),
        default="strict",
        help="Visual QA depth: fast skips rendering, review renders lighter contact-sheet QA, strict keeps the full visual gate.",
    )
    parser.add_argument("--quiet", action="store_true", help="Print only high-signal paths and final run summary; detailed artifacts are written to files.")
    parser.add_argument("--render-width", type=int, default=1800, help="Maximum rendered page width for visual QA PNGs.")
    parser.add_argument("--render-height", type=int, default=2600, help="Maximum rendered page height for visual QA PNGs.")
    parser.add_argument(
        "--cleanup-after-delivery",
        action="store_true",
        help=(
            "After a successful strict Stage 2 delivery, delete the confirmed format table and generated working "
            "analysis files while retaining source documents, final DOCX, and final visual-QA artifacts."
        ),
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    operation = "analyze" if args.analyze_only else "apply"
    tracker = PerformanceTracker(operation, args.qa_level)
    tracker.set_metric("pdf_conversion_passes", 0)
    tracker.set_metric("png_raster_passes", 0)
    if args.cleanup_after_delivery and args.qa_level != "strict":
        raise RuntimeError("--cleanup-after-delivery is allowed only with --qa-level strict final delivery.")
    args.analysis_dir.mkdir(parents=True, exist_ok=True)
    visual_qa_dir = (args.visual_qa_dir or args.analysis_dir / "visual-qa").resolve()
    policy = qa_policy(args.qa_level, args.render_width, args.render_height)
    render_width = int(policy["render_width"])
    render_height = int(policy["render_height"])
    tracker.set_metric("render_enabled", bool(policy["render_enabled"]))
    input_paths = [path for path in (args.sample, args.target) if path is not None]
    with tracker.measure("startup_validation"):
        if run_requires_libreoffice(args.qa_level, input_paths):
            soffice: Path | None = require_libreoffice()
            tracker.set_metric("libreoffice_required", True)
            tracker.set_metric("libreoffice_available", True)
        else:
            soffice = find_soffice()
            tracker.set_metric("libreoffice_required", False)
            tracker.set_metric("libreoffice_available", soffice is not None)
    if soffice is not None:
        emit(f"LibreOffice: {soffice}", quiet=args.quiet)
    else:
        emit(libreoffice_optional_message(input_paths), quiet=args.quiet, important=True)
    emit(
        f"QA level: {args.qa_level}"
        + (" (visual rendering skipped)" if not policy["render_enabled"] else f" ({render_width}x{render_height})"),
        quiet=args.quiet,
    )
    if not args.sample and not args.format_table:
        print("Provide --sample to create a format table, or --format-table to apply a confirmed table.", file=sys.stderr)
        return 2

    profile = default_profile()
    if args.reuse_profile:
        profile = load_profile_json(args.reuse_profile.resolve())
        emit(f"Reused profile JSON: {args.reuse_profile}", quiet=args.quiet, important=True)
    sample_label = args.sample.stem if args.sample else (args.reuse_profile.stem if args.reuse_profile else args.format_table.stem)
    profile_json = args.profile_json or args.analysis_dir / f"{sample_label}.format-profile.json"
    applying_confirmed_table = bool(args.sample and args.format_table and args.target and not args.analyze_only and args.export_format_table is None)
    sample_docx = None
    sample_report = None
    shared_sample_hit = False
    if args.sample:
        sample_render_dir = visual_qa_dir / f"{args.sample.stem}_sample"
        with tracker.measure("sample_cache_lookup"):
            cached_sample = load_cached_sample_analysis(
                args.sample.resolve(),
                args.analysis_dir.resolve(),
                profile_json.resolve(),
                sample_render_dir,
                render_width,
                render_height,
                args.use_sample_page_size,
                args.qa_level,
            )
        tracker.set_metric("local_sample_cache_hit", cached_sample is not None)
        if cached_sample is None and not args.no_shared_sample_cache:
            with tracker.measure("shared_sample_cache_lookup"):
                cached_sample = load_shared_sample_analysis(
                    args.sample.resolve(),
                    args.shared_sample_cache_dir.resolve(),
                    render_width,
                    render_height,
                    args.use_sample_page_size,
                    args.qa_level,
                )
            shared_sample_hit = cached_sample is not None
        tracker.set_metric("sample_cache_hit", cached_sample is not None)
        tracker.set_metric("shared_sample_cache_hit", shared_sample_hit)
        if cached_sample is not None:
            profile = cached_sample["profile"]
            sample_docx = Path(cached_sample["sample_docx"])
            sample_report = cached_sample["sample_report"]
            supplemented = []
            cache_dir = (
                shared_sample_cache_entry_dir(
                    args.shared_sample_cache_dir.resolve(),
                    args.sample.resolve(),
                    render_width,
                    render_height,
                    args.use_sample_page_size,
                    args.qa_level,
                )
                if shared_sample_hit
                else args.analysis_dir.resolve()
            )
            emit(f"Sample cache hit: {sample_cache_manifest_path(cache_dir)}", quiet=args.quiet)
        else:
            with tracker.measure("sample_conversion"):
                sample_docx = convert_legacy_word(args.sample.resolve(), args.analysis_dir.resolve(), soffice)
            with tracker.measure("sample_analysis"):
                legacy_source = args.sample.suffix.lower() in LEGACY_WORD_SUFFIXES
                profile = update_profile_from_observations(profile, sample_docx, legacy_source=legacy_source)
                supplemented = [] if applying_confirmed_table else supplement_missing_legacy_roles(profile, args.sample.resolve(), args.analysis_dir.resolve())
        if not applying_confirmed_table:
            if sample_report is None:
                with tracker.measure("sample_visual_qa"):
                    if policy["render_enabled"]:
                        sample_report = render_for_visual_qa(
                            sample_docx.resolve(),
                            sample_render_dir,
                            soffice,
                            "sample",
                            render_width,
                            render_height,
                            qa_level=args.qa_level,
                            edge_overflow_is_error=bool(policy["edge_overflow_is_error"]),
                            contact_sheet_max_pages=int(policy["contact_sheet_max_pages"]),
                        )
                        tracker.increment("pdf_conversion_passes")
                        tracker.increment("png_raster_passes")
                    else:
                        sample_report = write_fast_visual_qa_report(sample_docx.resolve(), sample_render_dir, soffice, "sample")
            tracker.set_metric("sample_rendered_pages", int(sample_report.get("page_count", 0)))
        if args.use_sample_page_size:
            profile["page"]["preserve_target_page_size"] = False
        emit(f"Sample analyzed: {args.sample}", quiet=args.quiet)
        if supplemented:
            emit("Legacy structural supplements: " + ", ".join(supplemented), quiet=args.quiet)
        if sample_report is not None:
            emit(
                f"Sample visual QA: {sample_report['render_dir']} ({sample_report['page_count']} pages, {args.qa_level})",
                quiet=args.quiet,
                important=args.analyze_only,
            )
        else:
            emit("Sample visual QA skipped during apply because a confirmed format table is already loaded.", quiet=args.quiet)

    if args.format_table:
        with tracker.measure("format_table_load"):
            profile = apply_format_table(profile, args.format_table.resolve())
        emit(f"Confirmed format table loaded: {args.format_table}", quiet=args.quiet)

    format_table = args.export_format_table
    if args.sample and (args.analyze_only or format_table is not None):
        format_table = format_table or args.analysis_dir / f"{args.sample.stem}_格式确认表.xlsx"
        with tracker.measure("format_table_export"):
            export_format_table(profile, format_table)
        emit(f"Editable format table: {format_table}", quiet=args.quiet, important=True)

    with tracker.measure("profile_and_lint_write"):
        profile_json.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
        lint_txt, lint_json, lint_issues = write_profile_lint_report(profile, args.analysis_dir.resolve())
    if args.sample and not args.format_table and sample_docx is not None and not shared_sample_hit:
        visual_report = None
        if sample_report and sample_report.get("render_dir"):
            visual_report = Path(sample_report["render_dir"]) / "visual-report.json"
        write_sample_cache_manifest(
            args.sample.resolve(),
            args.analysis_dir.resolve(),
            sample_docx.resolve(),
            profile_json.resolve(),
            visual_report,
            render_width,
            render_height,
            args.use_sample_page_size,
            args.qa_level,
        )
        if not args.no_shared_sample_cache:
            with tracker.measure("shared_sample_cache_publish"):
                publish_shared_sample_analysis(
                    args.sample.resolve(),
                    args.shared_sample_cache_dir.resolve(),
                    sample_docx.resolve(),
                    profile_json.resolve(),
                    Path(sample_report["render_dir"]) if sample_report else sample_render_dir,
                    render_width,
                    render_height,
                    args.use_sample_page_size,
                    args.qa_level,
                )

    emit(f"Profile JSON: {profile_json}", quiet=args.quiet, important=True)
    emit(f"Profile lint report: {lint_txt} ({len(lint_issues)} issue(s))", quiet=args.quiet, important=True)
    if not args.quiet:
        print(summarize_profile(profile))

    if args.analyze_only:
        performance_path = tracker.write(args.analysis_dir.resolve())
        emit(f"Performance report: {performance_path}", quiet=args.quiet, important=True)
        return 0
    if not args.target:
        print("--target is required unless --analyze-only is set", file=sys.stderr)
        return 2
    with tracker.measure("target_conversion"):
        target_docx = convert_legacy_word(args.target.resolve(), args.analysis_dir.resolve(), soffice)
    output = args.output or args.target.with_name(f"{args.target.stem}_formatted_from_sample.docx")
    with tracker.measure("format_application"):
        comment_count = apply_profile(
            target_docx.resolve(),
            output.resolve(),
            profile,
            skip_first_table=not args.format_cover_table,
            comment_mode=args.comment_mode,
            comment_author=args.comment_author,
            max_comments=args.max_comments,
            apply_engine=args.apply_engine,
            quiet=args.quiet,
        )
    with tracker.measure("output_docx_validation"):
        validate_docx_output(output.resolve(), comment_count, args.comment_mode)
    output_render_dir = visual_qa_dir / f"{output.stem}_output"
    if policy["render_enabled"]:
        output_report, toc_validation_report = render_output_with_toc_validation(
            output.resolve(),
            output_render_dir,
            soffice,
            policy,
            args.qa_level,
            comment_count,
            args.comment_mode,
            tracker=tracker,
        )
        emit(
            f"TOC page results validated: {toc_validation_report['entry_count']} entry/entries; "
            f"refreshed {toc_validation_report['refreshed_count']}.",
            quiet=args.quiet,
            important=True,
        )
    else:
        with tracker.measure("output_fast_qa_report"):
            output_report = write_fast_visual_qa_report(output.resolve(), output_render_dir, soffice, "formatted-output")
        tracker.set_metric("output_rendered_pages", 0)
        emit("TOC displayed-page validation skipped in fast QA; rerun strict before final delivery.", quiet=args.quiet)
    emit(f"Formatted copy: {output}", quiet=args.quiet, important=True)
    emit(f"Format-change comments: {comment_count} ({args.comment_mode})", quiet=args.quiet, important=True)
    emit(f"Output visual QA: {output_report['render_dir']} ({output_report['page_count']} pages, {args.qa_level})", quiet=args.quiet, important=True)
    emit("Visual QA report: " + str(Path(output_report["render_dir"]) / "visual-report.txt"), quiet=args.quiet, important=True)
    if output_report.get("contact_sheet"):
        emit("Contact sheet: " + str(output_report["contact_sheet"]), quiet=args.quiet, important=True)
    if args.cleanup_after_delivery:
        with tracker.measure("delivery_cleanup"):
            removed = cleanup_after_successful_delivery(
                delivery_cleanup_candidates(args, target_docx, profile_json, lint_txt, lint_json),
                protected=[
                    args.target.resolve(),
                    output.resolve(),
                    Path(output_report["render_dir"]).resolve(),
                ],
            )
        emit(
            f"Delivery cleanup removed {len(removed)} generated working file(s), including the confirmed format table when present.",
            quiet=args.quiet,
            important=True,
        )
    performance_path = tracker.write(args.analysis_dir.resolve())
    emit(f"Performance report: {performance_path}", quiet=args.quiet, important=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
